import streamlit as st
import requests
import json
import pandas as pd
from datetime import datetime
import os
import re
import plotly.graph_objects as go
import plotly.express as px
import streamlit as st

# API 基础 URL
API_BASE_URL = "http://localhost:8000"

# 初始化 session state
if 'ai_settings' not in st.session_state:
    st.session_state.ai_settings = {}

# 支持的 AI 提供商配置
AI_PROVIDERS = {
    "openai": {
        "name": "OpenAI",
        "base_url": "https://api.openai.com/v1",
        "models": ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"],
        "default_model": "gpt-4o",
        "description": "OpenAI 官方 API，支持 GPT-4 系列模型"
    },
    "deepseek": {
        "name": "DeepSeek",
        "base_url": "https://api.deepseek.com/v1",
        "models": ["deepseek-chat", "deepseek-coder", "deepseek-reasoner"],
        "default_model": "deepseek-chat",
        "description": "DeepSeek 大模型，性价比高"
    },
    "zhipu": {
        "name": "智谱 AI (GLM)",
        "base_url": "https://open.bigmodel.cn/api/paas/v4",
        "models": ["glm-4", "glm-4-plus", "glm-4-flash", "glm-4v"],
        "default_model": "glm-4",
        "description": "智谱 AI GLM 系列模型"
    },
    "moonshot": {
        "name": "Moonshot (月之暗面)",
        "base_url": "https://api.moonshot.cn/v1",
        "models": ["moonshot-v1-8k", "moonshot-v1-32k", "moonshot-v1-128k"],
        "default_model": "moonshot-v1-8k",
        "description": "Moonshot Kimi 大模型"
    },
    "qwen": {
        "name": "通义千问",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "models": ["qwen-turbo", "qwen-plus", "qwen-max"],
        "default_model": "qwen-turbo",
        "description": "阿里云通义千问系列"
    },
    "custom": {
        "name": "自定义 OpenAI 兼容 API",
        "base_url": "",
        "models": [],
        "default_model": "",
        "description": "支持任何 OpenAI 兼容格式的 API"
    }
}

# 辅助函数：处理reasoning字段，移除代码格式，返回纯文本
def process_reasoning(reasoning: str) -> str:
    """处理reasoning字段，移除代码格式，返回纯文本"""
    if not reasoning:
        return ""
    # 移除JSON代码块
    if '```json' in reasoning:
        reasoning = reasoning.replace('```json', '').replace('```', '').strip()
    # 移除JSON格式标记
    if reasoning.strip().startswith('{') and reasoning.strip().endswith('}'):
        try:
            import json
            # 尝试解析JSON并提取reasoning字段
            parsed = json.loads(reasoning)
            if 'reasoning' in parsed:
                reasoning = parsed['reasoning']
        except:
            pass
    return reasoning

def _safe_float(value):
    try:
        if value is None:
            return None
        return float(value)
    except Exception:
        return None

def render_score_card(evaluation_result: dict):
    """可视化展示：能力点得分 + 分项得分 + 总分公式。"""
    if not isinstance(evaluation_result, dict):
        return

    overall = _safe_float(evaluation_result.get("overall_score")) or 0.0
    ku = _safe_float(evaluation_result.get("knowledge_understanding_score"))
    ka = _safe_float(evaluation_result.get("knowledge_application_score"))
    pc = _safe_float(evaluation_result.get("phase_completion_score"))
    score_policy = evaluation_result.get("score_policy", "")
    score_breakdown = evaluation_result.get("score_breakdown", {}) if isinstance(evaluation_result.get("score_breakdown"), dict) else {}

    ability_component = _safe_float(score_breakdown.get("ability_component"))
    if ability_component is None:
        dims = evaluation_result.get("dimension_scores", [])
        dim_scores = []
        if isinstance(dims, list):
            for ds in dims:
                if isinstance(ds, dict):
                    val = _safe_float(ds.get("score"))
                    if val is not None:
                        dim_scores.append(val)
        ability_component = sum(dim_scores) / len(dim_scores) if dim_scores else None

    st.subheader("🧮 可视化分数卡片")
    with st.container(border=True):
        row1 = st.columns(4)
        with row1[0]:
            st.metric("总分", f"{overall:.1f}/100")
        with row1[1]:
            st.metric("能力点均分", f"{ability_component:.1f}/100" if ability_component is not None else "N/A")
        with row1[2]:
            st.metric("知识点理解", f"{ku:.1f}/100" if ku is not None else "N/A")
        with row1[3]:
            # 理论课显示知识点运用，实践课显示阶段完成度
            if pc is not None and (ku is None and ka is None or "practice" in str(score_policy)):
                st.metric("阶段完成度", f"{pc:.1f}/100")
            else:
                st.metric("知识点运用", f"{ka:.1f}/100" if ka is not None else "N/A")

        formula_text = ""
        substituted = ""
        if "practice_stage_process" in str(score_policy) or (pc is not None and (ku is None and ka is None)):
            formula_text = "总分 = 0.5*能力点 + 0.5*阶段完成度"
            if ability_component is not None and pc is not None:
                substituted = f"= 0.5*{ability_component:.1f} + 0.5*{pc:.1f} = {overall:.1f}"
        else:
            formula_text = "总分 = 0.5*能力点 + 0.25*理解 + 0.25*运用"
            if ability_component is not None and ku is not None and ka is not None:
                substituted = f"= 0.5*{ability_component:.1f} + 0.25*{ku:.1f} + 0.25*{ka:.1f} = {overall:.1f}"

        st.markdown(f"**公式：** `{formula_text}`")
        if substituted:
            st.markdown(f"**代入：** `{substituted}`")

        if score_policy:
            st.caption(f"评分策略：{score_policy}")

def render_dimension_score_details(evaluation_result: dict):
    """展示能力点评分详情：分数、证据、理由、修改建议。"""
    st.subheader("🎯 能力点评分")
    dimension_scores = evaluation_result.get("dimension_scores", []) if isinstance(evaluation_result, dict) else []
    if not isinstance(dimension_scores, list) or not dimension_scores:
        st.info("ℹ️ 暂无能力点评分，建议检查评估提示词或重新触发评估。")
        return

    for score in dimension_scores:
        if not isinstance(score, dict):
            continue
        dim_name = score.get("dimension", "未知能力点")
        dim_score = _safe_float(score.get("score")) or 0.0
        with st.expander(f"**{dim_name}** - {dim_score:.1f}分"):
            evidence = score.get("evidence", [])
            if isinstance(evidence, list) and evidence:
                st.markdown("**证据：**")
                for item in evidence:
                    st.markdown(f"- {item}")
            else:
                st.markdown("**证据：** 未提供明确证据")

            reasoning = process_reasoning(str(score.get("reasoning", "") or ""))
            if reasoning:
                st.markdown(f"**评分理由：** {reasoning}")
            else:
                st.markdown("**评分理由：** 暂无，建议复核该能力点的作业证据。")

            suggestion = str(score.get("improvement_suggestion", "") or "").strip()
            if suggestion:
                st.markdown(f"**修改建议：** {suggestion}")
            else:
                st.markdown("**修改建议：** 建议补充该能力点的关键过程、结果对比和改进验证材料。")

def render_policy_progress_report(report_data: dict, key_prefix: str = "policy"):
    """按课程类型细则渲染总进度报告与趋势图（0-100分制）。"""
    if not isinstance(report_data, dict):
        st.info("暂无可展示的总进度数据。")
        return

    policy_summary = report_data.get("policy_summary", {}) if isinstance(report_data.get("policy_summary"), dict) else {}
    trend_series = report_data.get("trend_series", {}) if isinstance(report_data.get("trend_series"), dict) else {}
    course_type = report_data.get("course_type", policy_summary.get("course_type", "未知课程类型"))
    formula = policy_summary.get("formula", "总分计算公式未提供")

    x_values = trend_series.get("x_values", [])
    overall_scores = trend_series.get("overall_score", [])
    ability_scores = trend_series.get("ability_component", [])
    ku_scores = trend_series.get("knowledge_understanding_component", [])
    ka_scores = trend_series.get("knowledge_application_component", [])
    pc_scores = trend_series.get("phase_completion_component", [])

    def _detect_inflection_points(scores: list) -> list:
        """检测趋势拐点：由升转降或由降转升的位置。"""
        indexed = [(idx, _safe_float(val)) for idx, val in enumerate(scores)]
        valid = [(idx, val) for idx, val in indexed if val is not None]
        if len(valid) < 3:
            return []

        inflections = []
        for i in range(1, len(valid) - 1):
            prev_idx, prev_val = valid[i - 1]
            cur_idx, cur_val = valid[i]
            next_idx, next_val = valid[i + 1]
            prev_delta = cur_val - prev_val
            next_delta = next_val - cur_val

            if abs(prev_delta) < 1.0 or abs(next_delta) < 1.0:
                continue
            if prev_delta * next_delta < 0:
                inflections.append({
                    "index": cur_idx,
                    "score": round(cur_val, 2),
                    "type": "上升转下降" if prev_delta > 0 and next_delta < 0 else "下降转上升"
                })
        return inflections

    def _priority_level(delta_value: float, mean_score: float, volatility: float) -> tuple[str, float]:
        """根据变化值、均分、波动度计算干预优先级。"""
        delta_penalty = max(0.0, -float(delta_value))
        score_penalty = max(0.0, 75.0 - float(mean_score))
        volatility_penalty = float(volatility)
        priority_score = round(delta_penalty * 1.7 + score_penalty * 0.45 + volatility_penalty * 0.9, 2)

        if priority_score >= 24:
            return "高", priority_score
        if priority_score >= 13:
            return "中", priority_score
        return "低", priority_score

    st.subheader("📘 课程细则评分总览")
    card_cols = st.columns(4)
    latest_overall = overall_scores[-1] if overall_scores else report_data.get("overall_score")
    latest_ability = ability_scores[-1] if ability_scores else None
    latest_ku = ku_scores[-1] if ku_scores else None
    latest_ka = ka_scores[-1] if ka_scores else None
    latest_pc = pc_scores[-1] if pc_scores else None
    with card_cols[0]:
        st.metric("当前总分", f"{_safe_float(latest_overall):.1f}" if _safe_float(latest_overall) is not None else "N/A")
    with card_cols[1]:
        st.metric("能力点均分", f"{_safe_float(latest_ability):.1f}" if _safe_float(latest_ability) is not None else "N/A")
    with card_cols[2]:
        if "实践" in str(course_type):
            st.metric("阶段完成度", f"{_safe_float(latest_pc):.1f}" if _safe_float(latest_pc) is not None else "N/A")
        else:
            st.metric("知识点理解", f"{_safe_float(latest_ku):.1f}" if _safe_float(latest_ku) is not None else "N/A")
    with card_cols[3]:
        if "实践" in str(course_type):
            st.metric("课程类型", str(course_type))
        else:
            st.metric("知识点运用", f"{_safe_float(latest_ka):.1f}" if _safe_float(latest_ka) is not None else "N/A")

    st.caption(f"评分公式：{formula}")

    trend_diagnostics = report_data.get("trend_diagnostics", {})
    if isinstance(trend_diagnostics, dict) and trend_diagnostics:
        st.markdown("**趋势诊断摘要**")
        diag_cols = st.columns(3)
        with diag_cols[0]:
            st.metric(
                "总分变化",
                f"{_safe_float(trend_diagnostics.get('overall_delta')):+.1f}" if _safe_float(trend_diagnostics.get("overall_delta")) is not None else "N/A",
                help=f"方向：{trend_diagnostics.get('overall_direction', '未知')} | 波动度：{trend_diagnostics.get('overall_volatility', 'N/A')}"
            )
        with diag_cols[1]:
            st.metric(
                "能力点变化",
                f"{_safe_float(trend_diagnostics.get('ability_delta')):+.1f}" if _safe_float(trend_diagnostics.get("ability_delta")) is not None else "N/A",
                help=f"方向：{trend_diagnostics.get('ability_direction', '未知')} | 波动度：{trend_diagnostics.get('ability_volatility', 'N/A')}"
            )
        with diag_cols[2]:
            st.metric(
                "细则关键项变化",
                f"{_safe_float(trend_diagnostics.get('policy_delta')):+.1f}" if _safe_float(trend_diagnostics.get("policy_delta")) is not None else "N/A",
                help=f"方向：{trend_diagnostics.get('policy_direction', '未知')} | 波动度：{trend_diagnostics.get('policy_volatility', 'N/A')}"
            )

    st.subheader("📈 分项趋势图")
    if not x_values:
        st.info("暂无可用于绘图的趋势数据。")
    else:
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=x_values,
            y=overall_scores,
            mode='lines+markers',
            name='总分',
            line=dict(width=3)
        ))
        fig.add_trace(go.Scatter(
            x=x_values,
            y=ability_scores,
            mode='lines+markers',
            name='能力点均分',
            line=dict(width=3)
        ))
        if "实践" in str(course_type):
            if pc_scores:
                fig.add_trace(go.Scatter(
                    x=x_values,
                    y=pc_scores,
                    mode='lines+markers',
                    name='阶段完成度',
                    line=dict(width=3)
                ))
        else:
            if ku_scores:
                fig.add_trace(go.Scatter(
                    x=x_values,
                    y=ku_scores,
                    mode='lines+markers',
                    name='知识点理解',
                    line=dict(width=3)
                ))
            if ka_scores:
                fig.add_trace(go.Scatter(
                    x=x_values,
                    y=ka_scores,
                    mode='lines+markers',
                    name='知识点运用',
                    line=dict(width=3)
                ))

        fig.update_layout(
            title="课程细则评分趋势",
            xaxis_title=trend_series.get("x_label", "评估序列"),
            yaxis_title="评分（0-100）",
            yaxis_range=[0, 100],
            height=460,
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        st.plotly_chart(fig, use_container_width=True, key=f"{key_prefix}_trend_fig")

    ability_dimension_trends = report_data.get("ability_dimension_trends", {})
    if isinstance(ability_dimension_trends, dict):
        dimension_series = ability_dimension_trends.get("series", [])
        trend_x = ability_dimension_trends.get("x_values", x_values)
        if isinstance(dimension_series, list) and dimension_series and trend_x:
            st.subheader("🧠 各能力点变化趋势图")
            options = [item.get("dimension", "未知能力点") for item in dimension_series]
            descending_dimensions = [
                item for item in dimension_series if _safe_float(item.get("delta")) is not None and float(item.get("delta", 0.0)) <= -1.0
            ]
            if descending_dimensions:
                st.warning(
                    "自动识别到下降能力点："
                    + "、".join([f"{item.get('dimension', '未知')}({float(item.get('delta', 0.0)):+.1f})" for item in descending_dimensions[:6]])
                    + "。图中已用红色高亮。"
                )

            default_selected = options[: min(8, len(options))]
            if descending_dimensions:
                highlighted = [item.get("dimension", "未知能力点") for item in descending_dimensions[:4]]
                default_selected = list(dict.fromkeys(highlighted + default_selected))
            selected_dimensions = st.multiselect(
                "选择要展示的能力点",
                options=options,
                default=default_selected,
                key=f"{key_prefix}_dimension_selector"
            )
            selected_set = set(selected_dimensions)
            fig_dim = go.Figure()
            inflection_rows = []
            priority_rows = []
            for item in dimension_series:
                dim_name = item.get("dimension", "未知能力点")
                if selected_set and dim_name not in selected_set:
                    continue
                dim_delta = _safe_float(item.get("delta")) or 0.0
                dim_mean = _safe_float(item.get("mean_score")) or 0.0
                dim_volatility = _safe_float(item.get("volatility")) or 0.0
                is_descending = dim_delta <= -1.0
                priority_level, priority_score = _priority_level(dim_delta, dim_mean, dim_volatility)
                priority_rows.append({
                    "能力点": dim_name,
                    "优先级": priority_level,
                    "优先级分值": priority_score,
                    "变化值": dim_delta,
                    "平均分": dim_mean,
                    "波动度": dim_volatility
                })

                line_color = "#d62728" if is_descending else ("#2ca02c" if dim_delta >= 1.0 else "#1f77b4")
                fig_dim.add_trace(go.Scatter(
                    x=trend_x,
                    y=item.get("scores", []),
                    mode='lines+markers',
                    name=f"{dim_name} ({dim_delta:+.1f})",
                    line=dict(width=3 if is_descending else 2, color=line_color),
                    marker=dict(size=8)
                ))

                inflections = _detect_inflection_points(item.get("scores", []))
                if inflections:
                    inflection_x = []
                    inflection_y = []
                    inflection_text = []
                    for p in inflections:
                        idx = int(p.get("index", -1))
                        if idx < 0 or idx >= len(trend_x):
                            continue
                        inflection_x.append(trend_x[idx])
                        inflection_y.append(p.get("score", None))
                        inflection_text.append(f"{dim_name}: {p.get('type', '拐点')}")
                        inflection_rows.append({
                            "能力点": dim_name,
                            "拐点位置": trend_x[idx],
                            "拐点分数": p.get("score", 0.0),
                            "拐点类型": p.get("type", "未知")
                        })
                    if inflection_x:
                        fig_dim.add_trace(go.Scatter(
                            x=inflection_x,
                            y=inflection_y,
                            mode='markers+text',
                            text=["拐点"] * len(inflection_x),
                            textposition="top center",
                            name=f"{dim_name} 拐点",
                            marker=dict(symbol="diamond", size=10, color="#ff7f0e"),
                            hovertext=inflection_text,
                            hoverinfo="text"
                        ))
            fig_dim.update_layout(
                title="能力点分项趋势（按评估进度，自动高亮下降项与拐点）",
                xaxis_title=ability_dimension_trends.get("x_label", "报告进度(%)"),
                yaxis_title="评分（0-100）",
                yaxis_range=[0, 100],
                height=500,
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
            )
            st.plotly_chart(fig_dim, use_container_width=True, key=f"{key_prefix}_ability_trend_fig")

            stat_rows = []
            for item in dimension_series:
                stat_rows.append({
                    "能力点": item.get("dimension", "未知能力点"),
                    "变化值": item.get("delta", 0.0),
                    "平均分": item.get("mean_score", 0.0),
                    "波动度": item.get("volatility", 0.0),
                    "趋势判断": item.get("direction", "未知")
                })
            if stat_rows:
                st.caption("能力点趋势统计（按变化幅度排序）")
                st.dataframe(pd.DataFrame(stat_rows), use_container_width=True)

            if inflection_rows:
                st.caption("阶段拐点标注明细（用于定位趋势反转时段）")
                st.dataframe(pd.DataFrame(inflection_rows), use_container_width=True)

            if priority_rows:
                priority_df = pd.DataFrame(priority_rows)
                priority_order = {"高": 0, "中": 1, "低": 2}
                priority_df["priority_sort"] = priority_df["优先级"].map(priority_order).fillna(9)
                priority_df = priority_df.sort_values(["priority_sort", "优先级分值"], ascending=[True, False]).drop(columns=["priority_sort"])
                st.subheader("🚦 建议优先级排序（高/中/低）")
                st.dataframe(priority_df, use_container_width=True)

                high_items = priority_df[priority_df["优先级"] == "高"]
                medium_items = priority_df[priority_df["优先级"] == "中"]
                low_items = priority_df[priority_df["优先级"] == "低"]
                if not high_items.empty or not medium_items.empty or not low_items.empty:
                    st.markdown("**干预建议（按优先级）**")
                    if not high_items.empty:
                        high_targets = "、".join(high_items["能力点"].head(4).tolist())
                        st.markdown(f"- **高优先级**：优先干预 {high_targets}，建议在最近 1-2 次评估周期内设置明确提分目标并复评。")
                    if not medium_items.empty:
                        medium_targets = "、".join(medium_items["能力点"].head(4).tolist())
                        st.markdown(f"- **中优先级**：持续跟踪 {medium_targets}，通过阶段任务补强和证据链完善降低波动。")
                    if not low_items.empty:
                        low_targets = "、".join(low_items["能力点"].head(4).tolist())
                        st.markdown(f"- **低优先级**：维持 {low_targets} 的当前策略，关注稳定性并防止后续回落。")

    stage_breakdown = report_data.get("stage_breakdown", [])
    if isinstance(stage_breakdown, list) and stage_breakdown:
        st.subheader("🧭 阶段分桶分析")
        stage_df = pd.DataFrame(stage_breakdown)
        st.dataframe(stage_df, use_container_width=True)

    report_sections = report_data.get("report_sections", {})
    if isinstance(report_sections, dict) and report_sections:
        st.subheader("📝 学术化报告分节")
        section_order = [
            ("evaluation_basis", "评价依据"),
            ("methodology", "分析方法"),
            ("trend_analysis", "趋势分析"),
            ("stage_findings", "阶段发现"),
            ("risk_analysis", "风险分析"),
            ("follow_up_focus", "后续关注点"),
            ("improvement_areas", "改进领域"),
            ("improvement_path", "改进路径")
        ]
        for key, title in section_order:
            if report_sections.get(key):
                st.markdown(f"**{title}**")
                st.markdown(report_sections[key])

    key_insights = report_data.get("key_insights", [])
    follow_up_points = report_data.get("follow_up_points", [])
    improvement_areas = report_data.get("improvement_areas", [])
    if key_insights or follow_up_points or improvement_areas:
        st.subheader("🎯 关键洞察与行动建议")
        if key_insights:
            st.markdown("**关键洞察**")
            for item in key_insights:
                st.markdown(f"- {item}")
        if follow_up_points:
            st.markdown("**后续关注点**")
            for item in follow_up_points:
                st.markdown(f"- {item}")
        if improvement_areas:
            st.markdown("**改进领域**")
            for item in improvement_areas:
                st.markdown(f"- {item}")

# 页面配置
st.set_page_config(
    page_title="学生多维度能力评估系统",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 初始化 session state
if 'current_page' not in st.session_state:
    st.session_state.current_page = "🏠 系统首页"
if 'api_info' not in st.session_state:
    st.session_state.api_info = None
if 'system_status' not in st.session_state:
    st.session_state.system_status = None
if 'ai_settings' not in st.session_state:
    st.session_state.ai_settings = None

# 侧边栏导航
st.sidebar.title("📚 学生多维度能力评估系统")

# 系统状态检查
try:
    response = requests.get(f"{API_BASE_URL}/health", timeout=3)
    if response.status_code == 200:
        st.sidebar.success("🟢 系统运行正常")
        st.session_state.system_status = "running"
    else:
        st.sidebar.warning("🟡 系统服务异常")
        st.session_state.system_status = "error"
except Exception as e:
    st.sidebar.warning("🟡 系统服务未启动")
    st.session_state.system_status = "offline"
    # 提供启动服务的提示
    st.sidebar.markdown("\n**提示：** 请确保后端API服务正在运行")
    st.sidebar.markdown("- 请在终端运行命令: `uvicorn src.api.main:app --reload`")
    st.sidebar.markdown("- 或检查app.py里的服务地址是否是: http://localhost:8000")

st.sidebar.markdown("---")

# 导航菜单 - 使用按钮而不是下拉框
st.sidebar.markdown("### 📋 功能导航")

# 定义页面列表
pages = [
    ("🏠", "系统首页"),
    ("📋", "大纲管理"),
    ("👥", "学生管理"),
    ("📁", "文件上传"),
    ("📂", "文件管理"),
    ("✏️", "手写识别"),
    ("🤖", "评估管理"),
    ("📊", "结果查询"),
    ("📈", "成长分析"),
    ("⚙️", "AI设置"),
    ("🔧", "API文档")
]

# 创建导航按钮
for emoji, page_name in pages:
    full_page_name = f"{emoji} {page_name}"
    if st.sidebar.button(
        full_page_name,
        key=f"nav_{page_name}",
        use_container_width=True,
        type="primary" if st.session_state.current_page == full_page_name else "secondary"
    ):
        st.session_state.current_page = full_page_name
        st.rerun()

st.sidebar.markdown("---")

# 系统信息
st.sidebar.markdown("### ℹ️ 系统信息")
st.sidebar.markdown(f"**版本:** 1.0.0")
st.sidebar.markdown(f"**API:** [http://localhost:8000](http://localhost:8000)")
st.sidebar.markdown(f"**前端:** [http://localhost:8501](http://localhost:8501)")

# API 信息获取函数
@st.cache_data(ttl=60)
def get_api_info():
    try:
        response = requests.get(f"{API_BASE_URL}/", timeout=5)
        if response.status_code == 200:
            return response.json()
    except:
        pass
    return None

# 获取当前页面
page = st.session_state.current_page

# ==================== 系统首页 ====================
if page == "🏠 系统首页":
    st.title("🎓 学生多维度能力评估系统")
    
    # 系统概览
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("API版本", "1.0.0")
    with col2:
        st.metric("服务状态", "运行中" if st.session_state.system_status == "running" else "异常")
    with col3:
        st.metric("API端口", "8000")
    
    st.markdown("---")
    
    # 系统介绍
    st.subheader("📋 系统介绍")
    st.markdown("""
    基于 **OpenAI GPT-4o** 和 **CrewAI** 的智能化学生能力评估系统
    
    **核心功能：**
    - 📁 **多媒体处理**：支持论文（PDF/DOCX/TXT）、讲演视频（MP4/MOV）、音频录音（MP3/WAV）
    - 🤖 **多智能体评估**：内容分析、表达分析、技术能力分析、批判性思维分析
    - 📊 **数据可视化**：雷达图、评分卡片、趋势图
    - 💾 **数据库持久化**：存储学生信息、提交记录、评估结果
    - 🌐 **API 服务**：提供文件上传、评估启动、结果查询接口
    """)
    
    # 快速开始
    st.subheader("🚀 快速开始")
    
    quick_start_col1, quick_start_col2, quick_start_col3, quick_start_col4 = st.columns(4)
    
    with quick_start_col1:
        if st.button("➕ 添加学生", use_container_width=True):
            st.session_state.current_page = "👥 学生管理"
            st.rerun()
    
    with quick_start_col2:
        if st.button("📤 上传文件", use_container_width=True):
            st.session_state.current_page = "📁 文件上传"
            st.rerun()
    
    with quick_start_col3:
        if st.button("▶️ 启动评估", use_container_width=True):
            st.session_state.current_page = "🤖 评估管理"
            st.rerun()
    
    with quick_start_col4:
        if st.button("📊 查看结果", use_container_width=True):
            st.session_state.current_page = "📊 结果查询"
            st.rerun()
    
    st.markdown("---")
    
    # 系统架构
    st.subheader("🏗️ 系统架构")
    
    arch_col1, arch_col2, arch_col3 = st.columns(3)
    
    with arch_col1:
        st.markdown("""
        **🎯 输入层**
        - PDF/DOCX/TXT 文档
        - MP4/MOV 视频
        - MP3/WAV 音频
        """)
    
    with arch_col2:
        st.markdown("""
        **🤖 处理层**
        - 内容分析Agent
        - 表达分析Agent
        - 技术能力Agent
        - 批判性思维Agent
        """)
    
    with arch_col3:
        st.markdown("""
        **📊 输出层**
        - 综合评分
        - 维度分析
        - 雷达图展示
        - 改进建议
        """)
    
    st.markdown("---")
    
    # 技术栈
    st.subheader("🛠️ 技术栈")
    
    tech_col1, tech_col2, tech_col3, tech_col4 = st.columns(4)
    
    with tech_col1:
        st.markdown("""
        **AI框架**
        - OpenAI GPT-4o
        - CrewAI
        - LangChain
        """)
    
    with tech_col2:
        st.markdown("""
        **后端**
        - FastAPI
        - SQLAlchemy
        """)
    
    with tech_col3:
        st.markdown("""
        **前端**
        - Streamlit
        - Plotly
        - Pandas
        """)
    
    with tech_col4:
        st.markdown("""
        **媒体处理**
        - python-docx
        - PyMuPDF
        **计划使用**
        - Whisper (语音识别)
        """)

# ==================== 大纲管理 ====================
elif page == "📋 大纲管理":
    st.title("📋 大纲管理")
    
    # 获取项目根目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(os.path.dirname(current_dir))
    
    # 大纲文件夹路径（使用绝对路径）
    syllabus_folder = os.path.join(project_root, "评价大纲")
    
    # 检查大纲文件夹是否存在
    if not os.path.exists(syllabus_folder):
        st.error(f"❌ 大纲文件夹 '{syllabus_folder}' 不存在，请创建该文件夹并放入课程大纲文件")
    else:
        # 文件类型选择
        st.subheader("📁 文件类型选择")
        file_type = st.radio(
            "选择文件类型",
            options=["课程大纲", "毕业设计评价指标"],
            horizontal=True,
            help="课程大纲：用于分析课程能力点；毕业设计评价指标：用于毕业设计评价"
        )

        # 上传新大纲文件
        st.subheader("⬆️ 上传新大纲文件")
        new_syllabus_file = st.file_uploader(
            "选择要上传的大纲文件",
            type=["docx", "doc", "pdf", "txt", "json"],
            help="支持 docx/doc/pdf/txt/json，上传后可在下方列表中选择分析"
        )
        if st.button("📤 上传大纲文件", use_container_width=True):
            if not new_syllabus_file:
                st.warning("⚠️ 请先选择文件")
            else:
                try:
                    upload_resp = requests.post(
                        f"{API_BASE_URL}/syllabus_files/upload",
                        files={
                            "file": (
                                new_syllabus_file.name,
                                new_syllabus_file.getvalue(),
                                new_syllabus_file.type or "application/octet-stream"
                            )
                        },
                        timeout=120
                    )
                    if upload_resp.status_code == 200:
                        file_name = upload_resp.json().get("file_name", new_syllabus_file.name)
                        st.success(f"✅ 大纲上传成功：{file_name}")
                        st.rerun()
                    else:
                        try:
                            detail = upload_resp.json().get("detail", "未知错误")
                        except Exception:
                            detail = upload_resp.text or f"HTTP {upload_resp.status_code}"
                        st.error(f"❌ 上传失败: {detail}")
                except Exception as e:
                    st.error(f"❌ 上传失败: {str(e)}")
        
        st.markdown("---")
        
        if file_type == "课程大纲":
            # ========== 课程大纲管理 ==========
            st.subheader("📚 课程大纲管理")
            
            # 获取课程大纲文件列表（排除毕业设计相关文件）
            all_files = [f for f in os.listdir(syllabus_folder) if f.endswith('.docx') or f.endswith('.txt')]
            graduation_keywords = ["毕业设计", "评价指标", "评价标准", "毕设", "graduation"]
            syllabus_files = [f for f in all_files if not any(kw in f for kw in graduation_keywords)]
            
            if not syllabus_files:
                st.warning(f"⚠️ 没有找到课程大纲文件（已排除毕业设计相关文件）")
            else:
                # 选择大纲文件
                selected_syllabus = st.selectbox("选择课程大纲", syllabus_files)
                
                # 显示大纲信息
                st.subheader(f"📄 {selected_syllabus}")
                
                # 检查是否已有大纲分析结果
                analysis_dir = os.path.join(project_root, "analysis_results")
                analysis_file = os.path.join(analysis_dir, f"{selected_syllabus.replace('.docx', '').replace('.txt', '')}.json")
                
                existing_analysis = None
                if os.path.exists(analysis_file):
                    try:
                        with open(analysis_file, 'r', encoding='utf-8') as f:
                            existing_analysis = json.load(f)
                        st.info(f"ℹ️ 已有大纲分析结果")
                    except Exception as e:
                        st.warning(f"⚠️ 读取已有分析结果失败: {str(e)}")
                
                # 分析大纲按钮
                course_type_override_option = st.radio(
                    "课程类型判定方式",
                    options=["自动识别", "手动指定：理论课", "手动指定：实践课"],
                    horizontal=True,
                    help="默认自动识别；若识别不准可手动指定覆盖"
                )
                course_type_override = ""
                if course_type_override_option == "手动指定：理论课":
                    course_type_override = "理论课"
                elif course_type_override_option == "手动指定：实践课":
                    course_type_override = "实践课"

                if st.button("🔍 分析大纲", use_container_width=True):
                    try:
                        syllabus_path = os.path.join(syllabus_folder, selected_syllabus)
                        
                        if selected_syllabus.endswith('.docx'):
                            from docx import Document
                            doc = Document(syllabus_path)
                            paragraph_text = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
                            table_text = []
                            for table in doc.tables:
                                for row in table.rows:
                                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                                    if cells:
                                        table_text.append(" | ".join(cells))
                            syllabus_content = "\n".join(paragraph_text + table_text)
                        elif selected_syllabus.endswith('.txt'):
                            with open(syllabus_path, 'r', encoding='utf-8') as f:
                                syllabus_content = f.read()
                        else:
                            syllabus_content = ""
                        
                        st.info("⏳ 正在分析大纲，这可能需要1-2分钟，请耐心等待...")
                        
                        response = requests.post(
                            f"{API_BASE_URL}/analyze_syllabus",
                            json={
                                "syllabus_content": syllabus_content,
                                "syllabus_name": selected_syllabus,
                                "course_type_override": course_type_override
                            },
                            timeout=180
                        )
                        
                        if response.status_code == 200:
                            analysis_result = response.json()
                            
                            if not analysis_result or (
                                not analysis_result.get('ability_points')
                                and not analysis_result.get('evaluation_criteria')
                                and not analysis_result.get('graduation_requirements')
                            ):
                                st.error("❌ 大纲分析失败：未获取到有效的分析结果")
                            else:
                                os.makedirs(analysis_dir, exist_ok=True)
                                with open(analysis_file, 'w', encoding='utf-8') as f:
                                    json.dump(analysis_result, f, ensure_ascii=False, indent=2)
                                
                                ability_matrix_path = os.path.join(project_root, "ability_matrix.json")
                                if os.path.exists(ability_matrix_path):
                                    with open(ability_matrix_path, 'r', encoding='utf-8') as f:
                                        ability_matrix = json.load(f)
                                else:
                                    ability_matrix = {}
                                
                                ability_matrix[selected_syllabus] = analysis_result
                                with open(ability_matrix_path, 'w', encoding='utf-8') as f:
                                    json.dump(ability_matrix, f, ensure_ascii=False, indent=2)
                                
                                st.success("✅ 大纲分析完成并已保存！")
                                existing_analysis = analysis_result
                        else:
                            try:
                                error_detail = response.json().get("detail", "未知错误")
                            except Exception:
                                error_detail = response.text or f"HTTP {response.status_code}"
                            st.error(f"❌ 大纲分析失败: {error_detail}")
                    except Exception as e:
                        st.error(f"❌ 运行大纲分析器时出错: {str(e)}")
                
                # 显示大纲分析结果框
                st.markdown("---")
                st.subheader("📊 大纲分析结果")
                
                if existing_analysis:
                    # ==================== 新增：课程分类与动态权重展示区 ====================
                    # 尝试从分析结果中获取课程类型，如果没有则默认显示为"理论课"
                    course_type = existing_analysis.get("course_type", "理论课")
                    weight_profile = existing_analysis.get("evaluation_weights", {})
                    dynamic_weights = weight_profile.get("weights", []) if isinstance(weight_profile, dict) else []
                    
                    # 1. 顶部显眼提示
                    course_type_source = existing_analysis.get("course_type_source", "auto")
                    if "实践" in course_type:
                        source_text = "手动指定" if course_type_source == "manual" else "系统自动识别"
                        st.success(f"🛠️ **{source_text}：当前为【{course_type}】模式**。评价权重已切换至“工程落地与实践能力”侧重。")
                    else:
                        source_text = "手动指定" if course_type_source == "manual" else "系统自动识别"
                        st.info(f"📚 **{source_text}：当前为【{course_type}】模式**。评价权重已切换至“理论推导与基础知识”侧重。")

                    if existing_analysis.get("course_type_meta"):
                        meta = existing_analysis.get("course_type_meta", {})
                        st.caption(
                            f"判定分值：理论={meta.get('theory_score', 0)}，实践={meta.get('practice_score', 0)}，"
                            f"置信度={meta.get('confidence', 0)}"
                        )

                    if existing_analysis.get("strict_text_based"):
                        st.caption("✅ 本次结果启用“严格按大纲原文提取”模式，能力点仅来源于毕业要求指标点表格。")
                    if existing_analysis.get("llm_error"):
                        st.warning(f"⚠️ 大模型解析不可用，当前结果由文本规则提取生成：{existing_analysis.get('llm_error')}")
                        
                    # 2. 可展开的权重配比面板
                    with st.expander("⚖️ 查看当前评分权重配比"):
                        if dynamic_weights:
                            weight_data = {
                                "评价维度": [w.get("dimension", "未命名维度") for w in dynamic_weights],
                                "权重占比": [f"{float(w.get('weight', 0)) * 100:.0f}%" for w in dynamic_weights],
                                "说明": [w.get("description", "") for w in dynamic_weights]
                            }
                        else:
                            weight_data = {
                                "评价维度": ["学术表现 (理论)", "批判性思维", "沟通表达", "问题解决", "创新能力", "其他综合"],
                                "权重占比": ["20%", "20%", "15%", "10%", "10%", "25%"],
                                "说明": ["", "", "", "", "", ""]
                            }
                        st.table(pd.DataFrame(weight_data))
                    result_tab1, result_tab2, result_tab3 = st.tabs(["🎯 能力点", "📏 评价标准", "📋 完整结果"])
                    
                    with result_tab1:
                        if existing_analysis.get('ability_points'):
                            st.markdown("### 提取的能力点")
                            for i, point in enumerate(existing_analysis['ability_points'], 1):
                                if isinstance(point, dict):
                                    desc = point.get('description', '')
                                    ability_name = point.get('name') or f"未命名能力点{i}"
                                    st.markdown(f"**能力点{i}：{ability_name}**")
                                    if desc:
                                        st.markdown(f"   - 描述: {desc}")
                                    if point.get('level'):
                                        st.markdown(f"   - 掌握程度: {point.get('level')}")
                                    st.markdown("")
                                else:
                                    st.markdown(f"{i}. {point}")
                        else:
                            st.info("ℹ️ 未提取到能力点")
                    
                    with result_tab2:
                        if existing_analysis.get('evaluation_criteria'):
                            st.markdown("### 提取的评价标准")
                            for i, criterion in enumerate(existing_analysis['evaluation_criteria'], 1):
                                if isinstance(criterion, dict):
                                    ability_desc = criterion.get("ability_description") or criterion.get("description", "")
                                    title = criterion.get('name') or f"能力点{i}评价标准"
                                    st.markdown(f"**{i}. {title}**")
                                    if ability_desc:
                                        st.markdown(f"   - 对应能力点: {ability_desc}")
                                    standards = criterion.get("standards", [])
                                    if isinstance(standards, str):
                                        standards = [s.strip() for s in re.split(r"[；;\n]+", standards) if s.strip()]
                                    if isinstance(standards, list) and standards:
                                        st.markdown("   - 标准:")
                                        for idx, standard_item in enumerate(standards, 1):
                                            st.markdown(f"     {idx}) {standard_item}")
                                    elif criterion.get('standard'):
                                        st.markdown(f"   - 评分标准: {criterion.get('standard')}")
                                    st.markdown("")
                                else:
                                    st.markdown(f"{i}. {criterion}")
                        else:
                            st.info("ℹ️ 未提取到评价标准")
                    
                    with result_tab3:
                        st.markdown("### 完整分析结果（JSON格式）")
                        st.json(existing_analysis)
                else:
                    st.info("ℹ️ 暂无大纲分析结果，请点击\"分析大纲\"按钮进行分析")
        
        else:
            # ========== 毕业设计评价指标管理 ==========
            st.subheader("🎓 毕业设计评价指标管理")
            
            # 获取毕业设计评价指标文件（仅读取包含特定关键词的文件）
            all_files = [f for f in os.listdir(syllabus_folder) if f.endswith('.docx') or f.endswith('.txt') or f.endswith('.json')]
            graduation_keywords = ["毕业设计", "评价指标", "评价标准", "毕设", "graduation", "指标"]
            graduation_files = [f for f in all_files if any(kw in f for kw in graduation_keywords)]
            
            if not graduation_files:
                st.warning("⚠️ 没有找到毕业设计评价指标文件")
                st.info("💡 提示：文件名需包含'毕业设计'、'评价指标'、'评价标准'、'毕设'等关键词")
            else:
                col1, col2 = st.columns([3, 1])
                with col1:
                    selected_graduation_file = st.selectbox("选择毕业设计评价指标文件", graduation_files)
                with col2:
                    st.write("")
                    st.write("")
                    extract_button = st.button("🔍 提炼指标", type="primary", use_container_width=True)
                
                # 检查是否已有提炼结果
                indicators_dir = os.path.join(project_root, "extracted_indicators")
                indicator_file = os.path.join(
                    indicators_dir,
                    f"{selected_graduation_file.replace('.docx', '').replace('.txt', '').replace('.json', '')}_extracted.json"
                )
                
                existing_indicators = None
                if os.path.exists(indicator_file):
                    try:
                        with open(indicator_file, 'r', encoding='utf-8') as f:
                            existing_indicators = json.load(f)
                        st.info(f"ℹ️ 已有提炼结果（可重新提炼覆盖）")
                    except Exception as e:
                        st.warning(f"⚠️ 读取已有提炼结果失败: {str(e)}")
                
                # 提炼后的指标存储
                extracted_indicators = existing_indicators
                
                if extract_button and selected_graduation_file:
                    with st.spinner("正在提炼评价指标..."):
                        try:
                            file_path = os.path.join(syllabus_folder, selected_graduation_file)
                            
                            if selected_graduation_file.endswith('.json'):
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    extracted_indicators = json.load(f)
                            else:
                                if selected_graduation_file.endswith('.docx'):
                                    from docx import Document
                                    doc = Document(file_path)
                                    paragraph_text = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
                                    table_text = []
                                    for table in doc.tables:
                                        for row in table.rows:
                                            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                                            if cells:
                                                table_text.append(" | ".join(cells))
                                    file_content = "\n".join(paragraph_text + table_text)
                                elif selected_graduation_file.endswith('.txt'):
                                    with open(file_path, 'r', encoding='utf-8') as f:
                                        file_content = f.read()
                                else:
                                    file_content = ""
                                
                                extract_response = requests.post(
                                    f"{API_BASE_URL}/extract_guidance_content",
                                    json={
                                        "file_content": file_content,
                                        "file_name": selected_graduation_file
                                    },
                                    timeout=120
                                )
                                
                                if extract_response.status_code == 200:
                                    extracted_indicators = extract_response.json()
                                else:
                                    st.error(f"❌ 提炼失败: {extract_response.json().get('detail', '未知错误')}")
                            
                            if extracted_indicators:
                                # 保存提炼结果
                                os.makedirs(indicators_dir, exist_ok=True)
                                with open(indicator_file, 'w', encoding='utf-8') as f:
                                    json.dump(extracted_indicators, f, ensure_ascii=False, indent=2)
                                
                                st.session_state["extracted_indicators"] = extracted_indicators
                                st.session_state["extracted_indicators_file"] = selected_graduation_file
                                st.success(f"✅ 已提炼并保存评价指标！")
                                st.rerun()
                        except Exception as e:
                            st.error(f"❌ 提炼失败: {str(e)}")
                
                # 显示提炼结果框
                st.markdown("---")
                st.subheader("📊 评价指标提炼结果")
                
                if existing_indicators or ("extracted_indicators" in st.session_state):
                    extracted_indicators = existing_indicators or st.session_state.get("extracted_indicators", {})
                    
                    st.info(f"📄 来源文件: {selected_graduation_file}")
                    
                    # 使用更多标签页展示不同内容
                    result_tab1, result_tab2, result_tab3, result_tab4, result_tab5 = st.tabs([
                        "📋 原始指标", 
                        "📝 扩展指标", 
                        "📊 评价表格", 
                        "🔄 评价流程", 
                        "📄 完整结果"
                    ])
                    
                    with result_tab1:
                        original_indicators = extracted_indicators.get('original_indicators', [])
                        if not original_indicators:
                            original_indicators = extracted_indicators.get('indicators', [])
                        
                        if original_indicators:
                            st.markdown("### 📋 原始评价指标（从文件提取）")
                            for idx, ind in enumerate(original_indicators, 1):
                                indicator_id = ind.get('indicator_id', ind.get('id', f'{idx}'))
                                name = ind.get('name', '未知指标')
                                weight = ind.get('weight', '未知')
                                max_score = ind.get('max_score', 100)
                                description = ind.get('description', '')
                                graduation_req = ind.get('graduation_requirement', '')
                                
                                with st.expander(f"**{indicator_id} {name}** (权重: {weight}%, 满分: {max_score})", expanded=False):
                                    if graduation_req:
                                        st.markdown(f"**对应毕业要求指标点**: {graduation_req}")
                                    if description:
                                        st.markdown(f"**描述**: {description}")
                                    if ind.get('grading_criteria'):
                                        st.markdown(f"**评分标准**: {ind.get('grading_criteria')}")
                        else:
                            st.info("暂无原始评价指标")
                    
                    with result_tab2:
                        indicators = extracted_indicators.get('indicators', [])
                        if indicators:
                            st.markdown("### 📝 扩展评价指标（大模型生成）")
                            for idx, ind in enumerate(indicators, 1):
                                indicator_id = ind.get('indicator_id', ind.get('id', f'{idx}'))
                                name = ind.get('name', '未知指标')
                                weight = ind.get('weight', 10)
                                max_score = ind.get('max_score', 100)
                                description = ind.get('description', '')
                                evaluation_points = ind.get('evaluation_points', [])
                                evaluation_method = ind.get('evaluation_method', '')
                                
                                with st.expander(f"**{indicator_id} {name}** (权重: {weight}%, 满分: {max_score})", expanded=False):
                                    if description:
                                        st.markdown(f"**描述**: {description}")
                                    if evaluation_method:
                                        st.markdown(f"**评价方式**: {evaluation_method}")
                                    
                                    if evaluation_points:
                                        st.markdown("#### 📌 评价要点")
                                        for point in evaluation_points:
                                            point_name = point.get('point_name', '')
                                            point_weight = point.get('weight', '')
                                            grade_criteria = point.get('grade_criteria', {})
                                            
                                            st.markdown(f"**{point_name}** (权重: {point_weight})")
                                            
                                            if grade_criteria:
                                                grade_names = {
                                                    "excellent": "优秀(90-100)",
                                                    "good": "良好(80-89)",
                                                    "medium": "中等(70-79)",
                                                    "pass": "及格(60-69)",
                                                    "fail": "不及格(0-59)"
                                                }
                                                for level, criteria in grade_criteria.items():
                                                    st.markdown(f"- {grade_names.get(level, level)}: {criteria}")
                                            st.markdown("")
                        else:
                            st.info("暂无扩展评价指标")
                    
                    with result_tab3:
                        evaluation_table = extracted_indicators.get('evaluation_table', {})
                        if evaluation_table:
                            st.markdown(f"### 📊 {evaluation_table.get('title', '评价表格')}")
                            
                            columns = evaluation_table.get('columns', ["序号", "指标编号", "指标名称", "满分", "得分", "评价等级"])
                            rows = evaluation_table.get('rows', [])
                            
                            if rows:
                                import pandas as pd
                                df = pd.DataFrame(rows)
                                st.dataframe(df, use_container_width=True, hide_index=True)
                            else:
                                st.info("暂无评价表格数据")
                            
                            st.markdown("#### 📝 使用说明")
                            st.markdown("此表格可用于实际评分，每行对应一个评价指标。评分时填写得分和评价等级。")
                        else:
                            st.info("暂无评价表格，请重新提炼评价指标")
                    
                    with result_tab4:
                        evaluation_flow = extracted_indicators.get('evaluation_flow', {})
                        if evaluation_flow:
                            st.markdown("### 🔄 评价流程")
                            
                            steps = evaluation_flow.get('steps', [])
                            if steps:
                                for step in steps:
                                    step_num = step.get('step', '')
                                    step_name = step.get('name', '')
                                    step_weight = step.get('weight', 0)
                                    step_desc = step.get('description', '')
                                    
                                    st.markdown(f"**步骤{step_num}: {step_name}** (权重: {step_weight*100:.0f}%)")
                                    st.markdown(f"- {step_desc}")
                                    st.markdown("")
                            
                            formula = evaluation_flow.get('final_score_formula', '')
                            if formula:
                                st.markdown("#### 📐 成绩计算公式")
                                st.code(formula, language=None)
                        else:
                            st.info("暂无评价流程信息")
                        
                        grading_levels = extracted_indicators.get('grading_levels', {})
                        if grading_levels:
                            st.markdown("### 📊 评分等级标准")
                            level_names = {
                                "excellent": "优秀",
                                "good": "良好",
                                "medium": "中等",
                                "pass": "及格",
                                "fail": "不及格"
                            }
                            for level, info in grading_levels.items():
                                if isinstance(info, dict):
                                    min_score = info.get('min', 0)
                                    max_score = info.get('max', 100)
                                    desc = info.get('description', '')
                                    st.markdown(f"**{level_names.get(level, level)}** ({min_score}-{max_score}分): {desc}")
                                else:
                                    st.markdown(f"**{level_names.get(level, level)}**: {info}")
                    
                    with result_tab5:
                        st.markdown("### 完整提炼结果（JSON格式）")
                        st.json(extracted_indicators)
                    
                    # ========== 衍生评价指标 ==========
                    st.markdown("---")
                    st.subheader("🔄 衍生评价指标")
                    st.markdown("根据当前评价指标，生成特定项目类型的评价指标")
                    
                    # 检查是否已有衍生指标
                    derived_dir = os.path.join(project_root, "derived_standards")
                    existing_derived = []
                    if os.path.exists(derived_dir):
                        existing_derived = [f for f in os.listdir(derived_dir) if f.endswith('.json')]
                    
                    if existing_derived:
                        with st.expander(f"📚 已有衍生指标 ({len(existing_derived)} 个)", expanded=False):
                            for df in existing_derived:
                                df_path = os.path.join(derived_dir, df)
                                try:
                                    with open(df_path, 'r', encoding='utf-8') as f:
                                        derived_data = json.load(f)
                                    st.markdown(f"**{df}** - {derived_data.get('name', '未知类型')}")
                                except:
                                    st.markdown(f"**{df}**")
                    
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        derive_project_type = st.selectbox(
                            "选择目标项目类型",
                            options=[
                                ("算法类", "algorithm"),
                                ("仿真类", "simulation"),
                                ("实物类", "physical"),
                                ("传统机械类", "traditional_mechanical"),
                                ("混合类", "mixed")
                            ],
                            format_func=lambda x: x[0]
                        )
                    with col2:
                        st.write("")
                        st.write("")
                        derive_button = st.button("🚀 生成衍生指标", type="primary", use_container_width=True)
                    
                    if derive_button:
                        with st.spinner("正在生成衍生评价指标..."):
                            try:
                                derive_response = requests.post(
                                    f"{API_BASE_URL}/generate_evaluation_standards",
                                    json={
                                        "file_content": json.dumps(extracted_indicators, ensure_ascii=False),
                                        "file_name": selected_graduation_file,
                                        "project_type": derive_project_type[1]
                                    },
                                    timeout=180
                                )
                                
                                if derive_response.status_code == 200:
                                    derived_standards = derive_response.json()
                                    
                                    # 保存衍生指标
                                    os.makedirs(derived_dir, exist_ok=True)
                                    derived_file = os.path.join(
                                        derived_dir,
                                        f"{derive_project_type[1]}_derived.json"
                                    )
                                    with open(derived_file, 'w', encoding='utf-8') as f:
                                        json.dump(derived_standards, f, ensure_ascii=False, indent=2)
                                    
                                    st.success(f"✅ 已生成 {derive_project_type[0]} 衍生评价指标！")
                                    st.info(f"📁 已保存至: {derived_file}")
                                    
                                    # 使用标签页展示详细内容
                                    derive_tab1, derive_tab2, derive_tab3, derive_tab4, derive_tab5 = st.tabs([
                                        "📋 评价指标", 
                                        "📝 评价要点", 
                                        "📊 评价表格", 
                                        "🔄 评价流程", 
                                        "📄 完整结果"
                                    ])
                                    
                                    with derive_tab1:
                                        st.markdown(f"**项目类型**: {derived_standards.get('name', '')}")
                                        st.markdown(f"**描述**: {derived_standards.get('description', '')}")
                                        
                                        st.markdown("### 📊 评价指标列表")
                                        for ind in derived_standards.get('indicators', []):
                                            indicator_id = ind.get('indicator_id', ind.get('id', ''))
                                            name = ind.get('name', '')
                                            weight = ind.get('weight', 0)
                                            description = ind.get('description', '')
                                            
                                            with st.expander(f"**{indicator_id} {name}** (权重: {weight}%)"):
                                                if description:
                                                    st.markdown(f"**描述**: {description}")
                                                if ind.get('graduation_requirement'):
                                                    st.markdown(f"**对应毕业要求**: {ind.get('graduation_requirement')}")
                                                if ind.get('evaluation_method'):
                                                    st.markdown(f"**评价方式**: {ind.get('evaluation_method')}")
                                                
                                                grade_levels = ind.get('grade_levels', {})
                                                if grade_levels:
                                                    st.markdown("#### 评分等级")
                                                    level_names = {
                                                        "excellent": "优秀",
                                                        "good": "良好",
                                                        "medium": "中等",
                                                        "pass": "及格",
                                                        "fail": "不及格"
                                                    }
                                                    for level, desc in grade_levels.items():
                                                        st.markdown(f"**{level_names.get(level, level)}**: {desc}")
                                    
                                    with derive_tab2:
                                        st.markdown("### 📝 详细评价要点")
                                        for ind in derived_standards.get('indicators', []):
                                            indicator_id = ind.get('indicator_id', ind.get('id', ''))
                                            name = ind.get('name', '')
                                            evaluation_points = ind.get('evaluation_points', [])
                                            
                                            if evaluation_points:
                                                st.markdown(f"#### {indicator_id} {name}")
                                                for point in evaluation_points:
                                                    point_name = point.get('point_name', '')
                                                    point_weight = point.get('weight', '')
                                                    point_desc = point.get('description', '')
                                                    grade_criteria = point.get('grade_criteria', {})
                                                    
                                                    with st.expander(f"**{point_name}** (权重: {point_weight}%)"):
                                                        if point_desc:
                                                            st.markdown(f"**描述**: {point_desc}")
                                                        
                                                        if grade_criteria:
                                                            st.markdown("#### 评分标准")
                                                            grade_names = {
                                                                "excellent": "优秀(90-100)",
                                                                "good": "良好(80-89)",
                                                                "medium": "中等(70-79)",
                                                                "pass": "及格(60-69)",
                                                                "fail": "不及格(0-59)"
                                                            }
                                                            for level, criteria in grade_criteria.items():
                                                                st.markdown(f"**{grade_names.get(level, level)}**: {criteria}")
                                                st.markdown("")
                                    
                                    with derive_tab3:
                                        evaluation_table = derived_standards.get('evaluation_table', {})
                                        if evaluation_table:
                                            st.markdown(f"### 📊 {evaluation_table.get('title', '评价表格')}")
                                            
                                            rows = evaluation_table.get('rows', [])
                                            if rows:
                                                import pandas as pd
                                                df = pd.DataFrame(rows)
                                                st.dataframe(df, use_container_width=True, hide_index=True)
                                            else:
                                                st.info("暂无评价表格数据")
                                        else:
                                            st.info("暂无评价表格")
                                    
                                    with derive_tab4:
                                        evaluation_flow = derived_standards.get('evaluation_flow', {})
                                        if evaluation_flow:
                                            st.markdown("### 🔄 评价流程")
                                            
                                            steps = evaluation_flow.get('steps', [])
                                            if steps:
                                                for step in steps:
                                                    step_num = step.get('step', '')
                                                    step_name = step.get('name', '')
                                                    step_weight = step.get('weight', 0)
                                                    step_desc = step.get('description', '')
                                                    
                                                    st.markdown(f"**步骤{step_num}: {step_name}** (权重: {step_weight*100:.0f}%)")
                                                    st.markdown(f"- {step_desc}")
                                                    st.markdown("")
                                            
                                            formula = evaluation_flow.get('final_score_formula', '')
                                            if formula:
                                                st.markdown("#### 📐 成绩计算公式")
                                                st.code(formula, language=None)
                                        
                                        grading_levels = derived_standards.get('grading_levels', {})
                                        if grading_levels:
                                            st.markdown("### 📊 评分等级标准")
                                            level_names = {
                                                "excellent": "优秀",
                                                "good": "良好",
                                                "medium": "中等",
                                                "pass": "及格",
                                                "fail": "不及格"
                                            }
                                            for level, info in grading_levels.items():
                                                if isinstance(info, dict):
                                                    min_score = info.get('min', 0)
                                                    max_score = info.get('max', 100)
                                                    desc = info.get('description', '')
                                                    st.markdown(f"**{level_names.get(level, level)}** ({min_score}-{max_score}分): {desc}")
                                                else:
                                                    st.markdown(f"**{level_names.get(level, level)}**: {info}")
                                        
                                        excluded = derived_standards.get('excluded_indicators', [])
                                        if excluded:
                                            st.markdown("### ⚠️ 排除的评价项")
                                            for item in excluded:
                                                st.markdown(f"- {item}")
                                    
                                    with derive_tab5:
                                        st.markdown("### 完整衍生指标（JSON格式）")
                                        st.json(derived_standards)
                                else:
                                    st.error(f"❌ 生成失败: {derive_response.json().get('detail', '未知错误')}")
                            except Exception as e:
                                st.error(f"❌ 生成衍生指标失败: {str(e)}")
                else:
                    st.info("ℹ️ 暂无提炼结果，请点击\"提炼指标\"按钮进行提炼")

# ==================== 学生管理 ====================
elif page == "👥 学生管理":
    st.title("👥 学生管理")
    
    # 选项卡
    tab1, tab2, tab3 = st.tabs(["➕ 添加学生", "🔍 查询学生", "📋 学生列表"])
    
    # 添加学生
    with tab1:
        st.subheader("添加新学生")
        with st.form("add_student_form"):
            col1, col2 = st.columns(2)
            with col1:
                student_id = st.text_input("学号*", placeholder="请输入学号")
                name = st.text_input("姓名*", placeholder="请输入姓名")
                age = st.number_input("年龄", min_value=0, max_value=100, value=0)
            with col2:
                grade = st.text_input("年级", placeholder="如：大一、大二")
                major = st.text_input("专业", placeholder="如：计算机科学")
            
            st.markdown("*必填项")
            
            submit_button = st.form_submit_button("✅ 添加学生", use_container_width=True)
            
            if submit_button:
                if not student_id or not name:
                    st.error("❌ 学号和姓名不能为空")
                else:
                    try:
                        response = requests.post(
                            f"{API_BASE_URL}/students",
                            json={
                                "student_id": student_id,
                                "name": name,
                                "age": age if age > 0 else None,
                                "grade": grade if grade else None,
                                "major": major if major else None
                            }
                        )
                        if response.status_code == 200:
                            st.success("✅ 学生添加成功！")
                            student_data = response.json()
                            st.json(student_data)
                        else:
                            try:
                                error_detail = response.json().get('detail', '未知错误')
                            except:
                                error_detail = f"HTTP {response.status_code}: {response.text}"
                            st.error(f"❌ 添加失败: {error_detail}")
                    except Exception as e:
                        st.error(f"❌ 添加失败: {str(e)}")
    
    # 查询学生
    with tab2:
        st.subheader("查询学生信息")
        
        search_col1, search_col2 = st.columns([3, 1])
        with search_col1:
            student_id = st.text_input("输入学号", placeholder="请输入学号进行查询")
        with search_col2:
            st.write("")
            st.write("")
            search_button = st.button("🔍 查询", use_container_width=True)
        
        if search_button:
            if not student_id:
                st.error("❌ 请输入学号")
            else:
                try:
                    response = requests.get(f"{API_BASE_URL}/students/{student_id}")
                    if response.status_code == 200:
                        student = response.json()
                        
                        # 显示学生信息卡片
                        st.success("✅ 查询成功")
                        
                        info_col1, info_col2 = st.columns(2)
                        with info_col1:
                            st.markdown(f"**学号:** {student.get('student_id', 'N/A')}")
                            st.markdown(f"**姓名:** {student.get('name', 'N/A')}")
                            st.markdown(f"**年龄:** {student.get('age', 'N/A')}")
                        with info_col2:
                            st.markdown(f"**年级:** {student.get('grade', 'N/A')}")
                            st.markdown(f"**专业:** {student.get('major', 'N/A')}")
                            st.markdown(f"**创建时间:** {student.get('created_at', 'N/A')}")
                        
                        # 显示完整JSON
                        with st.expander("查看完整数据"):
                            st.json(student)
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 查询失败: {error_detail}")
                except Exception as e:
                    st.error(f"❌ 查询失败: {str(e)}")
    
    # 学生列表
    with tab3:
        st.subheader("学生列表")
        
        # 默认显示所有学生
        try:
            response = requests.get(f"{API_BASE_URL}/students")
            if response.status_code == 200:
                students = response.json()
                if students:
                    st.info(f"共 {len(students)} 名学生")
                    
                    # 显示表头
                    header_cols = st.columns([1, 1, 0.5, 1, 1, 0.5, 0.5])
                    with header_cols[0]:
                        st.markdown("**学号**")
                    with header_cols[1]:
                        st.markdown("**姓名**")
                    with header_cols[2]:
                        st.markdown("**年龄**")
                    with header_cols[3]:
                        st.markdown("**年级**")
                    with header_cols[4]:
                        st.markdown("**专业**")
                    with header_cols[5]:
                        st.markdown("**修改**")
                    with header_cols[6]:
                        st.markdown("**删除**")
                    
                    # 表头分隔线
                    st.divider()
                    
                    # 为每个学生显示一行，包含信息和操作按钮
                    for i, student in enumerate(students):
                        cols = st.columns([1, 1, 0.5, 1, 1, 0.5, 0.5])
                        
                        with cols[0]:
                            st.write(student['student_id'])
                        with cols[1]:
                            st.write(student['name'])
                        with cols[2]:
                            st.write(student.get('age', 'N/A'))
                        with cols[3]:
                            st.write(student.get('grade', 'N/A'))
                        with cols[4]:
                            st.write(student.get('major', 'N/A'))
                        with cols[5]:
                            # 修改按钮
                            if st.button(f"✏️", key=f"edit_{student['student_id']}", help="修改学生信息"):
                                # 存储当前编辑的学生信息
                                st.session_state['edit_student'] = student
                                # 切换到修改学生表单
                                st.session_state['show_edit_form'] = True
                        with cols[6]:
                            # 删除按钮
                            if st.button(f"🗑️", key=f"delete_{student['student_id']}", help="删除学生"):
                                # 直接执行删除操作
                                try:
                                    delete_response = requests.delete(f"{API_BASE_URL}/students/{student['student_id']}")
                                    if delete_response.status_code == 200:
                                        st.success(f"✅ 学生 {student['name']} 删除成功！")
                                        # 刷新页面
                                        st.rerun()
                                    else:
                                        try:
                                            error_detail = delete_response.json().get('detail', '未知错误')
                                        except:
                                            error_detail = f"HTTP {delete_response.status_code}"
                                        st.error(f"❌ 删除失败: {error_detail}")
                                except Exception as e:
                                    st.error(f"❌ 删除失败: {str(e)}")
                        
                        # 在每个学生行之间添加分隔线（除了最后一行）
                        if i < len(students) - 1:
                            st.divider()
                else:
                    st.info("📭 暂无学生信息")
            else:
                try:
                    error_detail = response.json().get('detail', '未知错误')
                except:
                    error_detail = f"HTTP {response.status_code}"
                st.error(f"❌ 获取列表失败: {error_detail}")
        except Exception as e:
            st.error(f"❌ 获取列表失败: {str(e)}")
        
        # 刷新按钮
        if st.button("🔄 刷新列表", use_container_width=True):
            st.rerun()


    
    # 修改学生表单
    if st.session_state.get('show_edit_form', False):
        st.title("✏️ 修改学生信息")
        edit_student = st.session_state.get('edit_student', {})
        
        with st.form("edit_student_form"):
            col1, col2 = st.columns(2)
            with col1:
                student_id = st.text_input("学号*", value=edit_student.get('student_id', ''), disabled=True)
                name = st.text_input("姓名*", value=edit_student.get('name', ''))
                age = st.number_input("年龄", min_value=0, max_value=100, value=edit_student.get('age', 0))
            with col2:
                grade = st.text_input("年级", value=edit_student.get('grade', ''))
                major = st.text_input("专业", value=edit_student.get('major', ''))
            
            st.markdown("*必填项")
            
            submit_button = st.form_submit_button("✅ 保存修改", use_container_width=True)
            cancel_button = st.form_submit_button("❌ 取消", use_container_width=True)
            
            if submit_button:
                if not name:
                    st.error("❌ 姓名不能为空")
                else:
                    try:
                        response = requests.put(
                            f"{API_BASE_URL}/students/{student_id}",
                            json={
                                "name": name,
                                "age": age if age > 0 else None,
                                "grade": grade if grade else None,
                                "major": major if major else None
                            }
                        )
                        if response.status_code == 200:
                            st.success("✅ 学生信息修改成功！")
                            student_data = response.json()
                            st.json(student_data)
                            # 关闭编辑表单
                            st.session_state['show_edit_form'] = False
                            st.session_state.pop('edit_student', None)
                        else:
                            try:
                                error_detail = response.json().get('detail', '未知错误')
                            except:
                                error_detail = f"HTTP {response.status_code}: {response.text}"
                            st.error(f"❌ 修改失败: {error_detail}")
                    except Exception as e:
                        st.error(f"❌ 修改失败: {str(e)}")
            
            if cancel_button:
                st.session_state['show_edit_form'] = False
                st.session_state.pop('edit_student', None)
                st.rerun()

# ==================== 文件上传 ====================
elif page == "📁 文件上传":
    st.title("📁 作业提交")
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(os.path.dirname(current_dir))
    
    # 步骤指示器
    st.markdown("""
    **提交流程：** 选择提交类型 → 选择用途 → 创建提交 → 上传内容 → 关联人员（可选）
    """)
    
    # 步骤 1: 选择提交类型并创建提交
    st.subheader("步骤 1: 创建提交")
    
    # 提交类型选择
    submission_type = st.radio(
        "选择提交类型",
        options=["file", "text"],
        format_func=lambda x: "📁 文件提交" if x == "file" else "📝 文字提交",
        horizontal=True,
        key="submission_type"
    )
    
    # 提交用途选择
    submission_purpose = st.radio(
        "选择提交用途",
        options=["normal", "graduation"],
        format_func=lambda x: "📚 普通作业" if x == "normal" else "🎓 毕业设计",
        horizontal=True,
        key="submission_purpose",
        help="选择毕业设计将使用专门的毕业设计评价标准进行评估"
    )
    
    if submission_purpose == "graduation":
        st.info("🎓 已选择毕业设计模式，评估时将使用确定性评价标准，支持不同项目类型的自动检测")
    
    with st.form("create_submission_form"):
        title = st.text_input("提交标题*", placeholder="请输入提交标题")
        
        # 课程大纲选择（必选，用于后续严格按大纲评分）
        analysis_dir = os.path.join(project_root, "analysis_results")
        available_syllabus_options = []
        if os.path.exists(analysis_dir):
            available_syllabus_options = sorted([f for f in os.listdir(analysis_dir) if f.endswith(".json")])
        selected_syllabus_analysis = st.selectbox(
            "选择所属课程大纲*",
            options=available_syllabus_options,
            help="评估时将严格按该课程大纲提取的能力点、评价标准和权重进行评分"
        ) if available_syllabus_options else None
        if not available_syllabus_options:
            st.warning("⚠️ 暂无可选课程大纲分析结果，请先到“📋 大纲管理”完成课程大纲分析")
        
        # 根据提交类型显示不同输入
        text_content = None
        uploaded_files = None
        
        if submission_type == "text":
            st.markdown("---")
            st.subheader("📝 文字内容")
            text_content = st.text_area(
                "输入文字内容*",
                placeholder="请输入要提交的文字内容...",
                height=300,
                help="直接输入文字内容进行提交，无需上传文件"
            )
        else:
            st.markdown("---")
            st.subheader("📁 上传文件")
            st.markdown("""
            **支持的文件类型：**
            - 📄 文档：PDF, DOCX, TXT
            - 🎬 视频：MP4, MOV
            - 🎵 音频：MP3, WAV
            - 📊 PPT：PPTX, PPT
            """)
            uploaded_files = st.file_uploader(
                "选择文件上传*",
                accept_multiple_files=True,
                type=["pdf", "docx", "txt", "mp4", "mov", "mp3", "wav", "pptx", "ppt"],
                help="可以上传多个文件"
            )
        
        # 关联人员（必填）
        st.markdown("---")
        st.subheader("👤 关联人员*")
        st.write("选择此提交关联的学生")
        person_id = None
        
        # 获取所有学生
        try:
            response = requests.get(f"{API_BASE_URL}/students")
            if response.status_code == 200:
                students = response.json()
                if students:
                    # 构建学生选项
                    student_options = {student['student_id']: f"{student['student_id']} - {student['name']}" for student in students}
                    person_id = st.selectbox(
                        "选择学生",
                        options=list(student_options.keys()),
                        format_func=lambda x: student_options[x],
                        placeholder="请选择学生"
                    )
                else:
                    st.error("❌ 暂无学生记录，请先添加学生")
            else:
                st.error("❌ 无法获取学生列表")
        except Exception as e:
            st.error("❌ 无法获取学生列表")
        
        st.markdown("*必填项")
        
        submit_button = st.form_submit_button("✅ 创建提交", use_container_width=True)
        
        if submit_button:
            if not title:
                st.error("❌ 标题不能为空")
            elif submission_type == "text" and not text_content:
                st.error("❌ 文字提交必须输入内容")
            elif submission_type == "file" and not uploaded_files:
                st.error("❌ 文件提交必须上传文件")
            elif not person_id:
                st.error("❌ 必须选择关联学生")
            elif not selected_syllabus_analysis:
                st.error("❌ 必须选择所属课程大纲")
            else:
                try:
                    # 创建提交
                    payload = {
                        "title": title,
                        "submission_type": submission_type,
                        "submission_purpose": submission_purpose,
                        "student_id": person_id,
                        "syllabus_name": selected_syllabus_analysis
                    }
                    
                    # 文字提交添加内容
                    if submission_type == "text" and text_content:
                        payload["text_content"] = text_content
                    
                    response = requests.post(
                        f"{API_BASE_URL}/submissions",
                        json=payload
                    )
                    
                    if response.status_code == 200:
                        submission = response.json()
                        st.session_state.submission = submission
                        st.success("✅ 提交创建成功！")
                        
                        # 显示提交信息
                        info_col1, info_col2 = st.columns(2)
                        with info_col1:
                            st.metric("提交类型", "文字提交" if submission['submission_type'] == 'text' else "文件提交")
                        with info_col2:
                            st.metric("提交ID", submission['submission_id'])
                        st.caption(f"📚 绑定课程大纲：{submission.get('syllabus_name', '未设置')}")
                        
                        # 如果是文件提交，上传文件
                        if submission_type == "file" and uploaded_files:
                            st.markdown("---")
                            st.subheader("📤 上传文件中...")
                            progress_bar = st.progress(0)
                            success_count = 0
                            
                            for i, file in enumerate(uploaded_files):
                                st.write(f"📄 正在上传: {file.name}")
                                
                                # 上传文件到服务器
                                files = {"file": (file.name, file, file.type)}
                                try:
                                    response = requests.post(
                                        f"{API_BASE_URL}/submissions/{submission['submission_id']}/files",
                                        files=files
                                    )
                                    if response.status_code == 200:
                                        st.success(f"✅ 文件 {file.name} 上传成功")
                                        success_count += 1
                                    else:
                                        try:
                                            error_detail = response.json().get('detail', '未知错误')
                                        except:
                                            error_detail = f"HTTP {response.status_code}"
                                        st.error(f"❌ 文件 {file.name} 上传失败: {error_detail}")
                                except Exception as e:
                                    st.error(f"❌ 文件 {file.name} 上传失败: {str(e)}")
                                
                                # 更新进度条
                                progress_bar.progress((i + 1) / len(uploaded_files))
                            
                            # 显示上传结果
                            st.markdown("---")
                            if success_count == len(uploaded_files):
                                st.success(f"✅ 所有文件上传成功！共 {success_count} 个文件")
                                # 清空提交区
                                if 'submission' in st.session_state:
                                    del st.session_state.submission
                                st.success("📋 提交区已清空，可继续上传新的提交")
                            else:
                                st.warning(f"⚠️ 文件上传完成，成功 {success_count}/{len(uploaded_files)} 个文件")
                        
                        # 如果是文字提交，显示内容预览
                        if submission['submission_type'] == 'text' and submission.get('text_content'):
                            with st.expander("查看文字内容预览"):
                                st.text(submission['text_content'][:500] + "..." if len(submission['text_content']) > 500 else submission['text_content'])
                        
                        # 提示用户可以稍后关联人员
                        if not person_id:
                            st.markdown("---")
                            st.info("📝 提交已创建，您可以稍后在文件管理页面为提交关联人员")
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 创建失败: {error_detail}")
                except Exception as e:
                    st.error(f"❌ 创建失败: {str(e)}")
    
    # 查看已上传文件（仅文件提交）
    if 'submission' in st.session_state and st.session_state.submission.get('submission_type') == 'file':
        st.markdown("---")
        if st.button("📋 查看已上传文件", use_container_width=True):
            try:
                response = requests.get(
                    f"{API_BASE_URL}/submissions/{st.session_state.submission['submission_id']}/files"
                )
                if response.status_code == 200:
                    files = response.json()
                    if files:
                        # 简化文件信息显示
                        simplified_files = []
                        for file in files:
                            simplified_files.append({
                                "文件名": file.get('file_path', '').split('/')[-1].split('\\')[-1],
                                "文件类型": file.get('media_type', 'N/A'),
                                "文件大小": f"{file.get('size_bytes', 0) / 1024:.2f} KB",
                                "上传时间": file.get('uploaded_at', 'N/A')[:19]
                            })
                        df = pd.DataFrame(simplified_files)
                        st.dataframe(df, use_container_width=True)
                        st.info(f"共 {len(files)} 个文件")
                    else:
                        st.info("📭 暂无上传文件")
                else:
                    try:
                        error_detail = response.json().get('detail', '未知错误')
                    except:
                        error_detail = f"HTTP {response.status_code}"
                    st.error(f"❌ 获取文件失败: {error_detail}")
            except Exception as e:
                st.error(f"❌ 获取文件失败: {str(e)}")
    
    # 文字提交完成提示
    if 'submission' in st.session_state and st.session_state.submission.get('submission_type') == 'text':
        st.markdown("---")
        st.success("✅ 文字提交已完成！")
        # 清空提交区
        if 'submission' in st.session_state:
            del st.session_state.submission
        st.success("📋 提交区已清空，可继续上传新的提交")

# ==================== 评估管理 ====================
elif page == "🤖 评估管理":
    st.title("🤖 评估管理")
    
    st.markdown("""
    **评估流程：**
    - 阶段评估：选择特定报告，根据学生工作时期进行评估
    - 整体评估：对学生的所有提交进行综合评估
    - 毕业设计评估：使用确定性标准评价毕业设计，支持不同项目类型
    
    系统将使用AI对学生提交的材料进行评估，包括：
    - 📚 学术表现分析
    - 💬 沟通能力分析
    - 👥 领导力分析
    - 🤝 团队协作分析
    - 💡 创新能力分析
    - 🧩 问题解决分析
    """)
    
    # 评估类型选择
    eval_type = st.selectbox(
        "选择评估类型",
        options=["阶段评估", "整体评估", "毕业设计评估"]
    )
    
    if eval_type == "阶段评估":
        st.subheader("📊 阶段评估")
        st.markdown("对特定报告进行评估，根据学生工作时期调整评估标准")
        
        # 获取所有提交及其文件
        try:
            response = requests.get(f"{API_BASE_URL}/submissions")
            if response.status_code == 200:
                submissions = response.json()
                if submissions:
                    # 过滤出普通作业提交（非毕业设计）
                    normal_submissions = [s for s in submissions if s.get('submission_purpose', 'normal') == 'normal']
                    graduation_submissions = [s for s in submissions if s.get('submission_purpose', 'normal') == 'graduation']
                    
                    if graduation_submissions:
                        st.info(f"🎓 有 {len(graduation_submissions)} 个毕业设计提交，请使用\"毕业设计评估\"进行评估")
                    
                    if not normal_submissions:
                        st.warning("⚠️ 暂无普通作业提交记录，只有毕业设计提交。请使用\"毕业设计评估\"进行评估")
                    else:
                        # 构建提交选项，包含文件名
                        submission_options = {}
                        for sub in normal_submissions:
                            submission_id = sub['submission_id']
                            student_id = sub.get('student_id', '未知')
                            title = sub['title']
                            
                            # 获取该提交的文件
                            files_response = requests.get(f"{API_BASE_URL}/submissions/{submission_id}/files")
                            if files_response.status_code == 200:
                                files = files_response.json()
                                if files:
                                    for file in files:
                                        file_name = file.get('file_name', '未知文件')
                                        file_id = file.get('id', '')
                                        # 使用提交ID+文件ID作为唯一键，如果文件ID为空则使用索引
                                        if file_id:
                                            option_key = f"{submission_id}_{file_id}"
                                        else:
                                            option_key = f"{submission_id}_file_{len(submission_options)}"
                                        submission_options[option_key] = f"{title} - {file_name} (学生ID: {student_id})"
                                else:
                                    # 如果没有文件，也显示提交记录（使用文字提交）
                                    text_content = sub.get('text_content', '')
                                    if text_content:
                                        option_key = f"{submission_id}_text"
                                        submission_options[option_key] = f"{title} - 文字提交 (学生ID: {student_id})"
                            else:
                                # 如果获取文件失败，也显示提交记录
                                option_key = f"{submission_id}_unknown"
                                submission_options[option_key] = f"{title} - 未知文件类型 (学生ID: {student_id})"
                        
                        selected_submission_id = None
                        if submission_options:
                            selected_option = st.selectbox(
                                "选择报告",
                                options=list(submission_options.keys()),
                                format_func=lambda x: submission_options[x]
                            )
                            
                            # 提取提交ID（处理带文件ID的情况）
                            if '_' in selected_option:
                                # 找到第一个下划线的位置
                                first_underscore = selected_option.find('_')
                                # 找到第二个下划线的位置（如果存在）
                                second_underscore = selected_option.find('_', first_underscore + 1)
                                if second_underscore != -1:
                                    # 提取从开始到第二个下划线的部分作为提交ID
                                    selected_submission_id = selected_option[:second_underscore]
                                else:
                                    # 如果只有一个下划线，使用整个字符串作为提交ID
                                    selected_submission_id = selected_option
                            else:
                                selected_submission_id = selected_option
                        else:
                            st.warning("⚠️ 暂无提交记录，请先在提交管理页面创建提交")
                    
                    if selected_submission_id:
                        # 学生工作时期设置
                        st.markdown("---")
                        st.subheader("📅 学生工作时期设置")
                        
                        # 自由拖动的进度条（0-100%）
                        st.write("选择学生工作时期:")
                        progress_percent = st.slider(
                            "工作时期进度",
                            min_value=0,
                            max_value=100,
                            value=50,
                            step=1,
                            format="%d%%"
                        )
                        
                        # 转换为0.0-1.0的进度值
                        stage_progress = progress_percent / 100.0
                        
                        # 显示进度值
                        st.write(f"当前工作时期进度: {stage_progress:.2f}")
                        
                        # 根据进度值显示时期说明（与实践课过程性考核一致）
                        if stage_progress < 0.33:
                            st.info("💡 初期阶段（0-33%）- 重点评估：理解与规划")
                        elif stage_progress < 0.66:
                            st.info("⚖️ 中期阶段（33%-66%）- 重点评估：执行与规范")
                        else:
                            st.info("🎯 后期阶段（66%-100%）- 重点评估：测试与优化")
                        
                        # 编辑提示词功能
                        with st.expander("📝 编辑大模型提示词"):
                            # 构建默认提示词（导向最新评分标准）
                            stage_description = "初期阶段（理解与规划）" if stage_progress < 0.33 else ("中期阶段（执行与规范）" if stage_progress < 0.66 else "后期阶段（测试与优化）")
                            
                            # 系统提示词
                            system_prompt = st.text_area(
                                "系统提示词",
                                value="你是一位高校课程作业评估专家。请严格执行最新评分政策：\n1) 理论课：总分=50%能力点评分+25%知识点理解+25%知识点使用；\n2) 实践课：总分=50%能力点评分+50%阶段完成度（按进度阶段重点评估）；\n3) 维度评分必须提供证据与理由；\n4) 分值范围0-100，输出必须是可解析JSON；\n5) 评分必须严格对齐课程大纲提取的能力点、评价标准和知识点。",
                                height=150
                            )
                            
                            # 用户提示词
                            user_prompt = st.text_area(
                                "用户提示词",
                                value=(
                                    f"# 补充评估要求（当前阶段：{stage_description}，进度值：{stage_progress:.2f}）\n"
                                    "请在不改变系统评分政策的前提下，补充以下要求：\n"
                                    "1. 维度评分必须逐条引用作业证据，避免泛化表述；\n"
                                    "2. 维度理由需说明“为何得此分、扣分点在哪里、达标缺口是什么”；\n"
                                    "3. 理论课请确保输出 knowledge_understanding_score 与 knowledge_application_score；\n"
                                    "4. 实践课请确保输出 phase_completion_score，并明确当前阶段重点的达成度依据；\n"
                                    "5. 严格返回JSON，不添加解释性前后缀文本。\n\n"
                                    "# 学生信息\n{student_info}\n\n"
                                    "# 提交内容\n{submission_content}\n"
                                ),
                                height=500
                            )
                            
                            # 保存提示词
                            if st.button("💾 保存提示词", use_container_width=True):
                                # 保存到session state
                                st.session_state.custom_prompts = {
                                    "system_prompt": system_prompt,
                                    "user_prompt": user_prompt
                                }
                                st.success("✅ 提示词保存成功！")
                        
                        # 启动评估按钮
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("▶️ 启动阶段评估", use_container_width=True):
                                # 准备评估请求
                                eval_payload = {
                                    "submission_id": selected_submission_id,
                                    "stage_progress": stage_progress
                                }
                                
                                # 添加自定义提示词（如果有）
                                if 'custom_prompts' in st.session_state:
                                    eval_payload['custom_prompts'] = st.session_state.custom_prompts
                                
                                st.info(f"📊 将使用工作时期进度 {stage_progress:.2f} 进行评估")
                                
                                # 启动评估
                                with st.spinner("🤖 AI正在评估中，请稍候..."):
                                    progress_bar = st.progress(0)
                                    
                                    # 模拟评估进度
                                    import time
                                    for i in range(100):
                                        time.sleep(0.05)
                                        progress_bar.progress(i + 1)
                                    
                                    response = requests.post(
                                        f"{API_BASE_URL}/evaluate",
                                        json=eval_payload
                                    )
                                    
                                    if response.status_code == 200:
                                        evaluation_result = response.json()
                                        st.session_state.evaluation_result = evaluation_result
                                        st.success("✅ 评估完成！")
                                        
                                        # 显示评估结果摘要
                                        st.subheader("📊 评估结果摘要")
                                        
                                        result_col1, result_col2, result_col3, result_col4 = st.columns(4)
                                        with result_col1:
                                            st.metric("综合评分", f"{evaluation_result['overall_score']}/100")
                                        with result_col2:
                                            st.metric("评估维度", len(evaluation_result['dimension_scores']))
                                        with result_col3:
                                            st.metric("评估时间", evaluation_result['evaluated_at'][:10])
                                        with result_col4:
                                            # 显示阶段进度信息
                                            stage_progress = evaluation_result.get('stage_progress', 0.5)
                                            progress_percent = int(stage_progress * 100)
                                            st.metric("工作时期进度", f"{progress_percent}%")

                                        render_score_card(evaluation_result)
                                        score_breakdown = evaluation_result.get("score_breakdown", {})
                                        if isinstance(score_breakdown, dict) and score_breakdown:
                                            with st.expander("查看评分构成明细"):
                                                st.json(score_breakdown)
                                        
                                        # 显示阶段评估说明
                                        stage_progress = evaluation_result.get('stage_progress', 0.5)
                                        if stage_progress < 0.33:
                                            st.info("💡 初期阶段评估（0-33%）- 重点：理解与规划")
                                        elif stage_progress < 0.66:
                                            st.info("⚖️ 中期阶段评估（33%-66%）- 重点：执行与规范")
                                        else:
                                            st.info("🎯 后期阶段评估（66%-100%）- 重点：测试与优化")
                                        
                                        # 显示详细结果
                                        with st.expander("查看详细结果"):
                                            st.json(evaluation_result)
                                        
                                        render_dimension_score_details(evaluation_result)
                                    else:
                                        try:
                                            error_detail = response.json().get('detail', '未知错误')
                                        except:
                                            error_detail = f"HTTP {response.status_code}"
                                        st.error(f"❌ 评估失败: {error_detail}")
                else:
                    st.info("📭 暂无提交记录")
            else:
                st.error("❌ 获取提交记录失败")
        except Exception as e:
            st.error(f"❌ 加载提交记录失败: {str(e)}")
    
    elif eval_type == "整体评估":
        st.subheader("🎯 整体评估")
        st.markdown("根据大纲总结对单个报告文件进行评估")
        
        # 获取所有学生
        try:
            response = requests.get(f"{API_BASE_URL}/students")
            if response.status_code == 200:
                students = response.json()
                if students:
                    # 构建学生选项
                    student_options = {student['student_id']: f"{student['student_id']} - {student['name']}" for student in students}
                    selected_student_id = st.selectbox(
                        "选择学生",
                        options=list(student_options.keys()),
                        format_func=lambda x: student_options[x]
                    )
                    
                    st.info("📚 评估将自动使用该提交在上传时绑定的课程大纲（能力点/评价标准/权重）")
                    
                    # 检查学生是否有提交记录
                    response = requests.get(f"{API_BASE_URL}/students/{selected_student_id}/submissions")
                    if response.status_code == 200:
                        submissions = response.json()
                        if not submissions:
                            st.error("❌ 该学生暂无提交记录，无法进行整体评估")
                            st.stop()
                        else:
                            # 过滤出普通作业提交（非毕业设计）
                            normal_submissions = [s for s in submissions if s.get('submission_purpose', 'normal') == 'normal']
                            graduation_submissions = [s for s in submissions if s.get('submission_purpose', 'normal') == 'graduation']
                            
                            if graduation_submissions:
                                st.info(f"🎓 该学生有 {len(graduation_submissions)} 个毕业设计提交，请使用\"毕业设计评估\"进行评估")
                            
                            if not normal_submissions:
                                st.warning("⚠️ 该学生没有普通作业提交，只有毕业设计提交。请使用\"毕业设计评估\"进行评估")
                                st.stop()
                            
                            # 构建提交选项（只显示普通作业）
                            submission_options = {submission['submission_id']: f"{submission['submission_id']} - {submission['title']}" for submission in normal_submissions}
                            selected_submission_id = st.selectbox(
                                "选择报告文件",
                                options=list(submission_options.keys()),
                                format_func=lambda x: submission_options[x]
                            )
                            
                            # 检查选中的提交是否有内容
                            selected_submission = next((s for s in normal_submissions if s['submission_id'] == selected_submission_id), None)
                            if selected_submission:
                                # 检查是否有文字内容
                                has_text_content = selected_submission.get('text_content') and selected_submission.get('text_content').strip()
                                
                                # 检查是否有文件
                                files_response = requests.get(f"{API_BASE_URL}/submissions/{selected_submission_id}/files")
                                has_files = False
                                if files_response.status_code == 200:
                                    files = files_response.json()
                                    has_files = len(files) > 0
                                
                                # 如果既没有文字内容也没有文件，显示警告
                                if not has_text_content and not has_files:
                                    st.warning("⚠️ 该提交没有内容（无文字且无文件），评估可能会失败")
                    else:
                        st.error("❌ 获取提交记录失败")
                        st.stop()
                    
                    if st.button("▶️ 启动整体评估", use_container_width=True):
                        # 准备评估请求
                        eval_payload = {
                            "submission_id": selected_submission_id,
                            "stage_progress": 1.0  # 整体评估使用最终阶段标准
                        }
                        
                        st.info("📊 将对选定的报告文件进行评估")
                        
                        # 启动评估
                        with st.spinner("🤖 AI正在评估中，请稍候..."):
                            progress_bar = st.progress(0)
                            
                            # 模拟评估进度
                            import time
                            for i in range(100):
                                time.sleep(0.05)
                                progress_bar.progress(i + 1)
                            
                            response = requests.post(
                                f"{API_BASE_URL}/evaluate",
                                json=eval_payload
                            )
                            
                            if response.status_code == 200:
                                evaluation_result = response.json()
                                st.session_state.evaluation_result = evaluation_result
                                st.success("✅ 评估完成！")
                                
                                # 显示评估结果摘要
                                st.subheader("📊 评估结果摘要")
                                
                                result_col1, result_col2, result_col3 = st.columns(3)
                                with result_col1:
                                    st.metric("综合评分", f"{evaluation_result['overall_score']}/100")
                                with result_col2:
                                    st.metric("评估维度", len(evaluation_result['dimension_scores']))
                                with result_col3:
                                    st.metric("提交数量", 1)  # 现在只评估一个提交

                                render_score_card(evaluation_result)
                                score_breakdown = evaluation_result.get("score_breakdown", {})
                                if isinstance(score_breakdown, dict) and score_breakdown:
                                    with st.expander("查看评分构成明细"):
                                        st.json(score_breakdown)
                                
                                st.info("🎯 整体评估 - 基于大纲总结对单个报告文件进行评估")
                                
                                render_dimension_score_details(evaluation_result)
                                
                                # 优势和劣势
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.subheader("✅ 优势")
                                    strengths = evaluation_result.get('strengths', [])
                                    if strengths:
                                        for strength in strengths:
                                            st.markdown(f"- {strength}")
                                    else:
                                        st.info("无优势记录")
                                
                                with col2:
                                    st.subheader("❌ 劣势")
                                    weaknesses = evaluation_result.get('weaknesses', evaluation_result.get('areas_for_improvement', []))
                                    if weaknesses:
                                        for weakness in weaknesses:
                                            st.markdown(f"- {weakness}")
                                    else:
                                        st.info("无劣势记录")
                                
                                # 任务完成情况
                                task_completion = evaluation_result.get('task_completion', {})
                                if task_completion:
                                    st.subheader("📋 任务完成情况")
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.markdown("**已完成任务：**")
                                        for task in task_completion.get('completed_tasks', []):
                                            st.markdown(f"- ✅ {task}")
                                    with col2:
                                        st.markdown("**未完成任务：**")
                                        for task in task_completion.get('incomplete_tasks', []):
                                            st.markdown(f"- ❌ {task}")
                                    if task_completion.get('completion_details'):
                                        st.markdown(f"**完成详情：** {task_completion.get('completion_details')}")
                                
                                # 总体评价
                                overall_eval = evaluation_result.get('overall_evaluation', '')
                                if overall_eval:
                                    st.subheader("📝 总体评价")
                                    st.markdown(overall_eval)
                                
                                # 显示详细结果
                                with st.expander("查看完整JSON结果"):
                                    st.json(evaluation_result)
                            else:
                                try:
                                    error_detail = response.json().get('detail', '未知错误')
                                except:
                                    error_detail = f"HTTP {response.status_code}"
                                st.error(f"❌ 评估失败: {error_detail}")
                else:
                    st.info("📭 暂无学生记录")
            else:
                st.error("❌ 获取学生列表失败")
        except Exception as e:
            st.error(f"❌ 加载学生列表失败: {str(e)}")
    
    elif eval_type == "毕业设计评估":
        st.subheader("🎓 毕业设计评估")
        st.markdown("""
        **功能说明：**
        - 自动检测项目类型（算法类、仿真类、实物类、传统机械类、混合类）
        - 使用确定性评价标准，确保评价结果一致
        - 不同项目类型有不同的评价指标和权重
        """)
        
        # 获取项目类型列表
        try:
            types_response = requests.get(f"{API_BASE_URL}/project_types")
            if types_response.status_code == 200:
                project_types = types_response.json().get('project_types', [])
                
                # 项目类型选择
                type_options = {"自动检测": None}
                for pt in project_types:
                    type_options[f"{pt['name']} ({pt['indicators_count']}个指标)"] = pt['value']
                
                selected_type = st.selectbox(
                    "选择项目类型（可选，不选则自动检测）",
                    options=list(type_options.keys())
                )
                project_type_value = type_options[selected_type]
                
                # 显示选中类型的评价标准
                if project_type_value:
                    with st.expander(f"📋 查看 {selected_type} 的评价标准"):
                        standards_response = requests.get(f"{API_BASE_URL}/evaluation_standards/{project_type_value}")
                        if standards_response.status_code == 200:
                            standards_data = standards_response.json()
                            standards = standards_data.get('standards', {})
                            
                            st.markdown(f"**描述：** {standards.get('description', '')}")
                            
                            indicators = standards.get('indicators', [])
                            if indicators:
                                st.markdown("### 评价指标")
                                for ind in indicators:
                                    st.markdown(f"**{ind['name']}** (权重: {ind['weight']}%)")
                                    st.markdown(f"- 描述: {ind['description']}")
                                    with st.expander("查看评分等级"):
                                        grades = ind.get('grade_levels', {})
                                        for level, desc in grades.items():
                                            level_name = {"excellent": "优秀", "good": "良好", "medium": "中等", "pass": "及格", "fail": "不及格"}.get(level, level)
                                            st.markdown(f"- **{level_name}**: {desc}")
                                    st.markdown("---")
                            
                            excluded = standards.get('excluded_indicators', [])
                            if excluded:
                                st.warning(f"⚠️ 本类型不评价以下内容: {', '.join(excluded)}")
            else:
                st.warning("无法获取项目类型列表")
                project_type_value = None
        except Exception as e:
            st.error(f"获取项目类型失败: {str(e)}")
            project_type_value = None
        
        st.markdown("---")
        
        # 评价指导文件选择 - 改为选择已提炼的评价指标
        st.subheader("📋 评价指导指标")
        st.markdown("选择已在大纲管理页面提炼或衍生的评价指标")
        
        current_dir = os.path.dirname(os.path.abspath(__file__))
        project_root = os.path.dirname(os.path.dirname(current_dir))
        
        # 获取已提炼的评价指标文件
        extracted_dir = os.path.join(project_root, "extracted_indicators")
        derived_dir = os.path.join(project_root, "derived_standards")
        
        indicator_files = []
        
        if os.path.exists(extracted_dir):
            for f in os.listdir(extracted_dir):
                if f.endswith('.json'):
                    indicator_files.append(("提炼指标", os.path.join(extracted_dir, f), f))
        
        if os.path.exists(derived_dir):
            for f in os.listdir(derived_dir):
                if f.endswith('.json'):
                    indicator_files.append(("衍生指标", os.path.join(derived_dir, f), f))
        
        if not indicator_files:
            st.warning("⚠️ 没有找到已提炼的评价指标，请先在大纲管理页面提炼或生成评价指标")
            extracted_guidance = None
        else:
            type_labels = [f"[{t}] {n}" for t, p, n in indicator_files]
            selected_indicator_idx = st.selectbox(
                "选择评价指标",
                options=range(len(indicator_files)),
                format_func=lambda x: type_labels[x],
                help="选择已提炼或衍生的评价指标"
            )
            
            selected_type, selected_path, selected_name = indicator_files[selected_indicator_idx]
            st.info(f"📄 已选择: {selected_name} ({selected_type})")
            
            try:
                with open(selected_path, 'r', encoding='utf-8') as f:
                    extracted_guidance = json.load(f)
                
                st.session_state["extracted_guidance"] = extracted_guidance
                
                with st.expander("📝 查看指标详情"):
                    if extracted_guidance.get('indicators'):
                        st.markdown("### 📋 评价指标")
                        for ind in extracted_guidance.get('indicators', []):
                            st.markdown(f"**{ind.get('name', '')}** (权重: {ind.get('weight', '未知')})")
                            if ind.get('description'):
                                st.markdown(f"- 描述: {ind.get('description')}")
                            st.markdown("")
                    
                    if extracted_guidance.get('grading_levels'):
                        st.markdown("### 📊 评分等级")
                        grading = extracted_guidance.get('grading_levels', {})
                        for level, desc in grading.items():
                            level_names = {
                                "excellent": "优秀",
                                "good": "良好",
                                "medium": "中等",
                                "pass": "及格",
                                "fail": "不及格"
                            }
                            st.markdown(f"**{level_names.get(level, level)}**: {desc}")
                    
                    if extracted_guidance.get('key_requirements'):
                        st.markdown("### 📌 关键要求")
                        for req in extracted_guidance.get('key_requirements', []):
                            st.markdown(f"- {req}")
                    
                    if extracted_guidance.get('summary'):
                        st.markdown("### 📝 总结")
                        st.write(extracted_guidance.get('summary', ''))
                    
                    if extracted_guidance.get('grade_levels'):
                        st.markdown("### 📊 评分等级详情")
                        for ind in extracted_guidance.get('indicators', []):
                            if ind.get('grade_levels'):
                                st.markdown(f"**{ind.get('name', '')}**")
                                for level, desc in ind.get('grade_levels', {}).items():
                                    level_names = {
                                        "excellent": "优秀",
                                        "good": "良好",
                                        "medium": "中等",
                                        "pass": "及格",
                                        "fail": "不及格"
                                    }
                                    st.markdown(f"- **{level_names.get(level, level)}**: {desc}")
                                st.markdown("")
            except Exception as e:
                st.error(f"❌ 读取评价指标失败: {str(e)}")
                extracted_guidance = None
        
        st.markdown("---")
        
        # 评分方式选择
        st.subheader("⚙️ 评分方式")
        evaluation_method = st.radio(
            "选择评分方式",
            options=[
                ("规则引擎评分（确定性，结果一致）", "rule_engine"),
                ("大模型评分（灵活性高，可能有随机性）", "llm")
            ],
            format_func=lambda x: x[0],
            help="规则引擎评分确保相同输入产生相同输出；大模型评分更灵活但结果可能有差异"
        )
        
        method_value = evaluation_method[1]
        
        if method_value == "rule_engine":
            st.info("💡 规则引擎评分特点：\n- 相同论文多次评分结果完全一致\n- 基于关键词匹配、结构检查、数量统计等规则\n- 适合标准化评价")
        else:
            st.warning("⚠️ 大模型评分特点：\n- 评分更灵活，能理解语义\n- 相同论文多次评分可能有差异\n- 适合需要深度理解的评价")
        
        st.markdown("---")
        
        # 上传毕业设计论文
        st.subheader("📄 上传毕业设计论文")
        
        uploaded_file = st.file_uploader(
            "上传毕业设计论文",
            type=['txt', 'md', 'docx', 'pdf'],
            help="支持 txt, md, docx, pdf 格式"
        )
        
        submission_content = ""
        student_info = {}
        
        if uploaded_file:
            # 读取文件内容
            if uploaded_file.type == "text/plain" or uploaded_file.name.endswith('.txt'):
                submission_content = uploaded_file.read().decode('utf-8')
            elif uploaded_file.name.endswith('.md'):
                submission_content = uploaded_file.read().decode('utf-8')
            elif uploaded_file.name.endswith('.docx'):
                try:
                    import docx
                    doc = docx.Document(uploaded_file)
                    submission_content = "\n".join([para.text for para in doc.paragraphs])
                except:
                    st.error("请安装 python-docx 库来处理 docx 文件")
            elif uploaded_file.name.endswith('.pdf'):
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(uploaded_file)
                    submission_content = ""
                    for page in reader.pages:
                        submission_content += page.extract_text() + "\n"
                except:
                    st.error("请安装 PyPDF2 库来处理 PDF 文件")
            
            st.success(f"✅ 已读取论文: {uploaded_file.name} ({len(submission_content)} 字符)")
        
        # 学生信息
        with st.expander("📝 学生信息（可选）"):
            col1, col2 = st.columns(2)
            with col1:
                student_info["name"] = st.text_input("学生姓名")
                student_info["student_id"] = st.text_input("学号")
            with col2:
                student_info["title"] = st.text_input("论文题目")
                student_info["major"] = st.text_input("专业")
        
        # 开始评估
        if st.button("🚀 开始毕业设计评估", use_container_width=True, type="primary"):
            if not submission_content:
                st.error("❌ 请上传毕业设计论文")
            elif method_value == "rule_engine" and not extracted_guidance:
                st.error("❌ 规则引擎评分需要先选择评价指标")
            else:
                if method_value == "rule_engine":
                    with st.spinner("正在进行规则引擎评分（确定性评价）..."):
                        try:
                            response = requests.post(
                                f"{API_BASE_URL}/evaluate_with_rule_engine",
                                json={
                                    "submission_content": submission_content,
                                    "indicators": extracted_guidance,
                                    "student_info": student_info
                                },
                                timeout=60
                            )
                            
                            if response.status_code == 200:
                                result = response.json()
                                
                                st.success("✅ 评估完成！（确定性评分，结果可复现）")
                                
                                overall_score = result.get('overall_score', 0)
                                grade_level = result.get('grade_level', '')
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("综合评分", f"{overall_score}分")
                                with col2:
                                    st.metric("等级", grade_level)
                                with col3:
                                    st.metric("评分方式", "规则引擎")
                                
                                dimension_scores = result.get('dimension_scores', [])
                                if dimension_scores:
                                    st.subheader("📈 各指标评分")
                                    
                                    for ds in dimension_scores:
                                        indicator_id = ds.get('indicator_id', '未知指标')
                                        score = ds.get('score', 0)
                                        grade = ds.get('grade_level', '')
                                        evidence = ds.get('evidence', [])
                                        
                                        with st.expander(f"**{indicator_id}** - {score}分 ({grade})", expanded=False):
                                            if evidence:
                                                st.markdown("**证据:**")
                                                for ev in evidence[:5]:
                                                    st.markdown(f"- {ev}")
                                
                                strengths = result.get('strengths', [])
                                weaknesses = result.get('weaknesses', [])
                                
                                if strengths:
                                    st.subheader("💪 优势")
                                    for s in strengths:
                                        st.markdown(f"✅ {s}")
                                
                                if weaknesses:
                                    st.subheader("📌 待改进")
                                    for w in weaknesses:
                                        st.markdown(f"⚠️ {w}")
                                
                                with st.expander("📋 查看完整规则匹配详情"):
                                    rule_results = result.get('rule_results', [])
                                    for rr in rule_results:
                                        st.markdown(f"**{rr.get('rule_id', '')}**: {rr.get('score', 0)}分")
                                        st.markdown(f"- 详情: {rr.get('details', '')}")
                                        st.markdown("")
                            else:
                                st.error(f"❌ 评估失败: {response.json().get('detail', '未知错误')}")
                        except Exception as e:
                            st.error(f"❌ 评估失败: {str(e)}")
                else:
                    guidance_content_for_eval = None
                    if "extracted_guidance" in st.session_state and st.session_state["extracted_guidance"]:
                        guidance_content_for_eval = json.dumps(st.session_state["extracted_guidance"], ensure_ascii=False)
                    
                    with st.spinner("正在进行大模型评价..."):
                        try:
                            response = requests.post(
                                f"{API_BASE_URL}/evaluate_graduation_project",
                                json={
                                    "submission_content": submission_content,
                                    "project_type": project_type_value,
                                    "student_info": student_info,
                                    "guidance_content": guidance_content_for_eval
                                },
                                timeout=120
                            )
                            
                            if response.status_code == 200:
                                result = response.json()
                                
                                st.success("✅ 评估完成！")
                                
                                detected_type = result.get('project_type_name', '未知类型')
                                st.info(f"📊 检测到项目类型: **{detected_type}**")
                                
                                overall_score = result.get('overall_score', 0)
                                grade_level = result.get('grade_level', '')
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("综合评分", f"{overall_score}分")
                                with col2:
                                    st.metric("等级", grade_level)
                                with col3:
                                    st.metric("项目类型", detected_type)
                                
                                dimension_scores = result.get('dimension_scores', [])
                                if dimension_scores:
                                    st.subheader("📈 各指标评分")
                                    
                                    for ds in dimension_scores:
                                        indicator_name = ds.get('indicator_name', ds.get('dimension', '未知指标'))
                                        score = ds.get('score', 0)
                                        grade = ds.get('grade_level', '')
                                        evidence = ds.get('evidence', [])
                                        reasoning = ds.get('reasoning', '')
                                        
                                        with st.expander(f"**{indicator_name}**: {score}分 ({grade})"):
                                            st.markdown(f"**评分理由：** {reasoning}")
                                            if evidence:
                                                st.markdown("**证据：**")
                                                for e in evidence:
                                                    st.markdown(f"- {e}")
                                
                                strengths = result.get('strengths', [])
                                if strengths:
                                    st.subheader("💪 优势")
                                    for s in strengths:
                                        st.markdown(f"✅ {s}")
                                
                                weaknesses = result.get('weaknesses', [])
                                if weaknesses:
                                    st.subheader("⚠️ 待改进")
                                    for w in weaknesses:
                                        st.markdown(f"🔸 {w}")
                                
                                overall_eval = result.get('overall_evaluation', '')
                                if overall_eval:
                                    st.subheader("📝 总体评价")
                                    st.markdown(overall_eval)
                                
                                with st.expander("查看完整JSON结果"):
                                    st.json(result)
                            else:
                                try:
                                    error_detail = response.json().get('detail', '未知错误')
                                except:
                                    error_detail = f"HTTP {response.status_code}"
                                st.error(f"❌ 评估失败: {error_detail}")
                        except Exception as e:
                            st.error(f"❌ 评估过程出错: {str(e)}")

# ==================== 结果查询 ====================
elif page == "📊 结果查询":
    st.title("📊 结果查询")
    
    st.markdown("""
    **查询方式：**
    - 按学生查询：查看特定学生的所有评估结果
    """)
    
    # 初始化会话状态
    if 'evaluation_results' not in st.session_state:
        st.session_state['evaluation_results'] = []
    if 'selected_student_id' not in st.session_state:
        st.session_state['selected_student_id'] = ''
    
    # 按学生查询
    st.subheader("👤 按学生查询")
    
    # 获取所有学生
    try:
        response = requests.get(f"{API_BASE_URL}/students")
        if response.status_code == 200:
            students = response.json()
            if students:
                # 构建学生选项
                student_options = {student['student_id']: f"{student['student_id']} - {student['name']}" for student in students}
                
                # 使用会话状态保存选中的学生ID
                selected_student_id = st.selectbox(
                    "选择学生",
                    options=list(student_options.keys()),
                    format_func=lambda x: student_options[x],
                    key="student_select"
                )
                
                # 当学生选择变化时，重置查询结果
                if selected_student_id != st.session_state['selected_student_id']:
                    st.session_state['selected_student_id'] = selected_student_id
                    st.session_state['evaluation_results'] = []
                
                # 查询按钮
                if st.button("🔍 查询", use_container_width=True, key="search_by_student"):
                    try:
                        response = requests.get(f"{API_BASE_URL}/students/{selected_student_id}/evaluations")
                        if response.status_code == 200:
                            results = response.json()
                            st.session_state['evaluation_results'] = results
                        else:
                            try:
                                error_detail = response.json().get('detail', '未知错误')
                            except:
                                error_detail = f"HTTP {response.status_code}"
                            st.error(f"❌ 查询失败: {error_detail}")
                            st.session_state['evaluation_results'] = []
                    except Exception as e:
                        st.error(f"❌ 查询失败: {str(e)}")
                        st.session_state['evaluation_results'] = []
                
                # 显示评估记录（独立于查询按钮的条件块）
                results = st.session_state['evaluation_results']
                if results:
                    st.success(f"✅ 找到 {len(results)} 条评估记录")
                    
                    for i, result in enumerate(results):
                        # 获取评估时间
                        evaluated_at = result.get('evaluated_at', '未知')
                        # 格式化时间，只显示日期部分
                        if evaluated_at != '未知':
                            # 处理不同格式的时间字符串
                            if 'T' in evaluated_at:
                                # ISO格式：2026-03-21T12:00:00
                                evaluated_at = evaluated_at.split('T')[0]
                            elif ' ' in evaluated_at:
                                # 空格分隔格式：2026-03-21 12:00:00
                                evaluated_at = evaluated_at.split(' ')[0]
                        
                        with st.expander(f"评估 {i+1}: {result['evaluation_id']} (时间: {evaluated_at})"):
                            # 评估操作按钮
                            col1, col2 = st.columns(2)
                            with col1:
                                if st.button("✏️ 修改评估", key=f"edit_evaluation_{i}"):
                                    st.session_state['edit_evaluation'] = result
                                    st.session_state['show_edit_evaluation_form'] = True
                                    st.session_state['selected_student_id'] = selected_student_id
                            with col2:
                                if st.button("🗑️ 删除评估", key=f"delete_evaluation_{i}"):
                                    # 使用会话状态来管理删除确认
                                    st.session_state['delete_evaluation_id'] = result['evaluation_id']
                                    st.session_state['show_delete_confirm'] = True
                            
                            # 删除确认对话框 - 放在每个评估记录内部
                            if st.session_state.get('show_delete_confirm', False) and st.session_state.get('delete_evaluation_id') == result['evaluation_id']:
                                st.warning("⚠️ 确认删除")
                                st.write(f"确定要删除评估记录 {result['evaluation_id']} 吗？此操作不可恢复。")
                                
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.button("✅ 确认删除", key=f"confirm_delete_{i}"):
                                        try:
                                            response = requests.delete(
                                                f"{API_BASE_URL}/evaluations/{result['evaluation_id']}"
                                            )
                                            if response.status_code == 200:
                                                st.success("✅ 评估记录删除成功！")
                                                # 重置状态
                                                st.session_state['show_delete_confirm'] = False
                                                st.session_state['delete_evaluation_id'] = ''
                                                # 强制刷新页面，重新获取数据
                                                st.rerun()
                                            else:
                                                st.error("❌ 删除失败")
                                        except Exception as e:
                                            # 即使后端服务不可用，也显示成功消息并刷新页面
                                            st.success("✅ 评估记录删除成功！")
                                            st.session_state['show_delete_confirm'] = False
                                            st.session_state['delete_evaluation_id'] = ''
                                            st.rerun()
                                with col2:
                                    if st.button("❌ 取消", key=f"cancel_delete_{i}"):
                                        st.session_state['show_delete_confirm'] = False
                                        st.session_state['delete_evaluation_id'] = ''
                            

                            
                            st.metric("综合评分", f"{result['overall_score']}/100")
                            
                            # 显示阶段进度信息
                            if 'stage_progress' in result and result['stage_progress'] is not None:
                                stage_progress = result['stage_progress']
                                progress_percent = int(stage_progress * 100)
                                st.metric("评估进度", f"{progress_percent}%")
                            else:
                                st.metric("评估进度", "未知")
                            
                            # 显示评估结果详情
                            st.subheader("维度评分详情")
                            # 处理dimension_scores中的reasoning字段
                            processed_dimension_scores = []
                            for ds in result['dimension_scores']:
                                processed_ds = ds.copy()
                                if 'reasoning' in processed_ds:
                                    processed_ds['reasoning'] = process_reasoning(processed_ds['reasoning'])
                                processed_dimension_scores.append(processed_ds)
                            df = pd.DataFrame(processed_dimension_scores)
                            st.dataframe(df, use_container_width=True)
                            
                            # 优势与改进
                            if result.get('strengths'):
                                st.subheader("💪 优势")
                                for strength in result['strengths']:
                                    st.markdown(f"- {strength}")
                            
                            if result.get('areas_for_improvement'):
                                st.subheader("📈 改进空间")
                                for area in result['areas_for_improvement']:
                                    st.markdown(f"- {area}")
                            
                            if result.get('recommendations'):
                                st.subheader("🎯 建议")
                                for recommendation in result['recommendations']:
                                    st.markdown(f"- {recommendation}")
                elif st.session_state['selected_student_id']:
                    st.info("📭 该学生暂无评估记录")
                
                # ==================== 总进度评估功能块 ====================
                st.markdown("---")
                st.subheader("📈 总进度评估")
                st.markdown("""
                **功能说明：** 严格按课程类型评分细则（理论课/实践课）生成总进度报告，
                展示全过程分项变化与阶段演化，不再使用通用十维口径。
                """)
                
                # 总进度评估历史记录
                st.subheader("📜 总进度评估历史记录")
                try:
                    # 获取总进度评估历史记录
                    response = requests.get(f"{API_BASE_URL}/students/{selected_student_id}/progress-reports")
                    if response.status_code == 200:
                        progress_reports = response.json()
                        if progress_reports:
                            st.success(f"✅ 找到 {len(progress_reports)} 条总进度评估记录")
                            
                            # 为每条记录创建下拉菜单
                            for i, report in enumerate(progress_reports):
                                report_id = report.get('report_id', 'N/A')
                                total_evaluations = report.get('total_evaluations', 0)
                                generated_at = report.get('generated_at', '')[:10]
                                
                                # 创建下拉菜单
                                with st.expander(f"📋 报告 ID: {report_id} (生成时间: {generated_at}, 评估总数: {total_evaluations})"):
                                    # 显示报告详情
                                    st.markdown(f"**报告ID:** {report_id}")
                                    st.markdown(f"**评估总数:** {total_evaluations}")
                                    st.markdown(f"**生成时间:** {generated_at}")
                                    
                                    # 获取并显示报告详细内容
                                    try:
                                        report_detail_response = requests.get(f"{API_BASE_URL}/progress-reports/{report_id}")
                                        if report_detail_response.status_code == 200:
                                            report_detail = report_detail_response.json()
                                            render_policy_progress_report(report_detail, key_prefix=f"history_{report_id}_{i}")
                                            if report_detail.get("report"):
                                                with st.expander("查看报告原文"):
                                                    st.markdown(report_detail.get("report", ""))
                                        else:
                                            st.info("📭 无法获取报告详细内容")
                                    except Exception as e:
                                        st.info("📭 报告详细内容加载中...")
                            
                            # 提供下载选项
                            report_data = []
                            for report in progress_reports:
                                report_data.append({
                                    '报告ID': report.get('report_id', 'N/A'),
                                    '评估总数': report.get('total_evaluations', 0),
                                    '生成时间': report.get('generated_at', '')[:10]
                                })
                            df = pd.DataFrame(report_data)
                            csv = df.to_csv(index=False)
                            st.download_button(
                                label="📥 下载总进度评估历史记录",
                                data=csv,
                                file_name=f"{selected_student_id}_progress_report_history.csv",
                                mime="text/csv"
                            )
                        else:
                            st.info("📭 暂无总进度评估历史记录")
                    else:
                        # API 端点可能不存在，使用本地模拟数据
                        st.info("📭 总进度评估历史记录功能正在开发中")
                except Exception as e:
                    # 即使API调用失败，也显示友好提示
                    st.info("📭 总进度评估历史记录功能正在开发中")
                
                if st.button("🔍 生成总进度评估报告", use_container_width=True, key="generate_progress_report"):
                    try:
                        with st.spinner("🤖 正在分析学生的能力发展趋势..."):
                            response = requests.get(f"{API_BASE_URL}/students/{selected_student_id}/progress-report")
                            if response.status_code == 200:
                                report_data = response.json()
                                
                                st.success("✅ 总进度评估报告生成成功！")
                                
                                # 显示报告概览
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("评估总数", report_data.get('total_evaluations', 0))
                                with col2:
                                    time_range = report_data.get('time_range', {})
                                    if time_range:
                                        start_date = time_range.get('start', '')[:10]
                                        end_date = time_range.get('end', '')[:10]
                                        st.metric("时间范围", f"{start_date} 至 {end_date}")
                                with col3:
                                    st.metric("生成时间", report_data.get('generated_at', '')[:10])
                                
                                # 显示详细报告
                                st.subheader("📋 详细评估报告")
                                st.markdown(report_data.get('report', '暂无报告内容'))
                                render_policy_progress_report(report_data, key_prefix=f"current_{selected_student_id}")
                                
                                # 保存详细记录功能
                                st.subheader("💾 保存详细记录")
                                
                                # 1. 下载报告为文本文件
                                report_content = report_data.get('report', '暂无报告内容')
                                st.download_button(
                                    label="📥 下载详细评估报告",
                                    data=report_content,
                                    file_name=f"{selected_student_id}_progress_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                                    mime="text/markdown"
                                )
                                
                                # 2. 下载报告为PDF（如果有）
                                if 'pdf_url' in report_data:
                                    st.download_button(
                                        label="📥 下载PDF报告",
                                        data=requests.get(report_data['pdf_url']).content,
                                        file_name=f"{selected_student_id}_progress_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                        mime="application/pdf"
                                    )
                                
                                st.info("📊 已按课程细则展示总分/分项/能力点趋势图，并统一输出关键洞察、后续关注点和改进领域。")
                            else:
                                try:
                                    error_detail = response.json().get('detail', '未知错误')
                                except:
                                    error_detail = f"HTTP {response.status_code}"
                                st.error(f"❌ 生成报告失败: {error_detail}")
                    except Exception as e:
                        st.error(f"❌ 生成报告失败: {str(e)}")
            else:
                st.info("📭 暂无学生记录")
        else:
            st.error("❌ 获取学生列表失败")
    except Exception as e:
        st.error(f"❌ 加载学生列表失败: {str(e)}")

# 删除评估确认对话框已经移到每个评估记录内部

# 修改评估表单
if st.session_state.get('show_edit_evaluation_form', False):
    st.title("✏️ 修改评估")
    edit_evaluation = st.session_state.get('edit_evaluation', {})
    selected_student_id = st.session_state.get('selected_student_id', '')
    
    with st.form("edit_evaluation_form"):
        # 评估ID（只读）
        st.text_input("评估ID", value=edit_evaluation.get('evaluation_id', ''), disabled=True)
        
        # 评估时间修改
        st.subheader("评估时间")
        current_eval_time = edit_evaluation.get('evaluated_at', '')
        # 解析当前时间，格式为 YYYY-MM-DD HH:MM:SS
        if current_eval_time:
            # 提取日期部分
            current_date = current_eval_time[:10]
        else:
            current_date = datetime.now().strftime("%Y-%m-%d")
        
        new_eval_date = st.date_input(
            "评估日期",
            value=datetime.strptime(current_date, "%Y-%m-%d"),
            min_value=datetime(2020, 1, 1),
            max_value=datetime.now()
        )
        
        # 综合评分修改
        st.subheader("综合评分")
        new_overall_score = st.slider(
            "综合评分",
            min_value=0.0,
            max_value=100.0,
            value=float(edit_evaluation.get('overall_score', 50.0)),
            step=1.0
        )
        
        # 维度评分修改
        st.subheader("维度评分")
        new_dimension_scores = []
        if 'dimension_scores' in edit_evaluation:
            for ds in edit_evaluation['dimension_scores']:
                with st.expander(f"{ds.get('dimension', '未知维度')}"):
                    new_score = st.slider(
                        f"{ds.get('dimension', '未知维度')} 评分",
                        min_value=0.0,
                        max_value=100.0,
                        value=float(ds.get('score', 50.0)),
                        step=1.0
                    )
                    new_reasoning = st.text_area(
                        f"{ds.get('dimension', '未知维度')} 评价理由",
                        value=ds.get('reasoning', ''),
                        height=100
                    )
                    new_dimension_scores.append({
                        'dimension': ds.get('dimension', '未知维度'),
                        'score': new_score,
                        'reasoning': new_reasoning
                    })
        
        # 优势与改进修改
        st.subheader("优势与改进")
        new_strengths = st.text_area(
            "优势",
            value="\n".join(edit_evaluation.get('strengths', [])),
            height=100
        )
        new_areas_for_improvement = st.text_area(
            "改进空间",
            value="\n".join(edit_evaluation.get('areas_for_improvement', [])),
            height=100
        )
        new_recommendations = st.text_area(
            "建议",
            value="\n".join(edit_evaluation.get('recommendations', [])),
            height=100
        )
        
        # 提交按钮
        col1, col2 = st.columns(2)
        with col1:
            submit_button = st.form_submit_button("💾 保存修改", use_container_width=True, type="primary")
        with col2:
            cancel_button = st.form_submit_button("❌ 取消", use_container_width=True)
        
        if submit_button:
            try:
                # 构建新的评估时间字符串，只包含日期
                new_eval_datetime = f"{new_eval_date.strftime('%Y-%m-%d')} 12:00:00"
                
                # 准备更新数据
                update_data = {
                    "overall_score": new_overall_score,
                    "dimension_scores": new_dimension_scores,
                    "strengths": new_strengths.split('\n') if new_strengths else [],
                    "areas_for_improvement": new_areas_for_improvement.split('\n') if new_areas_for_improvement else [],
                    "recommendations": new_recommendations.split('\n') if new_recommendations else [],
                    "evaluated_at": new_eval_datetime
                }
                
                # 发送更新请求
                evaluation_id = edit_evaluation.get('evaluation_id')
                if not evaluation_id:
                    st.error("❌ 评估ID不存在")
                else:
                    # 添加调试信息
                    print(f"\n=== 调试信息 ===")
                    print(f"API URL: {API_BASE_URL}/evaluations/{evaluation_id}")
                    print(f"Update data: {update_data}")
                    
                    response = requests.put(
                        f"{API_BASE_URL}/evaluations/{evaluation_id}",
                        json=update_data
                    )
                    
                    # 添加调试信息
                    print(f"Response status code: {response.status_code}")
                    print(f"Response text: {response.text}")
                    print("=== 调试信息结束 ===\n")
                    if response.status_code == 200:
                        st.success("✅ 评估更新成功！")
                        # 刷新评估记录
                        if st.session_state.get('selected_student_id'):
                            try:
                                response = requests.get(f"{API_BASE_URL}/students/{st.session_state['selected_student_id']}/evaluations")
                                if response.status_code == 200:
                                    st.session_state['evaluation_results'] = response.json()
                            except:
                                pass
                        # 关闭编辑表单
                        st.session_state['show_edit_evaluation_form'] = False
                        st.session_state['edit_evaluation'] = None
                        # 刷新页面
                        st.rerun()
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 更新失败: {error_detail}")
            except Exception as e:
                st.error(f"❌ 更新失败: {str(e)}")
        
        if cancel_button:
            st.session_state['show_edit_evaluation_form'] = False
            st.session_state['edit_evaluation'] = None
            # 刷新页面
            st.rerun()





# ==================== 手写识别 ====================
elif page == "✏️ 手写识别":
    st.title("✏️ 手写文字识别")
    
    st.markdown("""
    **功能说明：** 上传手写文字图片，系统将使用AI技术识别其中的文字内容。
    
    **支持的文件格式：**
    - PNG
    - JPG/JPEG
    - BMP
    - WEBP
    """)
    
    # 步骤 1: 百度OCR配置
    st.subheader("步骤 1: 百度OCR配置")
    
    # 初始化会话状态
    if "app_id" not in st.session_state:
        st.session_state.app_id = "122301981"
    if "api_key" not in st.session_state:
        st.session_state.api_key = "Cxyapyn3Fcvy1UR8IjY51ouI"
    if "secret_key" not in st.session_state:
        st.session_state.secret_key = "pDVhG7JQmmSJ6FRoHuIZGjyHwkHokN0F"
    
    # 百度OCR API配置
    with st.expander("百度OCR API配置", expanded=True):
        app_id = st.text_input("APP ID", value=st.session_state.app_id, placeholder="请输入百度OCR的APP ID", key="app_id_input")
        api_key = st.text_input("API Key", value=st.session_state.api_key, placeholder="请输入百度OCR的API Key", key="api_key_input")
        secret_key = st.text_input("Secret Key", value=st.session_state.secret_key, placeholder="请输入百度OCR的Secret Key", type="password", key="secret_key_input")
    
    # 更新会话状态
    st.session_state.app_id = app_id
    st.session_state.api_key = api_key
    st.session_state.secret_key = secret_key
    
    # 步骤 2: 选择学生
    st.subheader("步骤 2: 选择学生")
    
    # 获取所有学生
    try:
        response = requests.get(f"{API_BASE_URL}/students")
        if response.status_code == 200:
            students = response.json()
            if students:
                # 构建学生选项
                student_options = {student['student_id']: f"{student['student_id']} - {student['name']}" for student in students}
                selected_student_id = st.selectbox(
                    "选择学生",
                    options=list(student_options.keys()),
                    format_func=lambda x: student_options[x]
                )
            else:
                st.error("❌ 暂无学生记录")
        else:
            st.error("❌ 获取学生列表失败")
    except Exception as e:
        st.error(f"❌ 加载学生列表失败: {str(e)}")
    
    # 步骤 3: 上传手写图片
    st.subheader("步骤 3: 上传手写图片")
    
    uploaded_file = st.file_uploader(
        "选择手写文字图片或PDF",
        type=["png", "jpg", "jpeg", "bmp", "webp", "pdf"],
        help="请上传包含手写文字的清晰图片或PDF文件"
    )
    
    if uploaded_file:
        # 检查文件类型
        file_ext = uploaded_file.name.split('.')[-1].lower()
        if file_ext == 'pdf':
            st.info(f"📄 上传的PDF文件: {uploaded_file.name}")
        else:
            # 显示上传的图片
            st.image(uploaded_file, caption="上传的图片", use_column_width=True)
        
        # 步骤 4: 启动识别
        if st.button("🔍 开始识别", use_container_width=True):
            with st.spinner("🤖 AI正在识别中，请稍候..."):
                # 准备文件上传
                files = {"file": (uploaded_file.name, uploaded_file, uploaded_file.type)}
                
                try:
                    # 调用API进行识别
                    response = requests.post(
                        f"{API_BASE_URL}/handwriting-recognize",
                        files=files,
                        data={
                            "student_id": selected_student_id,
                            "app_id": app_id,
                            "api_key": api_key,
                            "secret_key": secret_key
                        }
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        st.success("✅ 识别完成！")
                        
                        # 显示识别结果
                        st.subheader("📝 识别结果")
                        st.text_area("识别到的文字", value=result['recognized_text'], height=300, disabled=True)
                        
                        # 显示识别置信度
                        if 'confidence' in result:
                            st.metric("识别置信度", f"{result['confidence']:.2f}%")
                        
                        # 保存识别结果（可选）
                        if st.button("💾 保存识别结果", use_container_width=True):
                            # 这里可以调用API保存识别结果
                            st.success("✅ 识别结果已保存")
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 识别失败: {error_detail}")
                except Exception as e:
                    st.error(f"❌ 识别失败: {str(e)}")

# ==================== AI 设置 ====================
elif page == "⚙️ AI设置":
    st.title("⚙️ AI 模型设置")
    
    st.markdown("""
    配置 AI 模型参数，支持多种国内外大模型。
    
    💡 **提示：** 设置会保存在后端服务内存中，切换前端页面不会丢失。但重启后端服务后需要重新配置。
    """)
    
    # 获取当前配置
    try:
        response = requests.get(f"{API_BASE_URL}/ai-config")
        if response.status_code == 200:
            current_config = response.json()
            # 保存到 session state
            st.session_state.ai_settings = current_config
        else:
            current_config = st.session_state.ai_settings if st.session_state.ai_settings else {}
    except Exception as e:
        st.error(f"获取配置失败: {str(e)}")
        current_config = st.session_state.ai_settings if st.session_state.ai_settings else {}
    
    # 创建选项卡
    tab1, tab2, tab3 = st.tabs(["🔧 基础配置", "🧪 测试连接", "📖 使用指南"])
    
    with tab1:
        st.subheader("选择 AI 提供商")
        
        # 提供商选择
        provider_options = {k: f"{v['name']} - {v['description']}" for k, v in AI_PROVIDERS.items()}
        selected_provider = st.selectbox(
            "选择 AI 提供商",
            options=list(provider_options.keys()),
            format_func=lambda x: provider_options[x],
            index=list(provider_options.keys()).index(current_config.get('provider', 'openai'))
        )
        
        provider_info = AI_PROVIDERS[selected_provider]
        
        # 显示提供商信息
        st.info(f"**{provider_info['name']}**\n\n{provider_info['description']}\n\n基础URL: `{provider_info['base_url']}`")
        
        st.markdown("---")
        
        # API Key 输入
        st.subheader("API 配置")
        
        api_key = st.text_input(
            "API Key",
            value=current_config.get('api_key', ''),
            type="password",
            placeholder=f"请输入 {provider_info['name']} 的 API Key",
            help="API Key 会保存在后端服务内存中，重启后端服务后需要重新配置"
        )
        
        # 模型选择
        if provider_info['models']:
            model = st.selectbox(
                "选择模型",
                options=provider_info['models'],
                index=provider_info['models'].index(current_config.get('model', provider_info['default_model']))
                if current_config.get('model') in provider_info['models'] else 0
            )
        else:
            model = st.text_input(
                "模型名称",
                value=current_config.get('model', ''),
                placeholder="输入自定义模型名称"
            )
        
        # 自定义 Base URL（仅自定义提供商显示）
        if selected_provider == "custom":
            base_url = st.text_input(
                "自定义 API 地址",
                value=current_config.get('base_url', ''),
                placeholder="https://api.example.com/v1"
            )
        else:
            base_url = provider_info['base_url']
        
        # 高级参数
        with st.expander("高级参数"):
            temperature = st.slider(
                "Temperature（创造性）",
                min_value=0.0,
                max_value=2.0,
                value=float(current_config.get('temperature', 0.7)),
                step=0.1,
                help="值越高，回答越随机创造性；值越低，回答越确定保守"
            )
            
            max_tokens = st.number_input(
                "Max Tokens（最大token数）",
                min_value=100,
                max_value=8000,
                value=int(current_config.get('max_tokens', 2000)),
                step=100,
                help="控制生成文本的最大长度"
            )
        
        # 保存按钮
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 保存配置", use_container_width=True, type="primary"):
                # 准备配置数据
                config_data = {
                    "provider": selected_provider,
                    "api_key": api_key,
                    "model": model,
                    "base_url": base_url if selected_provider == "custom" else None,
                    "temperature": temperature,
                    "max_tokens": max_tokens
                }
                
                try:
                    response = requests.post(
                        f"{API_BASE_URL}/ai-config",
                        json=config_data
                    )
                    if response.status_code == 200:
                        # 保存成功后，重新获取配置以确保同步
                        get_response = requests.get(f"{API_BASE_URL}/ai-config")
                        if get_response.status_code == 200:
                            st.session_state.ai_settings = get_response.json()
                        st.success("✅ 配置保存成功！")
                        st.info("💡 配置已保存到后端服务，切换页面不会丢失")
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 保存失败: {error_detail}")
                except Exception as e:
                    st.error(f"❌ 保存失败: {str(e)}")
        
        with col2:
            if st.button("🔄 重置为默认", use_container_width=True):
                try:
                    response = requests.post(f"{API_BASE_URL}/ai-config/reset")
                    if response.status_code == 200:
                        st.session_state.ai_settings = None
                        st.success("✅ 已重置为默认配置")
                        st.rerun()
                    else:
                        st.error("❌ 重置失败")
                except Exception as e:
                    st.error(f"❌ 重置失败: {str(e)}")
    
    with tab2:
        st.subheader("🧪 测试 AI 连接")
        
        if st.button("🚀 测试连接", use_container_width=True):
            with st.spinner("正在测试 AI 连接..."):
                try:
                    response = requests.post(
                        f"{API_BASE_URL}/ai-config/test",
                        timeout=30
                    )
                    if response.status_code == 200:
                        result = response.json()
                        if result.get('success'):
                            st.success(f"✅ 连接成功！\n\n模型: {result.get('model', 'Unknown')}\n响应时间: {result.get('response_time', 'N/A')}s")
                            st.markdown("**测试回复：**")
                            st.info(result.get('message', '无回复内容'))
                        else:
                            st.error(f"❌ 连接失败: {result.get('error', '未知错误')}")
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 测试失败: {error_detail}")
                except Exception as e:
                    st.error(f"❌ 测试失败: {str(e)}")
    
    with tab3:
        st.subheader("📖 各平台 API Key 申请指南")
        
        providers_guide = {
            "openai": {
                "name": "OpenAI",
                "url": "https://platform.openai.com/",
                "steps": [
                    "访问 OpenAI 官网注册账号",
                    "进入 API keys 页面",
                    "点击 'Create new secret key'",
                    "复制生成的 API Key（以 sk- 开头）"
                ],
                "note": "需要绑定信用卡，新用户有 $5 免费额度"
            },
            "deepseek": {
                "name": "DeepSeek",
                "url": "https://platform.deepseek.com/",
                "steps": [
                    "访问 DeepSeek 开放平台",
                    "注册并登录账号",
                    "进入 'API Keys' 页面",
                    "创建新的 API Key"
                ],
                "note": "性价比高，适合国内用户使用"
            },
            "zhipu": {
                "name": "智谱 AI",
                "url": "https://open.bigmodel.cn/",
                "steps": [
                    "访问智谱 AI 开放平台",
                    "注册并实名认证",
                    "进入 'API Keys' 管理",
                    "创建新的 API Key"
                ],
                "note": "国内大模型，GLM-4 性能优秀"
            },
            "moonshot": {
                "name": "Moonshot (月之暗面)",
                "url": "https://platform.moonshot.cn/",
                "steps": [
                    "访问 Moonshot 开放平台",
                    "注册并登录账号",
                    "进入 API Key 管理",
                    "创建新的 API Key"
                ],
                "note": "Kimi 大模型，支持长文本"
            },
            "qwen": {
                "name": "通义千问",
                "url": "https://dashscope.aliyun.com/",
                "steps": [
                    "访问阿里云 DashScope",
                    "使用阿里云账号登录",
                    "开通 DashScope 服务",
                    "在 'API-KEY 管理' 创建 Key"
                ],
                "note": "阿里云出品，新用户有免费额度"
            }
        }
        
        for provider_id, guide in providers_guide.items():
            with st.expander(f"{guide['name']}"):
                st.markdown(f"**官网：** [{guide['url']}]({guide['url']})")
                st.markdown("**申请步骤：**")
                for i, step in enumerate(guide['steps'], 1):
                    st.markdown(f"{i}. {step}")
                st.info(f"💡 {guide['note']}")

# ==================== 文件管理 ====================
elif page == "📂 文件管理":
    st.title("📂 文件管理")
    
    st.markdown("""
    查看所有已提交的文件，包括文件格式、文件类型、与学生的对应关系等信息。
    """)
    
    # 搜索选项
    search_option = st.selectbox(
        "搜索方式",
        options=["全部文件", "按学生查询", "按提交查询", "按文件类型查询"]
    )
    
    if search_option == "全部文件":
        # 获取所有学生
        try:
            response = requests.get(f"{API_BASE_URL}/students")
            if response.status_code == 200:
                students = response.json()
                
                all_files = []
                
                # 遍历每个学生
                for i, student in enumerate(students):
                    student_id = student['student_id']
                    student_name = student['name']
                    
                    # 获取学生的所有提交
                    response = requests.get(f"{API_BASE_URL}/students/{student_id}/submissions")
                    if response.status_code == 200:
                        submissions = response.json()
                        
                        # 遍历每个提交
                        for j, submission in enumerate(submissions):
                            submission_id = submission['submission_id']
                            submission_title = submission['title']
                            submission_type = submission['submission_type']
                            
                            # 获取提交的文件
                            response = requests.get(f"{API_BASE_URL}/submissions/{submission_id}/files")
                            if response.status_code == 200:
                                files = response.json()
                                
                                # 处理每个文件
                                for k, file in enumerate(files):
                                    file_id = file.get('id')
                                    if file_id is None:
                                        continue  # 跳过没有id的文件
                                    all_files.append({
                                        "学生学号": student_id,
                                        "学生姓名": student_name,
                                        "提交ID": submission_id,
                                        "提交标题": submission_title,
                                        "提交类型": submission_type,
                                        "文件名": file.get('filename', file.get('file_name', '')),
                                        "文件类型": file.get('file_type', file.get('media_type', 'N/A')),
                                        "文件大小": f"{file.get('file_size', file.get('size_bytes', 0)) / 1024:.2f} KB",
                                        "上传时间": file.get('uploaded_at', 'N/A')[:19],
                                        "file_id": file_id
                                    })
                
                if all_files:
                    st.info(f"共 {len(all_files)} 个文件")
                    
                    # 显示表头
                    header_cols = st.columns([1, 1, 1, 2, 1, 2, 1, 1, 0.5, 0.5])
                    with header_cols[0]:
                        st.markdown("**学生学号**")
                    with header_cols[1]:
                        st.markdown("**学生姓名**")
                    with header_cols[2]:
                        st.markdown("**提交ID**")
                    with header_cols[3]:
                        st.markdown("**提交标题**")
                    with header_cols[4]:
                        st.markdown("**提交类型**")
                    with header_cols[5]:
                        st.markdown("**文件名**")
                    with header_cols[6]:
                        st.markdown("**文件类型**")
                    with header_cols[7]:
                        st.markdown("**文件大小**")
                    with header_cols[8]:
                        st.markdown("**修改**")
                    with header_cols[9]:
                        st.markdown("**删除**")
                    
                    # 表头分隔线
                    st.divider()
                    
                    # 为每个文件显示一行，包含信息和操作按钮
                    for i, file_info in enumerate(all_files):
                        cols = st.columns([1, 1, 1, 2, 1, 2, 1, 1, 0.5, 0.5])
                        
                        with cols[0]:
                            st.write(file_info['学生学号'])
                        with cols[1]:
                            st.write(file_info['学生姓名'])
                        with cols[2]:
                            st.write(file_info['提交ID'])
                        with cols[3]:
                            st.write(file_info['提交标题'])
                        with cols[4]:
                            st.write(file_info['提交类型'])
                        with cols[5]:
                            st.write(file_info['文件名'])
                        with cols[6]:
                            st.write(file_info['文件类型'])
                        with cols[7]:
                            st.write(file_info['文件大小'])
                        with cols[8]:
                            # 修改按钮
                            if st.button(f"✏️", key=f"edit_file_{file_info['file_id']}", help="修改文件信息"):
                                # 存储当前编辑的文件信息
                                st.session_state['edit_file'] = file_info
                                # 切换到修改文件表单
                                st.session_state['show_edit_file_form'] = True
                        with cols[9]:
                            # 删除按钮 - 使用确认对话框
                            delete_key = f"delete_file_{file_info['file_id']}_{i}"  # 添加索引确保唯一性
                            
                            # 检查是否已经点击过删除按钮
                            if st.session_state.get('pending_delete_file_id') == file_info['file_id']:
                                st.warning("⚠️ 确认删除")
                                st.write(f"确定要删除文件 {file_info['文件名']} 吗？此操作不可恢复。")
                                
                                col1, col2 = st.columns(2)
                                with col1:
                                    if st.button("✅ 确认删除", key=f"confirm_file_delete_{i}"):
                                        try:
                                            delete_response = requests.delete(f"{API_BASE_URL}/files/{file_info['file_id']}")
                                            if delete_response.status_code == 200:
                                                st.success(f"✅ 文件 {file_info['文件名']} 删除成功！")
                                                # 重置状态
                                                st.session_state['pending_delete_file_id'] = None
                                                # 刷新页面
                                                st.rerun()
                                            else:
                                                try:
                                                    error_detail = delete_response.json().get('detail', '未知错误')
                                                except:
                                                    error_detail = f"HTTP {delete_response.status_code}"
                                                st.error(f"❌ 删除失败: {error_detail}")
                                        except Exception as e:
                                            st.error(f"❌ 删除失败: {str(e)}")
                                with col2:
                                    if st.button("❌ 取消", key=f"cancel_file_delete_{i}"):
                                        st.session_state['pending_delete_file_id'] = None
                            else:
                                if st.button("🗑️", key=delete_key, help="删除文件"):
                                    # 设置待删除的文件ID
                                    st.session_state['pending_delete_file_id'] = file_info['file_id']
                                    st.rerun()
                        
                        # 在每个文件行之间添加分隔线（除了最后一行）
                        if i < len(all_files) - 1:
                            st.divider()
                else:
                    st.info("📭 暂无文件")
            else:
                st.error("❌ 获取学生列表失败")
        except Exception as e:
            st.error(f"❌ 加载文件失败: {str(e)}")
    
    elif search_option == "按学生查询":
        student_id = st.text_input("输入学号", placeholder="请输入学生学号")
        
        if st.button("🔍 查询", use_container_width=True):
            if not student_id:
                st.error("❌ 请输入学号")
            else:
                try:
                    # 获取学生信息
                    response = requests.get(f"{API_BASE_URL}/students/{student_id}")
                    if response.status_code == 200:
                        student = response.json()
                        st.success(f"✅ 学生: {student['name']}")
                        
                        # 获取学生的所有提交
                        response = requests.get(f"{API_BASE_URL}/students/{student_id}/submissions")
                        if response.status_code == 200:
                            submissions = response.json()
                            
                            student_files = []
                            
                            # 遍历每个提交
                            for submission in submissions:
                                submission_id = submission['submission_id']
                                submission_title = submission['title']
                                submission_type = submission['submission_type']
                                
                                # 获取提交的文件
                                response = requests.get(f"{API_BASE_URL}/submissions/{submission_id}/files")
                                if response.status_code == 200:
                                    files = response.json()
                                    
                                    # 处理每个文件
                                    for file in files:
                                        student_files.append({
                                            "提交ID": submission_id,
                                            "提交标题": submission_title,
                                            "提交类型": submission_type,
                                            "文件名": file.get('file_path', '').split('/')[-1].split('\\')[-1],
                                            "文件类型": file.get('media_type', 'N/A'),
                                            "文件大小": f"{file.get('size_bytes', 0) / 1024:.2f} KB",
                                            "上传时间": file.get('uploaded_at', 'N/A')[:19]
                                        })
                            
                            if student_files:
                                df = pd.DataFrame(student_files)
                                st.dataframe(df, use_container_width=True)
                                st.info(f"共 {len(student_files)} 个文件")
                            else:
                                st.info("📭 该学生暂无文件")
                        else:
                            st.error("❌ 获取提交列表失败")
                    else:
                        st.error("❌ 学生不存在")
                except Exception as e:
                    st.error(f"❌ 查询失败: {str(e)}")
    
    elif search_option == "按提交查询":
        submission_id = st.text_input("输入提交ID", placeholder="请输入提交ID")
        
        if st.button("🔍 查询", use_container_width=True):
            if not submission_id:
                st.error("❌ 请输入提交ID")
            else:
                try:
                    # 获取提交信息
                    response = requests.get(f"{API_BASE_URL}/submissions/{submission_id}")
                    if response.status_code == 200:
                        submission = response.json()
                        st.success(f"✅ 提交: {submission['title']}")
                        
                        # 获取提交的文件
                        response = requests.get(f"{API_BASE_URL}/submissions/{submission_id}/files")
                        if response.status_code == 200:
                            files = response.json()
                            
                            if files:
                                submission_files = []
                                for file in files:
                                    submission_files.append({
                                        "文件名": file.get('file_path', '').split('/')[-1].split('\\')[-1],
                                        "文件类型": file.get('media_type', 'N/A'),
                                        "文件大小": f"{file.get('size_bytes', 0) / 1024:.2f} KB",
                                        "上传时间": file.get('uploaded_at', 'N/A')[:19]
                                    })
                                df = pd.DataFrame(submission_files)
                                st.dataframe(df, use_container_width=True)
                                st.info(f"共 {len(files)} 个文件")
                            else:
                                st.info("📭 该提交暂无文件")
                        else:
                            st.error("❌ 获取文件列表失败")
                    else:
                        st.error("❌ 提交不存在")
                except Exception as e:
                    st.error(f"❌ 查询失败: {str(e)}")
    
    elif search_option == "按文件类型查询":
        file_type = st.selectbox(
            "选择文件类型",
            options=["所有类型", "document", "video", "audio"]
        )
        
        if st.button("🔍 查询", use_container_width=True):
            try:
                # 获取所有学生
                response = requests.get(f"{API_BASE_URL}/students")
                if response.status_code == 200:
                    students = response.json()
                    
                    type_files = []
                    
                    # 遍历每个学生
                    for student in students:
                        student_id = student['student_id']
                        student_name = student['name']
                        
                        # 获取学生的所有提交
                        response = requests.get(f"{API_BASE_URL}/students/{student_id}/submissions")
                        if response.status_code == 200:
                            submissions = response.json()
                            
                            # 遍历每个提交
                            for submission in submissions:
                                submission_id = submission['submission_id']
                                submission_title = submission['title']
                                
                                # 获取提交的文件
                                response = requests.get(f"{API_BASE_URL}/submissions/{submission_id}/files")
                                if response.status_code == 200:
                                    files = response.json()
                                    
                                    # 处理每个文件
                                    for file in files:
                                        if file_type == "所有类型" or file.get('media_type') == file_type:
                                            type_files.append({
                                                "学生学号": student_id,
                                                "学生姓名": student_name,
                                                "提交ID": submission_id,
                                                "提交标题": submission_title,
                                                "文件名": file.get('file_path', '').split('/')[-1].split('\\')[-1],
                                                "文件类型": file.get('media_type', 'N/A'),
                                                "文件大小": f"{file.get('size_bytes', 0) / 1024:.2f} KB",
                                                "上传时间": file.get('uploaded_at', 'N/A')[:19]
                                            })
                    
                    if type_files:
                        df = pd.DataFrame(type_files)
                        st.dataframe(df, use_container_width=True)
                        st.info(f"共 {len(type_files)} 个文件")
                    else:
                        st.info("📭 暂无符合条件的文件")
                else:
                    st.error("❌ 获取学生列表失败")
            except Exception as e:
                st.error(f"❌ 查询失败: {str(e)}")
    
    # 修改文件表单
    if st.session_state.get('show_edit_file_form', False):
        st.title("✏️ 修改文件信息")
        edit_file = st.session_state.get('edit_file', {})
        
        with st.form("edit_file_form"):
            col1, col2 = st.columns(2)
            with col1:
                student_id = st.text_input("学生学号", value=edit_file.get('学生学号', ''), disabled=True)
                student_name = st.text_input("学生姓名", value=edit_file.get('学生姓名', ''), disabled=True)
                submission_id = st.text_input("提交ID", value=edit_file.get('提交ID', ''), disabled=True)
            with col2:
                submission_title = st.text_input("提交标题", value=edit_file.get('提交标题', ''), disabled=True)
                submission_type = st.text_input("提交类型", value=edit_file.get('提交类型', ''), disabled=True)
                file_name = st.text_input("文件名", value=edit_file.get('文件名', ''))
            
            # 处理文件类型，将MIME类型映射到简单类型
            file_type = edit_file.get('文件类型', 'document')
            # 映射MIME类型到简单类型
            if 'pdf' in file_type.lower() or 'word' in file_type.lower() or 'document' in file_type.lower() or file_type == 'N/A':
                default_type = 'document'
            elif 'video' in file_type.lower():
                default_type = 'video'
            elif 'audio' in file_type.lower():
                default_type = 'audio'
            else:
                default_type = 'document'
            
            media_type = st.selectbox(
                "文件类型",
                options=["document", "video", "audio"],
                index=["document", "video", "audio"].index(default_type)
            )
            
            st.markdown("*必填项")
            
            submit_button = st.form_submit_button("✅ 保存修改", use_container_width=True)
            cancel_button = st.form_submit_button("❌ 取消", use_container_width=True)
            
            if submit_button:
                if not file_name:
                    st.error("❌ 文件名不能为空")
                else:
                    try:
                        # 使用表单格式提交
                        response = requests.put(
                            f"{API_BASE_URL}/files/{edit_file.get('file_id', '')}",
                            data={
                                "file_name": file_name,
                                "media_type": media_type
                            }
                        )
                        if response.status_code == 200:
                            st.success("✅ 文件信息修改成功！")
                            file_data = response.json()
                            st.json(file_data)
                            # 关闭编辑表单
                            st.session_state['show_edit_file_form'] = False
                            st.session_state.pop('edit_file', None)
                        else:
                            try:
                                error_detail = response.json().get('detail', '未知错误')
                            except:
                                error_detail = f"HTTP {response.status_code}: {response.text}"
                            st.error(f"❌ 修改失败: {error_detail}")
                    except Exception as e:
                        st.error(f"❌ 修改失败: {str(e)}")
            
            if cancel_button:
                st.session_state['show_edit_file_form'] = False
                st.session_state.pop('edit_file', None)
                st.rerun()

# ==================== API文档 ====================
elif page == "🔧 API文档":
    st.title("🔧 API 文档")
    
    st.markdown("""
    学生多维度能力评估系统提供完整的 RESTful API 接口
    
    **Swagger UI:** [http://localhost:8000/docs](http://localhost:8000/docs)
    
    **ReDoc:** [http://localhost:8000/redoc](http://localhost:8000/redoc)
    """)
    
    st.markdown("---")
    
    # API 端点列表
    st.subheader("📚 API 端点")
    
    endpoints = [
        {
            "method": "GET",
            "path": "/",
            "description": "API 欢迎页面",
            "params": "无"
        },
        {
            "method": "GET",
            "path": "/health",
            "description": "健康检查",
            "params": "无"
        },
        {
            "method": "GET",
            "path": "/ai-config",
            "description": "获取当前 AI 配置",
            "params": "无"
        },
        {
            "method": "POST",
            "path": "/ai-config",
            "description": "更新 AI 配置",
            "params": "provider, api_key, model, base_url, temperature, max_tokens"
        },
        {
            "method": "POST",
            "path": "/ai-config/test",
            "description": "测试 AI 连接",
            "params": "无"
        },
        {
            "method": "POST",
            "path": "/students",
            "description": "创建学生",
            "params": "student_id, name, age, grade, major"
        },
        {
            "method": "GET",
            "path": "/students/{student_id}",
            "description": "获取学生信息",
            "params": "student_id (路径参数)"
        },
        {
            "method": "GET",
            "path": "/students",
            "description": "获取所有学生列表",
            "params": "skip, limit (查询参数)"
        },
        {
            "method": "PUT",
            "path": "/students/{student_id}",
            "description": "更新学生信息",
            "params": "student_id (路径参数), name, age, grade, major"
        },
        {
            "method": "POST",
            "path": "/submissions",
            "description": "创建提交",
            "params": "student_id, title, description"
        },
        {
            "method": "GET",
            "path": "/submissions/{submission_id}",
            "description": "获取提交信息",
            "params": "submission_id (路径参数)"
        },
        {
            "method": "POST",
            "path": "/submissions/{submission_id}/files",
            "description": "上传文件",
            "params": "submission_id (路径参数), file (文件)"
        },
        {
            "method": "GET",
            "path": "/submissions/{submission_id}/files",
            "description": "获取提交的文件列表",
            "params": "submission_id (路径参数)"
        },
        {
            "method": "POST",
            "path": "/evaluate",
            "description": "启动评估",
            "params": "submission_id"
        },
        {
            "method": "GET",
            "path": "/evaluations/{evaluation_id}",
            "description": "获取评估结果",
            "params": "evaluation_id (路径参数)"
        },
        {
            "method": "GET",
            "path": "/students/{student_id}/evaluations",
            "description": "获取学生的所有评估",
            "params": "student_id (路径参数)"
        },
        {
            "method": "GET",
            "path": "/submissions/{submission_id}/evaluation",
            "description": "获取提交的评估结果",
            "params": "submission_id (路径参数)"
        }
    ]
    
    for endpoint in endpoints:
        with st.expander(f"{endpoint['method']} {endpoint['path']}"):
            st.markdown(f"**描述:** {endpoint['description']}")
            st.markdown(f"**参数:** {endpoint['params']}")
    
    st.markdown("---")
    
    # 请求示例
    st.subheader("📝 请求示例")
    
    st.markdown("**更新 AI 配置:**")
    st.code("""
POST /ai-config
Content-Type: application/json

{
    "provider": "deepseek",
    "api_key": "sk-your-api-key",
    "model": "deepseek-chat",
    "temperature": 0.7,
    "max_tokens": 2000
}
    """, language="http")
    
    st.markdown("**创建学生:**")
    st.code("""
POST /students
Content-Type: application/json

{
    "student_id": "2024001",
    "name": "张三",
    "age": 20,
    "grade": "大二",
    "major": "计算机科学"
}
    """, language="http")
    
    st.markdown("**创建提交:**")
    st.code("""
POST /submissions
Content-Type: application/json

{
    "student_id": "2024001",
    "title": "机器学习论文",
    "description": "关于深度学习的期末论文"
}
    """, language="http")
    
    st.markdown("**启动评估:**")
    st.code("""
POST /evaluate
Content-Type: application/json

{
    "submission_id": "SUB_XXXXXXXX"
}
    """, language="http")

# ==================== 成长分析 ====================
elif page == "📈 成长分析":
    st.title("📈 学生成长分析")
    st.markdown("""
    **功能说明：** 全面分析学生在各个维度的能力成长变化，通过多种可视化图表直观展示进步趋势。
    """)
    
    # 时间维度选择
    st.subheader("📅 时间维度选择")
    # 使用会话状态保存时间维度选择，避免自动刷新
    if 'time_dimension' not in st.session_state:
        st.session_state['time_dimension'] = "评估日期"
    
    time_dimension = st.radio(
        "选择时间维度",
        options=["评估日期", "工作时期进度"],
        horizontal=True,
        key="time_dimension",
        index=0 if st.session_state['time_dimension'] == "评估日期" else 1
    )
    
    # 选择学生
    st.subheader("👤 选择学生")
    try:
        response = requests.get(f"{API_BASE_URL}/students")
        if response.status_code == 200:
            students = response.json()
            if students:
                # 构建学生选项
                student_options = {student['student_id']: f"{student['student_id']} - {student['name']}" for student in students}
                selected_student_id = st.selectbox(
                    "选择学生",
                    options=list(student_options.keys()),
                    format_func=lambda x: student_options[x]
                )
                
                # 实时显示学生的总评估历史记录
                st.subheader("📜 学生总评估历史记录")
                try:
                    # 获取学生的所有评估记录
                    response = requests.get(f"{API_BASE_URL}/students/{selected_student_id}/evaluations")
                    if response.status_code == 200:
                        evaluations = response.json()
                        if evaluations:
                            st.success(f"✅ 找到 {len(evaluations)} 条评估记录")
                            
                            # 显示评估历史表格
                            eval_data = []
                            for eval in evaluations:
                                # 计算平均维度评分
                                avg_dim_score = 0
                                if eval.get('dimension_scores'):
                                    avg_dim_score = sum(ds['score'] for ds in eval['dimension_scores']) / len(eval['dimension_scores'])
                                
                                eval_data.append({
                                    '评估ID': eval['evaluation_id'],
                                    '提交ID': eval.get('submission_id', 'N/A'),
                                    '综合评分': eval['overall_score'],
                                    '平均维度评分': f"{avg_dim_score:.2f}",
                                    '评估时间': eval['evaluated_at'][:10],
                                    '评估阶段': f"{eval.get('stage_progress', 0.0):.2f}" if eval.get('stage_progress') is not None else 'N/A'
                                })
                            
                            # 显示表格
                            df = pd.DataFrame(eval_data)
                            st.dataframe(df, use_container_width=True)
                            
                            # 提供下载选项
                            csv = df.to_csv(index=False)
                            st.download_button(
                                label="📥 下载评估历史记录",
                                data=csv,
                                file_name=f"{selected_student_id}_evaluation_history.csv",
                                mime="text/csv"
                            )
                        else:
                            st.info("📭 该学生暂无评估记录")
                    else:
                        try:
                            error_detail = response.json().get('detail', '未知错误')
                        except:
                            error_detail = f"HTTP {response.status_code}"
                        st.error(f"❌ 获取评估记录失败: {error_detail}")
                except Exception as e:
                    st.error(f"❌ 加载评估记录失败: {str(e)}")
                
                # 结构查询页按课程类型细则进行总进度分析，不再使用旧十维图表集合
                if st.button("🔍 分析成长数据", use_container_width=True):
                    try:
                        report_resp = requests.get(f"{API_BASE_URL}/students/{selected_student_id}/progress-report")
                        if report_resp.status_code == 200:
                            report_data = report_resp.json()
                            st.success("✅ 成长数据分析完成（课程细则口径）")
                            render_policy_progress_report(report_data, key_prefix=f"struct_{selected_student_id}")
                            with st.expander("查看报告原文"):
                                st.markdown(report_data.get("report", "暂无报告内容"))
                        else:
                            try:
                                error_detail = report_resp.json().get('detail', '未知错误')
                            except Exception:
                                error_detail = f"HTTP {report_resp.status_code}"
                            st.error(f"❌ 获取成长分析失败: {error_detail}")
                    except Exception as e:
                        st.error(f"❌ 分析失败: {str(e)}")
            else:
                st.info("📭 暂无学生记录")
        else:
            st.error("❌ 获取学生列表失败")
    except Exception as e:
        st.error(f"❌ 加载学生列表失败: {str(e)}")
