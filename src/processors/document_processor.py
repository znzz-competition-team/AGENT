from typing import Dict, Any, List, Optional
from pathlib import Path
import logging
import pypdf
<<<<<<< HEAD
import re
=======
>>>>>>> 13733ce0a70eef683f89b9c58cf4bdf335da8e17
from docx import Document
import pandas as pd
from .base_processor import BaseProcessor

logger = logging.getLogger(__name__)

<<<<<<< HEAD
def clean_pdf_text(text):
    """清理PDF提取的文本，去除字符间的多余空格"""
    if not text:
        return ""
    
    clean_text = ''.join(c if c.isprintable() or c in '\n\t\r' else ' ' for c in text)
    
    lines = clean_text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        if not line.strip():
            continue
            
        if re.search(r'[\u4e00-\u9fff]', line):
            chars = list(line)
            result = []
            i = 0
            while i < len(chars):
                char = chars[i]
                if '\u4e00' <= char <= '\u9fff':
                    result.append(char)
                    i += 1
                    while i < len(chars) and chars[i] == ' ':
                        i += 1
                    while i < len(chars) and '\u4e00' <= chars[i] <= '\u9fff':
                        result.append(chars[i])
                        i += 1
                        while i < len(chars) and chars[i] == ' ':
                            i += 1
                else:
                    result.append(char)
                    i += 1
            cleaned_lines.append(''.join(result))
        else:
            cleaned_lines.append(re.sub(r' +', ' ', line))
    
    return '\n'.join(cleaned_lines)

=======
>>>>>>> 13733ce0a70eef683f89b9c58cf4bdf335da8e17
class DocumentProcessor(BaseProcessor):
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
    def process(self, file_path: str) -> Dict[str, Any]:
        try:
            self.validate_file(file_path)
            file_info = self.get_file_info(file_path)
            
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                content = self.process_pdf(file_path)
            elif file_extension in {'.docx', '.doc'}:
                content = self.process_docx(file_path)
            elif file_extension in {'.xlsx', '.xls'}:
                content = self.process_excel(file_path)
            elif file_extension == '.txt':
                content = self.process_text(file_path)
            else:
                raise ValueError(f"Unsupported document format: {file_extension}")
            
            result = {
                "file_info": file_info,
                "content": content,
                "status": "success"
            }
            
            self.logger.info(f"Document processing completed: {file_path}")
            return result
            
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            return {
                "file_path": file_path,
                "status": "error",
                "error_message": str(e)
            }
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        try:
            pdf_reader = pypdf.PdfReader(file_path)
            num_pages = len(pdf_reader.pages)
            
            text_content = []
            metadata = {}
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
<<<<<<< HEAD
                cleaned_text = clean_pdf_text(page_text) if page_text else ""
                text_content.append({
                    "page_number": page_num + 1,
                    "text": cleaned_text,
                    "char_count": len(cleaned_text)
=======
                text_content.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "char_count": len(page_text)
>>>>>>> 13733ce0a70eef683f89b9c58cf4bdf335da8e17
                })
            
            if pdf_reader.metadata:
                metadata = {
                    "title": pdf_reader.metadata.get('/Title', ''),
                    "author": pdf_reader.metadata.get('/Author', ''),
                    "creator": pdf_reader.metadata.get('/Creator', ''),
                    "producer": pdf_reader.metadata.get('/Producer', ''),
                    "creation_date": str(pdf_reader.metadata.get('/CreationDate', '')),
                    "modification_date": str(pdf_reader.metadata.get('/ModDate', ''))
                }
            
            full_text = '\n'.join([page["text"] for page in text_content])
            
            return {
                "type": "pdf",
                "num_pages": num_pages,
                "pages": text_content,
                "full_text": full_text,
                "metadata": metadata,
                "total_chars": len(full_text),
                "total_words": len(full_text.split())
            }
            
        except Exception as e:
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        try:
            doc = Document(file_path)
            
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                paragraphs.append({
                    "paragraph_index": i,
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal"
                })
            
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    "table_index": table_idx,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "data": table_data
                })
            
            full_text = '\n'.join([para["text"] for para in paragraphs])
            
            return {
                "type": "docx",
                "paragraphs": paragraphs,
                "tables": tables,
                "full_text": full_text,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "total_chars": len(full_text),
                "total_words": len(full_text.split())
            }
            
        except Exception as e:
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    def process_excel(self, file_path: str) -> Dict[str, Any]:
        try:
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            sheets_data = {}
            total_rows = 0
            total_columns = 0
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                sheets_data[sheet_name] = {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": df.columns.tolist(),
                    "data": df.to_dict('records'),
                    "summary": {
                        "numeric_columns": df.select_dtypes(include=['number']).columns.tolist(),
                        "text_columns": df.select_dtypes(include=['object']).columns.tolist()
                    }
                }
                
                total_rows += len(df)
                total_columns = max(total_columns, len(df.columns))
            
            return {
                "type": "excel",
                "sheet_names": sheet_names,
                "sheets": sheets_data,
                "total_sheets": len(sheet_names),
                "total_rows": total_rows,
                "total_columns": total_columns
            }
            
        except Exception as e:
            raise ValueError(f"Failed to process Excel: {str(e)}")
    
    def process_text(self, file_path: str) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            lines = text.split('\n')
            paragraphs = [line.strip() for line in lines if line.strip()]
            
            return {
                "type": "text",
                "full_text": text,
                "lines": lines,
                "paragraphs": paragraphs,
                "total_lines": len(lines),
                "total_paragraphs": len(paragraphs),
                "total_chars": len(text),
                "total_words": len(text.split())
            }
            
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    text = f.read()
                
                lines = text.split('\n')
                paragraphs = [line.strip() for line in lines if line.strip()]
                
                return {
                    "type": "text",
                    "full_text": text,
                    "lines": lines,
                    "paragraphs": paragraphs,
                    "total_lines": len(lines),
                    "total_paragraphs": len(paragraphs),
                    "total_chars": len(text),
                    "total_words": len(text.split()),
                    "encoding": "gbk"
                }
            except Exception as e:
                raise ValueError(f"Failed to read text file: {str(e)}")
    
    def extract_keywords(self, text: str, num_keywords: int = 10) -> List[str]:
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            import jieba
            
            words = jieba.cut(text)
            filtered_words = [word for word in words if len(word) > 1]
            
            if len(filtered_words) < num_keywords:
                return list(set(filtered_words))
            
            vectorizer = TfidfVectorizer(max_features=num_keywords)
            tfidf_matrix = vectorizer.fit_transform([' '.join(filtered_words)])
            feature_names = vectorizer.get_feature_names_out()
            
            return feature_names.tolist()
            
        except Exception as e:
            self.logger.warning(f"Could not extract keywords: {e}")
            return []
    
    def validate_file(self, file_path: str) -> bool:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        valid_extensions = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt'}
        if path.suffix.lower() not in valid_extensions:
            raise ValueError(f"Invalid document format: {path.suffix}")
        
        return True