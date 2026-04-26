"""
Word文档提取器模块 - 多库组合策略提取Word文档内容

支持：
1. python-docx - 主要方法，支持段落、表格、样式
2. docx2txt - 备用方法，简单文本提取
3. mammoth - 转换为HTML后提取（可选）
"""

import re
import logging
from typing import Tuple, List, Dict, Optional

logger = logging.getLogger(__name__)


class WordExtractor:
    """增强版Word文档提取器，使用多库组合策略"""
    
    def __init__(self, min_text_length: int = 50):
        self.min_text_length = min_text_length
    
    def extract(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        提取Word文档内容，返回(文本内容, 表格列表)
        
        使用多库组合策略：
        1. 优先使用python-docx提取（支持段落、表格、样式）
        2. 如果失败或内容太短，尝试docx2txt
        3. 清理和格式化文本
        """
        text = ""
        tables = []
        
        text, tables = self._extract_with_python_docx(file_path)
        
        if len(text.strip()) < self.min_text_length:
            backup_text = self._extract_with_docx2txt(file_path)
            if len(backup_text.strip()) > len(text.strip()):
                text = backup_text
                tables = []
        
        text = self._clean_text(text)
        
        return text, tables
    
    def _extract_with_python_docx(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        使用python-docx提取内容
        
        优点：
        - 支持段落样式识别
        - 支持表格提取
        - 支持页眉页脚
        """
        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document(file_path)
            text_parts = []
            tables = []
            
            for section in doc.sections:
                header = section.header
                if header:
                    header_text = "\n".join([p.text for p in header.paragraphs if p.text.strip()])
                    if header_text:
                        text_parts.append(f"[页眉] {header_text}")
            
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element is element:
                            para = p
                            break
                    
                    if para:
                        text = para.text.strip()
                        if text:
                            style_name = para.style.name if para.style else ""
                            
                            if "Heading" in style_name or "标题" in style_name:
                                level = self._get_heading_level(style_name)
                                text_parts.append(f"\n{'#' * level} {text}\n")
                            elif para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                                text_parts.append(f"\n**{text}**\n")
                            else:
                                text_parts.append(text)
            
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    text_parts.append(" | ".join(row_data))
                
                if table_data:
                    tables.append({
                        "index": i,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                        "data": table_data
                    })
            
            for section in doc.sections:
                footer = section.footer
                if footer:
                    footer_text = "\n".join([p.text for p in footer.paragraphs if p.text.strip()])
                    if footer_text:
                        text_parts.append(f"[页脚] {footer_text}")
            
            result = "\n".join(text_parts)
            logger.info(f"python-docx提取成功，共{len(result)}字符，{len(tables)}个表格")
            return result, tables
            
        except ImportError:
            logger.warning("python-docx未安装，跳过此方法")
            return "", []
        except Exception as e:
            logger.error(f"python-docx提取失败: {str(e)}")
            return "", []
    
    def _extract_with_docx2txt(self, file_path: str) -> str:
        """
        使用docx2txt提取文本（备用方法）
        
        优点：
        - 简单可靠
        - 可以提取嵌入的图片
        """
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            logger.info(f"docx2txt提取成功，共{len(text)}字符")
            return text
        except ImportError:
            logger.warning("docx2txt未安装，跳过此方法")
            return ""
        except Exception as e:
            logger.error(f"docx2txt提取失败: {str(e)}")
            return ""
    
    def _extract_with_mammoth(self, file_path: str) -> str:
        """
        使用mammoth提取文本（可选备用方法）
        
        优点：
        - 转换为HTML，保留格式
        - 支持复杂样式
        """
        try:
            import mammoth
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text = result.value
                logger.info(f"mammoth提取成功，共{len(text)}字符")
                return text
        except ImportError:
            logger.warning("mammoth未安装，跳过此方法")
            return ""
        except Exception as e:
            logger.error(f"mammoth提取失败: {str(e)}")
            return ""
    
    def _get_heading_level(self, style_name: str) -> int:
        """获取标题级别"""
        if not style_name:
            return 1
        
        match = re.search(r'(\d+)', style_name)
        if match:
            return int(match.group(1))
        
        if "标题 1" in style_name or "Heading 1" in style_name:
            return 1
        elif "标题 2" in style_name or "Heading 2" in style_name:
            return 2
        elif "标题 3" in style_name or "Heading 3" in style_name:
            return 3
        
        return 1
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本内容
        
        处理：
        1. 多余空格
        2. 多余空行
        3. 特殊字符
        4. 保留段落结构
        """
        if not text:
            return ""
        
        text = re.sub(r'[ \t]+', '', text)
        
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            if re.match(r'^[IVXivx]+$', line) and len(line) <= 5:
                continue
            if re.match(r'^\d+$', line) and len(line) <= 3:
                continue
            
            if line:
                cleaned_lines.append(line)
        
        result = '\n'.join(cleaned_lines)
        
        return result
    
    def extract_with_structure(self, file_path: str) -> Dict:
        """
        提取Word文档内容（带结构信息）
        
        返回：
        - text: 完整文本
        - tables: 表格列表
        - headings: 标题列表（带层级）
        - structure: 文档结构
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            text_parts = []
            tables = []
            headings = []
            structure = []
            
            current_section = {"title": "文档开头", "level": 0, "content": []}
            
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element is element:
                            para = p
                            break
                    
                    if para:
                        text = para.text.strip()
                        if text:
                            style_name = para.style.name if para.style else ""
                            
                            if "Heading" in style_name or "标题" in style_name:
                                level = self._get_heading_level(style_name)
                                headings.append({
                                    "text": text,
                                    "level": level
                                })
                                
                                if current_section["content"]:
                                    structure.append(current_section)
                                
                                current_section = {
                                    "title": text,
                                    "level": level,
                                    "content": []
                                }
                                
                                text_parts.append(f"\n{'#' * level} {text}\n")
                            else:
                                current_section["content"].append(text)
                                text_parts.append(text)
            
            if current_section["content"]:
                structure.append(current_section)
            
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        "index": i,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                        "data": table_data
                    })
            
            result = {
                "text": self._clean_text("\n".join(text_parts)),
                "tables": tables,
                "headings": headings,
                "structure": structure
            }
            
            logger.info(f"结构化提取成功: {len(result['text'])}字符, {len(tables)}表格, {len(headings)}标题")
            return result
            
        except Exception as e:
            logger.error(f"结构化提取失败: {str(e)}")
            text, tables = self.extract(file_path)
            return {
                "text": text,
                "tables": tables,
                "headings": [],
                "structure": []
            }


def extract_word_content(file_path: str) -> str:
    """
    便捷函数：提取Word文档内容
    
    Args:
        file_path: Word文档路径
        
    Returns:
        提取的文本内容
    """
    extractor = WordExtractor()
    text, _ = extractor.extract(file_path)
    return text


def extract_word_with_tables(file_path: str) -> Tuple[str, List[Dict]]:
    """
    便捷函数：提取Word文档内容和表格
    
    Args:
        file_path: Word文档路径
        
    Returns:
        (文本内容, 表格列表)
    """
    extractor = WordExtractor()
    return extractor.extract(file_path)
