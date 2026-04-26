from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form, Body
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
import os
import shutil
from datetime import datetime
import sys
import time
import logging
import re
from src.course_classifier import classify_course_type

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('api_debug.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

# 添加项目根目录到 Python 路径
current_file = os.path.abspath(__file__)
api_dir = os.path.dirname(current_file)
src_dir = os.path.dirname(api_dir)
project_root = os.path.dirname(src_dir)

# 确保 src 目录和项目根目录都在 Python 路径中
for path in [src_dir, project_root]:
    if path not in sys.path:
        sys.path.insert(0, path)

from src.config import settings, AI_PROVIDERS, get_ai_config
from src.database import get_db, DatabaseService, init_db
from src.database.models import Student, Submission, MediaFile, EvaluationResult, ProgressReport
from src.models.schemas import (
    StudentCreate, StudentUpdate, StudentResponse,
    SubmissionCreate, SubmissionResponse, SubmissionStatus, SubmissionType, SubmissionPurpose,
    MediaFileResponse,
    EvaluationResultResponse, DimensionScoreResponse,
    EvaluationRequest, EvaluationResponse, ProgressReportResponse,
    EvaluationDimension
)
from src.models.schemas import EvaluationResult as SchemaEvaluationResult
# 移除对CrewAI和MediaProcessor的依赖，避免CV2依赖
# from src.agents.crew_manager import StudentEvaluationCrew
# from src.utils.media_processor import MediaProcessor
from pydantic import BaseModel

from fastapi.middleware.cors import CORSMiddleware
from fastapi.requests import Request
from fastapi.responses import JSONResponse
import json
from typing import Union

# 增加文件大小限制
from starlette.middleware.base import BaseHTTPMiddleware

# 辅助函数：处理evidence字段，确保返回List[str]类型
def process_evidence(evidence: Union[str, list, None]) -> list:
    """处理evidence字段，确保返回List[str]类型"""
    if isinstance(evidence, str):
        # 如果是字符串，分割成列表
        return [item.strip() for item in evidence.split(",") if item.strip()]
    elif isinstance(evidence, list):
        # 如果已经是列表，直接返回
        return evidence
    else:
        # 如果是None或其他类型，返回空列表
        return []

def process_string_list(value: Union[str, list, None]) -> list:
    """处理字符串列表字段，确保返回List[str]类型"""
    if isinstance(value, str):
        # 如果是字符串，分割成列表
        return [item.strip() for item in value.split(",") if item.strip()]
    elif isinstance(value, list):
        # 如果已经是列表，直接返回
        return value
    else:
        # 如果是None或其他类型，返回空列表
        return []

# 文档内容提取函数
def extract_document_content(file_path: str) -> str:
    """提取文档内容"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == ".pdf":
            return extract_pdf_content(file_path)
        elif file_ext == ".docx":
            return extract_docx_content(file_path)
        elif file_ext == ".doc":
            return extract_doc_content(file_path)
        elif file_ext == ".txt":
            return extract_txt_content(file_path)
        else:
            return ""
    except Exception as e:
        logger.error(f"提取文件内容失败: {str(e)}")
        return ""

def extract_pdf_content(file_path: str) -> str:
    """
    提取PDF文件内容（增强版）
    
    使用增强的多策略方法：
    1. 优先使用增强版PDF提取器（保留结构、识别章节）
    2. 备用原PDF提取器
    3. 最后使用PyMuPDF/pdfplumber/PyPDF2
    """
    try:
        from src.utils.pdf_extractor_enhanced import extract_pdf_enhanced
        text = extract_pdf_enhanced(file_path, enable_ocr=False)
        if text and len(text.strip()) > 100:
            logger.info(f"使用增强版PDF提取器成功，共{len(text)}字符")
            return text
    except ImportError:
        logger.warning("增强版PDF提取器模块未找到，尝试原提取器")
    except Exception as e:
        logger.warning(f"增强版PDF提取器失败: {str(e)}，尝试原提取器")
    
    try:
        from src.utils.pdf_extractor import extract_pdf_content as extract_with_new
        text = extract_with_new(file_path)
        if text and len(text.strip()) > 50:
            logger.info(f"使用原PDF提取器成功，共{len(text)}字符")
            return text
    except ImportError:
        logger.warning("原PDF提取器模块未找到，使用备用方法")
    except Exception as e:
        logger.warning(f"原PDF提取器失败: {str(e)}，使用备用方法")
    
    try:
        import fitz
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        text = "\n".join(text_parts)
        if text.strip():
            logger.info(f"PyMuPDF提取成功，共{len(text)}字符")
            return text
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyMuPDF提取失败: {str(e)}")
    
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts)
        if text.strip():
            logger.info(f"pdfplumber提取成功，共{len(text)}字符")
            return text
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfplumber提取失败: {str(e)}")
    
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text()
        logger.info(f"PyPDF2提取成功，共{len(text)}字符")
        return text
    except Exception as e:
        logger.error(f"提取PDF内容失败: {str(e)}")
        return ""

def extract_docx_content(file_path: str) -> str:
    """
    提取DOCX文件内容（改进版）
    
    使用多库组合策略：
    1. 优先使用新的WordExtractor
    2. 备用python-docx
    3. 最后使用docx2txt
    """
    try:
        from src.utils.word_extractor import extract_word_content as extract_with_new
        text = extract_with_new(file_path)
        if text and len(text.strip()) > 50:
            logger.info(f"使用新Word提取器成功，共{len(text)}字符")
            return text
    except ImportError:
        logger.warning("新Word提取器模块未找到，使用备用方法")
    except Exception as e:
        logger.warning(f"新Word提取器失败: {str(e)}，使用备用方法")
    
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        text = "\n".join(text_parts)
        logger.info(f"python-docx提取成功，共{len(text)}字符")
        return text
    except Exception as e:
        logger.warning(f"python-docx提取失败: {str(e)}")
    
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        logger.info(f"docx2txt提取成功，共{len(text)}字符")
        return text
    except Exception as e:
        logger.error(f"docx2txt提取失败: {str(e)}")
        return ""

def extract_doc_content(file_path: str) -> str:
    """
    提取DOC文件内容（改进版）
    
    使用多库组合策略：
    1. 优先尝试使用新的WordExtractor（支持.docx）
    2. 使用docx2txt
    3. 使用antiword（Linux/Mac）
    """
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        if text and len(text.strip()) > 50:
            logger.info(f"docx2txt提取DOC成功，共{len(text)}字符")
            return text
    except Exception as e:
        logger.warning(f"docx2txt提取DOC失败: {str(e)}")
    
    try:
        import subprocess
        result = subprocess.run(['antiword', file_path], capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            text = result.stdout
            logger.info(f"antiword提取成功，共{len(text)}字符")
            return text
    except FileNotFoundError:
        logger.warning("antiword未安装")
    except Exception as e:
        logger.warning(f"antiword提取失败: {str(e)}")
    
    logger.error("所有DOC提取方法都失败")
    return ""

def extract_txt_content(file_path: str) -> str:
    """提取TXT文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        logger.error(f"提取TXT内容失败: {str(e)}")
        return ""

def extract_abstract(content: str) -> str:
    """
    从论文内容中提取中文摘要（改进版）
    
    改进点：
    1. 精确定位中文摘要，排除标题、关键词、英文摘要
    2. 在关键词、ABSTRACT等标记处停止
    3. 处理各种格式变体
    """
    if not content:
        return ""
    
    def clean_abstract_text(text: str) -> str:
        """清理摘要文本"""
        text = re.sub(r'[ \t]+', '', text)
        text = re.sub(r'\n+', '', text)
        text = text.strip()
        return text
    
    patterns = [
        r'摘\s*要\s*[：:]*\s*\n?(.*?)(?=\s*(?:关键词|Key\s*words|Keyword|ABSTRACT|Abstract|目录|引言|绪论|第一章|第1章))',
        r'(?:^|\n)\s*摘\s*要\s*[：:]*\s*\n?(.*?)(?=\n\s*(?:关键词|ABSTRACT|Abstract))',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            abstract = match.group(1).strip()
            abstract = clean_abstract_text(abstract)
            
            if len(abstract) > 50:
                return abstract
    
    lines = content.split('\n')
    abstract_start = -1
    abstract_end = -1
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        if '摘要' in line_stripped and 'ABSTRACT' not in line_stripped.upper():
            if '关键词' not in line_stripped:
                abstract_start = i + 1
                break
    
    if abstract_start == -1:
        return ""
    
    for j in range(abstract_start, min(abstract_start + 60, len(lines))):
        line_stripped = lines[j].strip()
        
        if not line_stripped:
            continue
        
        if re.search(r'关键词|Key\s*words|Keyword', line_stripped, re.IGNORECASE):
            abstract_end = j
            break
        
        if 'ABSTRACT' in line_stripped.upper() and '摘要' not in line_stripped:
            abstract_end = j
            break
        
        if re.match(r'^第[一二三四五六七八九十\d]+\s*章', line_stripped):
            abstract_end = j
            break
        
        if '目录' in line_stripped or '引言' in line_stripped or '绪论' in line_stripped:
            abstract_end = j
            break
        
        if re.match(r'^[IVXivx]+$', line_stripped) and len(line_stripped) <= 3:
            if j > abstract_start + 5:
                abstract_end = j
                break
    
    if abstract_end == -1:
        abstract_end = min(abstract_start + 50, len(lines))
    
    abstract_lines = []
    for k in range(abstract_start, abstract_end):
        line = lines[k].strip()
        if not line:
            continue
        
        if re.match(r'^[IVXivx]+$', line) and len(line) <= 3:
            continue
        if re.match(r'^\d+$', line) and len(line) <= 3:
            continue
        
        abstract_lines.append(line)
    
    if abstract_lines:
        abstract = ''.join(abstract_lines)
        abstract = clean_abstract_text(abstract)
        
        if len(abstract) > 50:
            return abstract
    
    return ""

def clean_text_content(text: str) -> str:
    """
    清理文本内容，移除多余空格
    
    与摘要提取使用相同的清理逻辑
    """
    if not text:
        return ""
    
    text = re.sub(r'[ \t]+', '', text)
    text = re.sub(r'\n+', '\n', text)
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if line:
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def extract_chapter(content: str, chapter_name: str) -> str:
    """
    从论文内容中提取指定章节
    
    Args:
        content: 论文全文内容
        chapter_name: 章节名称，如 "引言"、"绪论"、"结论" 等
        
    Returns:
        提取的章节内容
    """
    if not content or not chapter_name:
        return ""
    
    chapter_patterns = {
        "摘要": [r'摘\s*要\s*[：:]*\s*\n?(.*?)(?=\s*(?:关键词|Key\s*words|Keyword|ABSTRACT|Abstract))'],
        "引言": [r'(?:第[一二三四五六七八九十\d]+\s*章\s*)?引言\s*\n?(.*?)(?=\n\s*(?:第[一二三四五六七八九十\d]+\s*章|参考文献|致谢|结论|总结))'],
        "绪论": [r'(?:第[一二三四五六七八九十\d]+\s*章\s*)?绪论\s*\n?(.*?)(?=\n\s*(?:第[一二三四五六七八九十\d]+\s*章|参考文献|致谢|结论|总结))'],
        "结论": [r'(?:第[一二三四五六七八九十\d]+\s*章\s*)?结论\s*\n?(.*?)(?=\n\s*(?:参考文献|致谢|附录))'],
        "总结": [r'(?:第[一二三四五六七八九十\d]+\s*章\s*)?总结\s*\n?(.*?)(?=\n\s*(?:参考文献|致谢|附录))'],
        "参考文献": [r'参考\s*文献\s*\n?(.*?)(?=\n\s*(?:致谢|附录|$))'],
        "致谢": [r'致\s*谢\s*\n?(.*?)(?=\n\s*(?:附录|作者简介|$))'],
    }
    
    patterns = chapter_patterns.get(chapter_name, [re.escape(chapter_name) + r'\s*\n?(.*?)(?=\n\s*(?:第[一二三四五六七八九十\d]+\s*章|参考文献|致谢))'])
    
    for pattern in patterns:
        match = re.search(pattern, content, re.DOTALL | re.IGNORECASE)
        if match:
            chapter_content = match.group(1).strip()
            chapter_content = clean_text_content(chapter_content)
            if len(chapter_content) > 50:
                return chapter_content[:3000]
    
    lines = content.split('\n')
    chapter_start = -1
    chapter_end = -1
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if chapter_name in line_stripped:
            chapter_start = i + 1
            break
    
    if chapter_start == -1:
        return ""
    
    for j in range(chapter_start, min(chapter_start + 100, len(lines))):
        line_stripped = lines[j].strip()
        
        if re.match(r'^第[一二三四五六七八九十\d]+\s*章', line_stripped):
            chapter_end = j
            break
        if '参考文献' in line_stripped or '致谢' in line_stripped:
            chapter_end = j
            break
    
    if chapter_end == -1:
        chapter_end = min(chapter_start + 50, len(lines))
    
    chapter_lines = []
    for k in range(chapter_start, chapter_end):
        line = lines[k].strip()
        if line:
            chapter_lines.append(line)
    
    if chapter_lines:
        chapter_content = '\n'.join(chapter_lines)
        return chapter_content[:3000]
    
    return ""

def detect_project_type_by_llm(abstract: str, title: str = "", content: str = "") -> Dict:
    """
    混合检测论文类型：规则检测 + LLM辅助
    
    改进点：
    1. 优先使用多维度规则检测（关键词+章节+术语）
    2. 如果置信度不够高，再调用LLM精确判断
    3. 融合两种结果，提高准确率
    """
    try:
        from src.evaluation.thesis_type_detector import detect_thesis_type
        rule_result = detect_thesis_type(title, content, abstract)
        
        if rule_result["confidence"] > 0.6:
            logger.info(f"规则检测置信度足够: {rule_result['confidence']:.2f}")
            return rule_result
    except ImportError:
        logger.warning("论文类型检测器模块未找到，使用LLM检测")
        rule_result = None
    except Exception as e:
        logger.warning(f"规则检测失败: {str(e)}，使用LLM检测")
        rule_result = None
    
    if not abstract and not title and not content:
        return {"type": "mixed", "type_name": "混合类", "confidence": 0.5, "reason": "信息不足，无法判断"}
    
    from src.evaluation.llm_evaluator import LLMEvaluator
    
    content_preview = content[:2000] if content else ""
    
    prompt = f"""请根据以下论文的标题、摘要和关键内容，判断这篇毕业设计论文属于哪种类型。

论文标题：{title}

论文摘要：
{abstract}

关键内容预览：
{content_preview}

请从以下类型中选择最合适的一种：
1. algorithm - 算法类：主要涉及算法设计、模型开发、数据分析、机器学习等
2. simulation - 仿真类：主要涉及仿真分析、数值模拟、虚拟实验等
3. physical - 实物类：主要涉及硬件制作、电路设计、嵌入式系统、样机开发等
4. traditional_mechanical - 传统机械类：主要涉及机械结构设计、加工制造、传动机构等
5. mixed - 混合类：涉及以上多种类型的组合

请直接返回JSON格式：
{{"type": "类型代码", "type_name": "类型名称", "confidence": 0.95, "reason": "判断理由"}}
"""
    
    try:
        llm = LLMEvaluator()
        logger.info(f"开始调用大模型检测论文类型，标题: {title}")
        response = llm.generate_report(prompt, max_tokens=500)
        logger.info(f"大模型返回: {response}")
        
        import json
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group())
            logger.info(f"解析结果: {result}")
            
            if rule_result:
                rule_confidence = rule_result.get("confidence", 0.5)
                llm_confidence = result.get("confidence", 0.5)
                
                if rule_result.get("type") == result.get("type"):
                    result["confidence"] = min((rule_confidence + llm_confidence) / 2 + 0.1, 1.0)
                    result["reason"] = f"规则检测与LLM检测一致。{result.get('reason', '')}"
                else:
                    result["confidence"] = llm_confidence * 0.8
                    result["reason"] = f"LLM检测结果（与规则检测不一致）。{result.get('reason', '')}"
            
            return result
        else:
            logger.warning(f"无法从响应中提取JSON: {response}")
    except Exception as e:
        logger.error(f"大模型检测论文类型失败: {str(e)}", exc_info=True)
    
    if rule_result:
        return rule_result
    
    return {"type": "mixed", "type_name": "混合类", "confidence": 0.5, "reason": "无法确定具体类型"}

# 初始化数据库
init_db()

app = FastAPI(
    title="学生多维度能力评估系统 API",
    description="基于 OpenAI GPT-4o 和 CrewAI 的学生能力评估系统",
    version="1.0.0"
)

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 在生产环境中应该设置具体的域名
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 配置文件上传大小限制
app.max_request_size = 500 * 1024 * 1024  # 500MB

# 确保上传目录存在
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 创建数据库表
from src.database.database import Base, engine
Base.metadata.create_all(bind=engine)

# 依赖项
def get_database_service(db: Session = Depends(get_db)) -> DatabaseService:
    return DatabaseService(db)

# 根路径 - API 欢迎页面
@app.get("/")
async def root():
    return {
        "message": "欢迎使用学生多维度能力评估系统 API",
        "version": "1.0.0",
        "docs": "http://localhost:8000/docs",
        "health": "http://localhost:8000/health",
        "endpoints": {
            "students": "/students",
            "submissions": "/submissions",
            "evaluations": "/evaluate",
            "files": "/submissions/{submission_id}/files"
        }
    }

# 健康检查
@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}

# 学生相关路由
@app.post("/students", response_model=StudentResponse)
async def create_student(
    student: StudentCreate,
    db_service: DatabaseService = Depends(get_database_service)
):
    # 检查学生是否已存在
    existing_student = db_service.get_student_by_id(student.student_id)
    if existing_student:
        raise HTTPException(status_code=400, detail="学生已存在")
    
    # 创建学生
    new_student = db_service.create_student(
        student_id=student.student_id,
        name=student.name,
        age=student.age,
        grade=student.grade,
        major=student.major
    )
    
    return StudentResponse(
        id=new_student.id,
        student_id=new_student.student_id,
        name=new_student.name,
        age=new_student.age,
        grade=new_student.grade,
        major=new_student.major,
        created_at=new_student.created_at,
        updated_at=new_student.updated_at
    )

@app.get("/students/{student_id}", response_model=StudentResponse)
async def get_student(
    student_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    student = db_service.get_student_by_id(student_id)
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    return StudentResponse(
        id=student.id,
        student_id=student.student_id,
        name=student.name,
        age=student.age,
        grade=student.grade,
        major=student.major,
        created_at=student.created_at,
        updated_at=student.updated_at
    )

@app.get("/students", response_model=List[StudentResponse])
async def get_all_students(
    skip: int = 0,
    limit: int = 100,
    db_service: DatabaseService = Depends(get_database_service)
):
    students = db_service.get_all_students(skip=skip, limit=limit)
    return [
        StudentResponse(
            id=student.id,
            student_id=student.student_id,
            name=student.name,
            age=student.age,
            grade=student.grade,
            major=student.major,
            created_at=student.created_at,
            updated_at=student.updated_at
        )
        for student in students
    ]

@app.put("/students/{student_id}", response_model=StudentResponse)
async def update_student(
    student_id: str,
    student_update: StudentUpdate,
    db_service: DatabaseService = Depends(get_database_service)
):
    student = db_service.update_student(
        student_id=student_id,
        **student_update.dict(exclude_unset=True)
    )
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    return StudentResponse(
        id=student.id,
        student_id=student.student_id,
        name=student.name,
        age=student.age,
        grade=student.grade,
        major=student.major,
        created_at=student.created_at,
        updated_at=student.updated_at
    )

@app.delete("/students/{student_id}")
async def delete_student(
    student_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    try:
        deleted = db_service.delete_student(student_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="学生不存在")
        
        return {"message": "学生删除成功"}
    except Exception as e:
        logger.error(f"删除学生失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除学生失败: {str(e)}")

# 提交相关路由
@app.post("/submissions", response_model=SubmissionResponse)
async def create_submission(
    submission: SubmissionCreate,
    db_service: DatabaseService = Depends(get_database_service)
):
    # 检查学生是否存在（如果提供了 student_id）
    student = None
    if submission.student_id:
        student = db_service.get_student_by_id(submission.student_id)
        if not student:
            raise HTTPException(status_code=404, detail="学生不存在")
    
    # 验证提交类型
    if submission.submission_type == SubmissionType.TEXT and not submission.text_content:
        raise HTTPException(status_code=400, detail="文字提交必须提供内容")
    
    # 创建提交
    new_submission = db_service.create_submission(
        title=submission.title,
        description=submission.description,
        student_id=submission.student_id,
        submission_type=submission.submission_type.value,
        submission_purpose=submission.submission_purpose.value,
        text_content=submission.text_content
    )
    
    return SubmissionResponse(
        id=new_submission.id,
        submission_id=new_submission.submission_id,
        student_id=student.student_id if student else None,
        title=new_submission.title,
        description=new_submission.description,
        submission_type=SubmissionType(new_submission.submission_type),
        submission_purpose=SubmissionPurpose(new_submission.submission_purpose),
        text_content=new_submission.text_content,
        status=SubmissionStatus(new_submission.status),
        created_at=new_submission.created_at,
        updated_at=new_submission.updated_at
    )

@app.get("/submissions", response_model=List[SubmissionResponse])
async def get_all_submissions(
    skip: int = 0,
    limit: int = 100,
    db_service: DatabaseService = Depends(get_database_service)
):
    submissions = db_service.get_all_submissions(skip=skip, limit=limit)
    
    result = []
    for submission in submissions:
        student = None
        if submission.student_id:
            student = db_service.get_student_by_internal_id(submission.student_id)
        
        result.append(SubmissionResponse(
            id=submission.id,
            submission_id=submission.submission_id,
            student_id=student.student_id if student else None,
            title=submission.title,
            description=submission.description,
            submission_type=SubmissionType(submission.submission_type),
            submission_purpose=SubmissionPurpose(getattr(submission, 'submission_purpose', 'normal')),
            text_content=submission.text_content,
            status=SubmissionStatus(submission.status),
            created_at=submission.created_at,
            updated_at=submission.updated_at
        ))
    
    return result

@app.get("/submissions/{submission_id}", response_model=SubmissionResponse)
async def get_submission(
    submission_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    submission = db_service.get_submission_by_id(submission_id)
    if not submission:
        raise HTTPException(status_code=404, detail="提交不存在")
    
    student = None
    if submission.student_id:
        student = db_service.get_student_by_internal_id(submission.student_id)
        if not student:
            raise HTTPException(status_code=404, detail="学生不存在")
    
    return SubmissionResponse(
        id=submission.id,
        submission_id=submission.submission_id,
        student_id=student.student_id if student else None,
        title=submission.title,
        description=submission.description,
        submission_type=SubmissionType(submission.submission_type),
        submission_purpose=SubmissionPurpose(getattr(submission, 'submission_purpose', 'normal')),
        text_content=submission.text_content,
        status=SubmissionStatus(submission.status),
        created_at=submission.created_at,
        updated_at=submission.updated_at
    )

@app.get("/students/{student_id}/submissions", response_model=List[SubmissionResponse])
async def get_student_submissions(
    student_id: str,
    skip: int = 0,
    limit: int = 100,
    db_service: DatabaseService = Depends(get_database_service)
):
    submissions = db_service.get_submissions_by_student_id(student_id, skip=skip, limit=limit)
    
    return [
        SubmissionResponse(
            id=submission.id,
            submission_id=submission.submission_id,
            student_id=student_id,
            title=submission.title,
            description=submission.description,
            submission_type=SubmissionType(submission.submission_type),
            submission_purpose=SubmissionPurpose(getattr(submission, 'submission_purpose', 'normal')),
            text_content=submission.text_content,
            status=SubmissionStatus(submission.status),
            created_at=submission.created_at,
            updated_at=submission.updated_at
        )
        for submission in submissions
    ]

@app.delete("/submissions/{submission_id}")
async def delete_submission(
    submission_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """删除提交记录及其相关的媒体文件和评估结果"""
    deleted = db_service.delete_submission(submission_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="提交不存在")
    
    return {"message": "提交删除成功"}

# 手写文字识别路由
@app.post("/handwriting-recognize")
async def handwriting_recognize(
    student_id: str = Form(...),
    app_id: str = Form(...),
    api_key: str = Form(...),
    secret_key: str = Form(...),
    file: UploadFile = File(...),
    db_service: DatabaseService = Depends(get_database_service)
):
    """识别手写文字"""
    # 检查学生是否存在
    student = db_service.get_student_by_id(student_id)
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    # 保存上传的图片
    file_path = os.path.join(UPLOAD_DIR, f"handwriting_{student_id}_{file.filename}")
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # 检查文件类型
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in [".png", ".jpg", ".jpeg", ".bmp", ".webp", ".pdf"]:
            raise HTTPException(status_code=400, detail="不支持的文件类型")
        
        # 处理PDF文件
        if file_ext == ".pdf":
            # 使用PyMuPDF将PDF转换为图片，然后调用百度OCR识别
            logger.info("开始处理PDF文件")
            recognized_text = ""
            
            try:
                # 使用PyMuPDF将PDF转换为图片
                logger.info("使用PyMuPDF将PDF转换为图片")
                import fitz
                
                # 打开PDF文件
                doc = fitz.open(file_path)
                page_count = len(doc)
                logger.info(f"PDF文件共 {page_count} 页")
                
                # 将每一页转换为图片并调用百度OCR
                for page_num in range(page_count):
                    page = doc[page_num]
                    
                    # 将页面转换为图片
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 使用2倍分辨率提高识别率
                    img_data = pix.tobytes("png")
                    
                    # 调用百度OCR的手写文字识别接口
                    from aip import AipOcr
                    client = AipOcr(app_id, api_key, secret_key)
                    result = client.handwriting(img_data)
                    
                    # 打印识别结果用于调试
                    logger.info(f"第{page_num+1}页识别结果: {result}")
                    
                    # 提取识别结果
                    if 'words_result' in result and result['words_result']:
                        page_text = '\n'.join([item['words'] for item in result['words_result']])
                        recognized_text += f"第{page_num+1}页:\n{page_text}\n\n"
                    else:
                        # 如果没有识别到文字，记录错误信息
                        error_msg = result.get('error_msg', '未知错误')
                        error_code = result.get('error_code', '未知错误码')
                        logger.warning(f"第{page_num+1}页识别失败: {error_msg} (错误码: {error_code})")
                
                # 关闭PDF文档
                doc.close()
                
                logger.info(f"PDF识别完成，共处理 {page_count} 页")
                
                if not recognized_text:
                    raise HTTPException(status_code=400, detail="PDF文件中未提取到文字，请确保PDF包含可识别的文字内容")
                
                confidence = 90.0
                logger.info("PDF处理完成")
            except Exception as e:
                logger.error(f"PDF处理失败: {str(e)}")
                import traceback
                error_detail = traceback.format_exc()
                logger.error(f"详细错误: {error_detail}")
                raise HTTPException(status_code=500, detail=f"PDF处理失败: {str(e)}")
        else:
            # 使用百度OCR进行手写文字识别
            from aip import AipOcr
            
            # 使用前端传递的百度OCR API配置
            # 初始化百度OCR客户端
            client = AipOcr(app_id, api_key, secret_key)
            
            # 读取图片文件
            with open(file_path, 'rb') as f:
                image = f.read()
            
            # 调用百度OCR的手写文字识别接口
            result = client.handwriting(image)
            
            # 提取识别结果
            if 'words_result' in result:
                recognized_text = '\n'.join([item['words'] for item in result['words_result']])
                confidence = 95.0
            else:
                error_msg = result.get('error_msg', '未知错误')
                error_code = result.get('error_code', '未知错误码')
                raise HTTPException(status_code=400, detail=f"百度OCR识别失败：{error_msg} (错误码: {error_code})")
        
        # 保存识别记录
        record = {
            "student_id": student_id,
            "file_name": file.filename,
            "recognized_text": recognized_text,
            "confidence": confidence,
            "timestamp": datetime.now().isoformat()
        }
        db_service.add_handwriting_record(record)
        
        return {
            "student_id": student_id,
            "recognized_text": recognized_text,
            "confidence": confidence,
            "file_name": file.filename
        }
    except HTTPException:
        # 清理临时文件
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
        raise
    except Exception as e:
        # 清理临时文件
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
        # 记录详细错误信息
        import traceback
        error_detail = traceback.format_exc()
        print(f"识别失败: {error_detail}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {str(e)}")
    finally:
        # 清理临时文件
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass

# 文件上传路由
@app.post("/submissions/{submission_id}/files", response_model=MediaFileResponse)
async def upload_file(submission_id: int, file: UploadFile, db: Session = Depends(get_db)):
    # ... 之前原有的保存文件、提取 text_content 的代码 ...
    # 假设你已经把大纲内容提取到了变量 extracted_text 中

    # [新增逻辑] 自动分析课程类型
    detected_type = classify_course_type(extracted_text)

    # [新增逻辑] 将分类结果更新到数据库中
    submission = db.query(Submission).filter(Submission.id == submission_id).first()
    if submission:
        submission.course_type = detected_type
        db.commit()
        
    # 检查提交是否存在
    submission = db_service.get_submission_by_id(submission_id)
    if not submission:
        raise HTTPException(status_code=404, detail="提交不存在")
    
    # 保存文件
    file_path = os.path.join(UPLOAD_DIR, f"{submission_id}_{file.filename}")
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # 确定文件类型
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext in [".pdf", ".docx", ".doc", ".txt"]:
        media_type = "document"
    elif file_ext in [".mp4", ".mov"]:
        media_type = "video"
    elif file_ext in [".mp3", ".wav"]:
        media_type = "audio"
    else:
        raise HTTPException(status_code=400, detail="不支持的文件类型")
    
    # 获取文件大小
    size_bytes = os.path.getsize(file_path)
    
    # 创建媒体文件记录
    media_file = db_service.create_media_file(
        submission_id=submission_id,
        file_path=file_path,
        file_name=file.filename,
        media_type=media_type,
        size_bytes=size_bytes
    )
    
    return MediaFileResponse(
        id=media_file.id,
        submission_id=media_file.submission_id,
        file_path=media_file.file_path,
        file_name=media_file.file_name,
        media_type=media_file.media_type,
        size_bytes=media_file.size_bytes,
        duration=media_file.duration,
        processed=media_file.processed,
        uploaded_at=media_file.created_at
    )

@app.get("/submissions/{submission_id}/files", response_model=List[MediaFileResponse])
async def get_submission_files(
    submission_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    media_files = db_service.get_media_files_by_submission_id(submission_id)
    
    return [
        MediaFileResponse(
            id=file.id,
            submission_id=file.submission_id,
            file_path=file.file_path,
            file_name=file.file_name,
            media_type=file.media_type,
            size_bytes=file.size_bytes,
            duration=file.duration,
            processed=file.processed,
            uploaded_at=file.created_at
        )
        for file in media_files
    ]

# 文件修改路由
@app.put("/files/{file_id}", response_model=MediaFileResponse)
async def update_file(
    file_id: int,
    file_name: str = Form(..., description="文件名"),
    media_type: str = Form(..., description="文件类型") ,
    db_service: DatabaseService = Depends(get_database_service)
):
    # 验证输入
    if not file_name or not media_type:
        raise HTTPException(status_code=400, detail="文件名和文件类型为必填项")
    
    # 检查文件是否存在
    media_file = db_service.get_media_file_by_id(file_id)
    if not media_file:
        raise HTTPException(status_code=404, detail=f"文件ID {file_id} 不存在")
    
    try:
        # 更新文件信息
        updated_file = db_service.update_media_file(
            file_id=file_id,
            file_name=file_name,
            media_type=media_type
        )
        
        return MediaFileResponse(
            id=updated_file.id,
            submission_id=updated_file.submission_id,
            file_path=updated_file.file_path,
            file_name=updated_file.file_name,
            media_type=updated_file.media_type,
            size_bytes=updated_file.size_bytes,
            duration=updated_file.duration,
            processed=updated_file.processed,
            uploaded_at=updated_file.created_at
        )
    except Exception as e:
        logger.error(f"更新文件信息失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"更新文件信息失败: {str(e)}")

# 文件删除路由
@app.delete("/files/{file_id}")
async def delete_file(
    file_id: int,
    db_service: DatabaseService = Depends(get_database_service)
):
    deleted = db_service.delete_media_file(file_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return {"message": "文件删除成功"}

# 评估相关路由
@app.post("/analyze_syllabus")
async def analyze_syllabus(
    request: Dict = Body(...)
):
    """
    使用大模型分析课程大纲
    自动检测文档类型并使用对应的提示词进行分析
    """
    try:
        # 导入大模型评估器
        from src.evaluation.llm_evaluator import llm_evaluator
        from src.evaluation.syllabus_analyzer import SyllabusAnalyzer
        
        # 从请求中获取参数
        syllabus_content = request.get('syllabus_content', '')
        syllabus_name = request.get('syllabus_name', '')
        
        # 使用 SyllabusAnalyzer 检测文档类型并构建对应提示词
        analyzer = SyllabusAnalyzer("")
        doc_type = analyzer.detect_document_type(syllabus_content)
        
        print(f"检测到文档类型: {doc_type}")
        
        # 根据文档类型选择对应的提示词
        if doc_type == 'graduation_requirements':
            prompt = analyzer.build_graduation_requirements_prompt(syllabus_content)
        elif doc_type == 'course_evaluation':
            prompt = analyzer.build_course_evaluation_prompt(syllabus_content)
        else:
            # 未知类型使用通用提示词
            prompt = analyzer.build_syllabus_analysis_prompt(syllabus_content)
        
        # 调用大模型（增加max_tokens以获取更详细的分析结果）
        result = llm_evaluator.generate_report(prompt, max_tokens=4000)
        
        # 解析大模型返回结果
        try:
            # 尝试直接解析JSON
            analysis_result = json.loads(result)
            # 添加文档类型信息
            analysis_result['document_type'] = doc_type
            # 检查分析结果是否为空
            has_content = (
                analysis_result.get('ability_points') or 
                analysis_result.get('evaluation_criteria') or 
                analysis_result.get('graduation_requirements')
            )
            if not analysis_result or not has_content:
                raise Exception("未获取到有效的分析结果")
            return analysis_result
        except Exception as e:
            # 如果解析失败，尝试从文本中提取JSON
            try:
                # 查找JSON开始和结束的位置
                start_idx = result.find('{')
                end_idx = result.rfind('}') + 1
                if start_idx != -1 and end_idx != -1:
                    json_str = result[start_idx:end_idx]
                    analysis_result = json.loads(json_str)
                    # 添加文档类型信息
                    analysis_result['document_type'] = doc_type
                    # 检查分析结果是否为空
                    has_content = (
                        analysis_result.get('ability_points') or 
                        analysis_result.get('evaluation_criteria') or 
                        analysis_result.get('graduation_requirements')
                    )
                    if not analysis_result or not has_content:
                        raise Exception("未获取到有效的分析结果")
                    return analysis_result
                else:
                    raise Exception("未找到有效的JSON格式结果")
            except Exception as e2:
                # 如果仍然解析失败，直接报错
                raise Exception(f"解析大模型返回结果失败: {str(e2)}")
    except Exception as e:
        # 直接返回错误，不使用模拟数据
        error_message = str(e)
        if "API密钥未设置" in error_message:
            raise HTTPException(status_code=400, detail=f"大纲分析失败: API密钥未设置，请在AI设置页面中配置API密钥")
        elif "400" in error_message or "Bad Request" in error_message:
            raise HTTPException(status_code=400, detail=f"大纲分析失败: 请求参数错误，请检查输入内容")
        elif "timeout" in error_message.lower() or "timed out" in error_message.lower():
            raise HTTPException(status_code=504, detail=f"大纲分析失败: 请求超时，请稍后重试")
        else:
            raise HTTPException(status_code=500, detail=f"大纲分析失败: {error_message}")

@app.post("/analyze_graduation_project")
async def analyze_graduation_project(
    request: Dict = Body(...)
):
    """
    多轮迭代分析毕业设计大纲
    通过三轮调用大模型，逐步完善评价提示词
    """
    try:
        from src.evaluation.llm_evaluator import llm_evaluator
        from src.evaluation.syllabus_analyzer import SyllabusAnalyzer
        
        syllabus_content = request.get('syllabus_content', '')
        syllabus_name = request.get('syllabus_name', '')
        max_rounds = request.get('max_rounds', 3)
        
        analyzer = SyllabusAnalyzer("")
        results = {
            "syllabus_name": syllabus_name,
            "rounds": [],
            "final_result": None
        }
        
        # 第一轮：初步分析
        print(f"开始第一轮分析: {syllabus_name}")
        round1_prompt = analyzer.build_graduation_project_initial_analysis_prompt(syllabus_content)
        round1_result = llm_evaluator.generate_report(round1_prompt, max_tokens=3000)
        
        try:
            initial_result = json.loads(round1_result)
            results["rounds"].append({
                "round": 1,
                "type": "initial_analysis",
                "result": initial_result
            })
            print(f"第一轮分析完成")
        except Exception as e:
            raise Exception(f"第一轮分析结果解析失败: {str(e)}")
        
        # 第二轮：详细分析
        print(f"开始第二轮分析: {syllabus_name}")
        round2_prompt = analyzer.build_graduation_project_detailed_analysis_prompt(syllabus_content, initial_result)
        round2_result = llm_evaluator.generate_report(round2_prompt, max_tokens=4000)
        
        try:
            detailed_result = json.loads(round2_result)
            results["rounds"].append({
                "round": 2,
                "type": "detailed_analysis",
                "result": detailed_result
            })
            print(f"第二轮分析完成")
        except Exception as e:
            raise Exception(f"第二轮分析结果解析失败: {str(e)}")
        
        # 第三轮：提示词优化
        if max_rounds >= 3:
            print(f"开始第三轮分析: {syllabus_name}")
            round3_prompt = analyzer.build_graduation_project_prompt_refinement_prompt(syllabus_content, detailed_result)
            round3_result = llm_evaluator.generate_report(round3_prompt, max_tokens=4000)
            
            try:
                refined_result = json.loads(round3_result)
                results["rounds"].append({
                    "round": 3,
                    "type": "prompt_refinement",
                    "result": refined_result
                })
                print(f"第三轮分析完成")
                
                # 使用优化后的结果作为最终结果
                results["final_result"] = {
                    "ability_points": refined_result.get("optimized_ability_points", []),
                    "evaluation_criteria": refined_result.get("optimized_evaluation_criteria", []),
                    "evaluation_prompt_template": refined_result.get("evaluation_prompt_template", ""),
                    "optimization_notes": refined_result.get("optimization_notes", "")
                }
            except Exception as e:
                # 如果第三轮失败，使用第二轮结果
                print(f"第三轮分析结果解析失败: {str(e)}，使用第二轮结果")
                results["final_result"] = detailed_result
        else:
            results["final_result"] = detailed_result
        
        return results
        
    except Exception as e:
        error_message = str(e)
        if "API密钥未设置" in error_message:
            raise HTTPException(status_code=400, detail=f"毕业设计分析失败: API密钥未设置，请在AI设置页面中配置API密钥")
        elif "400" in error_message or "Bad Request" in error_message:
            raise HTTPException(status_code=400, detail=f"毕业设计分析失败: 请求参数错误，请检查输入内容")
        elif "timeout" in error_message.lower() or "timed out" in error_message.lower():
            raise HTTPException(status_code=504, detail=f"毕业设计分析失败: 请求超时，请稍后重试")
        else:
            raise HTTPException(status_code=500, detail=f"毕业设计分析失败: {error_message}")

@app.post("/evaluate_graduation_project")
async def evaluate_graduation_project(
    request: Dict = Body(...)
):
    """
    使用确定性评价标准评价毕业设计
    自动检测项目类型并使用对应的评价标准
    """
    try:
        from src.evaluation.llm_evaluator import llm_evaluator
        from src.evaluation.evaluation_standards import (
            ProjectType,
            PROJECT_TYPE_NAMES,
            detect_project_type
        )
        
        submission_content = request.get('submission_content', '')
        project_type = request.get('project_type', None)
        student_info = request.get('student_info', {})
        guidance_content = request.get('guidance_content', None)
        
        if not submission_content:
            raise HTTPException(status_code=400, detail="提交内容不能为空")
        
        title = student_info.get("title", "") if student_info else ""
        
        if project_type:
            try:
                detected_type = ProjectType(project_type)
            except ValueError:
                detected_type = detect_project_type(title, submission_content)
        else:
            detected_type = detect_project_type(title, submission_content)
        
        type_name = PROJECT_TYPE_NAMES.get(detected_type, "未知类型")
        print(f"检测到项目类型: {type_name}")
        
        result = llm_evaluator.evaluate_with_deterministic_standards(
            submission_content=submission_content,
            project_type=detected_type.value,
            student_info=student_info,
            guidance_content=guidance_content
        )
        
        result["detected_project_type"] = detected_type.value
        result["project_type_name"] = type_name
        
        return result
        
    except Exception as e:
        error_message = str(e)
        if "API密钥未设置" in error_message:
            raise HTTPException(status_code=400, detail=f"评价失败: API密钥未设置，请在AI设置页面中配置API密钥")
        elif "400" in error_message or "Bad Request" in error_message:
            raise HTTPException(status_code=400, detail=f"评价失败: 请求参数错误，请检查输入内容")
        elif "timeout" in error_message.lower() or "timed out" in error_message.lower():
            raise HTTPException(status_code=504, detail=f"评价失败: 请求超时，请稍后重试")
        else:
            raise HTTPException(status_code=500, detail=f"评价失败: {error_message}")

@app.get("/project_types")
async def get_project_types():
    """
    获取所有项目类型
    """
    from src.evaluation.evaluation_standards import ProjectType, PROJECT_TYPE_NAMES, EVALUATION_STANDARDS
    
    types = []
    for pt in ProjectType:
        if pt != ProjectType.UNKNOWN:
            type_config = EVALUATION_STANDARDS.get(pt.value, {})
            types.append({
                "value": pt.value,
                "name": PROJECT_TYPE_NAMES.get(pt, "未知类型"),
                "description": type_config.get("description", ""),
                "indicators_count": len(type_config.get("indicators", []))
            })
    
    return {"project_types": types}

@app.get("/evaluation_standards/{project_type}")
async def get_evaluation_standards_by_type(project_type: str):
    """
    获取指定项目类型的评价标准
    """
    from src.evaluation.evaluation_standards import (
        ProjectType,
        PROJECT_TYPE_NAMES,
        EVALUATION_STANDARDS
    )
    
    try:
        pt = ProjectType(project_type)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"无效的项目类型: {project_type}")
    
    standards = EVALUATION_STANDARDS.get(pt.value)
    if not standards:
        raise HTTPException(status_code=404, detail=f"未找到项目类型 {project_type} 的评价标准")
    
    return {
        "project_type": pt.value,
        "project_type_name": PROJECT_TYPE_NAMES.get(pt, "未知类型"),
        "standards": standards
    }

@app.post("/extract_guidance_content")
async def extract_guidance_content(
    request: Dict = Body(...)
):
    """
    使用大模型提炼评价指导文件的内容
    """
    try:
        from src.evaluation.llm_evaluator import llm_evaluator
        
        file_content = request.get('file_content', '')
        file_name = request.get('file_name', '')
        
        if not file_content:
            raise HTTPException(status_code=400, detail="文件内容不能为空")
        
        result = llm_evaluator.extract_guidance_content(
            file_content=file_content,
            file_name=file_name
        )
        
        return result
        
    except Exception as e:
        error_message = str(e)
        if "API密钥未设置" in error_message:
            raise HTTPException(status_code=400, detail=f"提炼失败: API密钥未设置，请在AI设置页面中配置API密钥")
        elif "timeout" in error_message.lower() or "timed out" in error_message.lower():
            raise HTTPException(status_code=504, detail=f"提炼失败: 请求超时，请稍后重试")
        else:
            raise HTTPException(status_code=500, detail=f"提炼失败: {error_message}")

@app.post("/generate_evaluation_standards")
async def generate_evaluation_standards(
    request: Dict = Body(...)
):
    """
    使用大模型生成项目评价指标
    """
    try:
        from src.evaluation.llm_evaluator import llm_evaluator
        
        file_content = request.get('file_content', '')
        file_name = request.get('file_name', '')
        project_type = request.get('project_type', 'mixed')
        
        if not file_content:
            raise HTTPException(status_code=400, detail="文件内容不能为空")
        
        result = llm_evaluator.generate_evaluation_standards(
            file_content=file_content,
            file_name=file_name,
            project_type=project_type
        )
        
        return result
        
    except Exception as e:
        error_message = str(e)
        if "API密钥未设置" in error_message:
            raise HTTPException(status_code=400, detail=f"生成失败: API密钥未设置，请在AI设置页面中配置API密钥")
        elif "timeout" in error_message.lower() or "timed out" in error_message.lower():
            raise HTTPException(status_code=504, detail=f"生成失败: 请求超时，请稍后重试")
        else:
            raise HTTPException(status_code=500, detail=f"生成失败: {error_message}")

@app.get("/syllabus_files")
async def get_syllabus_files():
    """
    获取大纲管理页面的文件列表
    """
    import os
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    syllabus_folder = os.path.join(project_root, "评价大纲")
    
    if not os.path.exists(syllabus_folder):
        return {"files": [], "folder_path": syllabus_folder}
    
    files = []
    for f in os.listdir(syllabus_folder):
        if f.endswith('.docx') or f.endswith('.doc') or f.endswith('.pdf') or f.endswith('.txt'):
            file_path = os.path.join(syllabus_folder, f)
            files.append({
                "name": f,
                "path": file_path,
                "size": os.path.getsize(file_path),
                "modified": os.path.getmtime(file_path)
            })
    
    return {"files": files, "folder_path": syllabus_folder}

@app.post("/read_syllabus_file")
async def read_syllabus_file(
    request: Dict = Body(...)
):
    """
    读取大纲文件内容
    """
    import os
    file_name = request.get('file_name', '')
    
    if not file_name:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    file_path = os.path.join(project_root, "评价大纲", file_name)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail=f"文件不存在: {file_name}")
    
    try:
        if file_name.endswith('.docx'):
            from docx import Document
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
        elif file_name.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif file_name.endswith('.pdf'):
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式: {file_name}")
        
        return {"file_name": file_name, "content": content}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取文件失败: {str(e)}")

@app.post("/evaluate_with_rule_engine")
async def evaluate_with_rule_engine(
    request: Dict = Body(...)
):
    """
    使用大模型进行评分（根据评价指标）
    支持校方固有评价体系融合评分
    总共调用2次大模型：
    1. 原始评分：根据评价指标评分
    2. 融合系数：固有评价体系评分
    """
    try:
        submission_content = request.get('submission_content', '')
        indicators = request.get('indicators', {})
        student_info = request.get('student_info', {})
        dimension_weights = request.get('dimension_weights', {})
        
        if not submission_content:
            raise HTTPException(status_code=400, detail="提交内容不能为空")
        
        if not indicators:
            raise HTTPException(status_code=400, detail="评价指标不能为空，请先选择评价指标")
        
        from src.evaluation.llm_evaluator import llm_evaluator
        
        print("正在进行原始评分（第1次大模型调用）...")
        result = llm_evaluator.evaluate_with_indicators(
            submission_content=submission_content,
            indicators=indicators,
            student_info=student_info
        )
        
        result["student_info"] = student_info
        
        strengths = []
        weaknesses = []
        
        for dim in result.get("dimension_scores", []):
            if dim.get("score", 0) >= 80:
                strengths.append(f"{dim.get('indicator_name', dim.get('indicator_id', ''))}表现良好({dim.get('score')}分)")
            elif dim.get("score", 0) < 60:
                weaknesses.append(f"{dim.get('indicator_name', dim.get('indicator_id', ''))}需要改进({dim.get('score')}分)")
        
        result["strengths"] = strengths if strengths else ["整体表现符合要求"]
        result["weaknesses"] = weaknesses if weaknesses else ["无明显短板"]
        
        coefficient_config = request.get('coefficient_config', {})
        use_custom_coefficients = request.get('use_custom_coefficients', False)
        
        if dimension_weights:
            try:
                print("正在进行校方固有评价体系评分（第2次大模型调用）...")
                institutional_result = llm_evaluator.evaluate_institutional_dimensions(
                    submission_content=submission_content,
                    dimension_weights=dimension_weights
                )
                
                result["institutional_evaluation"] = institutional_result
                
                original_score = result.get("overall_score", 0)
                fusion_result = llm_evaluator.calculate_fusion_score(
                    rule_engine_score=original_score,
                    institutional_result=institutional_result,
                    coefficient_config=coefficient_config if use_custom_coefficients else None
                )
                
                result["original_score"] = fusion_result["original_score"]
                result["fusion_coefficient"] = fusion_result["fusion_coefficient"]
                result["adjustment"] = fusion_result["adjustment"]
                result["fusion_details"] = {
                    "adjustment": fusion_result["adjustment"],
                    "dimension_coefficients": fusion_result["dimension_coefficients"],
                    "coefficient_config_used": fusion_result.get("coefficient_config_used", {})
                }
                
                result["overall_score"] = fusion_result["fusion_score"]
                
                if fusion_result["fusion_score"] >= 90:
                    result["grade_level"] = "优秀"
                elif fusion_result["fusion_score"] >= 80:
                    result["grade_level"] = "良好"
                elif fusion_result["fusion_score"] >= 70:
                    result["grade_level"] = "中等"
                elif fusion_result["fusion_score"] >= 60:
                    result["grade_level"] = "及格"
                else:
                    result["grade_level"] = "不及格"
                
                print(f"融合评分完成: 原始{fusion_result['original_score']}分 -> 融合后{fusion_result['fusion_score']}分")
                
            except Exception as e:
                print(f"固有评价体系评分失败: {str(e)}")
                result["institutional_evaluation_error"] = str(e)
        
        return result
        
    except Exception as e:
        error_message = str(e)
        raise HTTPException(status_code=500, detail=f"评分失败: {error_message}")

@app.post("/evaluate_sectioned")
async def evaluate_sectioned(
    request: Dict = Body(...)
):
    """
    分段评估论文（智能分段 + 章节衔接检测）
    适用于长篇论文，解决内容丢失问题
    
    调用流程：
    1. 大模型识别论文结构
    2. 提取各章节内容
    3. 各章节独立评估（带上下文）
    4. 章节衔接检测
    5. 固有评价体系评分
    6. 生成最终评价
    """
    try:
        submission_content = request.get('submission_content', '')
        indicators = request.get('indicators', {})
        student_info = request.get('student_info', {})
        dimension_weights = request.get('dimension_weights', {})
        coefficient_config = request.get('coefficient_config', {})
        use_custom_coefficients = request.get('use_custom_coefficients', False)
        
        if not submission_content:
            raise HTTPException(status_code=400, detail="提交内容不能为空")
        
        from src.evaluation.llm_evaluator import llm_evaluator
        from src.evaluation.sectioned_evaluator import SectionedEvaluator
        
        print("开始分段评估论文...")
        print("=" * 50)
        
        sectioned_evaluator = SectionedEvaluator(llm_evaluator)
        
        result = sectioned_evaluator.evaluate_thesis_sectioned(
            content=submission_content,
            indicators=indicators,
            student_info=student_info,
            dimension_weights=dimension_weights if dimension_weights else None
        )
        
        if dimension_weights and result.get("institutional_evaluation"):
            institutional_result = result["institutional_evaluation"]
            original_score = result.get("overall_score", 0)
            
            fusion_result = llm_evaluator.calculate_fusion_score(
                rule_engine_score=original_score,
                institutional_result=institutional_result,
                coefficient_config=coefficient_config if use_custom_coefficients else None
            )
            
            result["original_score"] = fusion_result["original_score"]
            result["fusion_coefficient"] = fusion_result["fusion_coefficient"]
            result["adjustment"] = fusion_result["adjustment"]
            result["fusion_details"] = {
                "adjustment": fusion_result["adjustment"],
                "dimension_coefficients": fusion_result["dimension_coefficients"],
                "coefficient_config_used": fusion_result.get("coefficient_config_used", {})
            }
            
            result["overall_score"] = fusion_result["fusion_score"]
            
            if fusion_result["fusion_score"] >= 90:
                result["grade_level"] = "优秀"
            elif fusion_result["fusion_score"] >= 80:
                result["grade_level"] = "良好"
            elif fusion_result["fusion_score"] >= 70:
                result["grade_level"] = "中等"
            elif fusion_result["fusion_score"] >= 60:
                result["grade_level"] = "及格"
            else:
                result["grade_level"] = "不及格"
            
            print(f"融合评分完成: 原始{fusion_result['original_score']}分 -> 融合后{fusion_result['fusion_score']}分")
        
        print("=" * 50)
        print(f"分段评估完成: {result.get('overall_score', 0)}分")
        
        return result
        
    except Exception as e:
        error_message = str(e)
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"分段评估失败: {error_message}")

@app.get("/evaluations/{evaluation_id}", response_model=EvaluationResponse)
async def get_evaluation(
    evaluation_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    # 获取评估结果
    evaluation = db_service.get_evaluation_result_by_id(evaluation_id)
    if not evaluation:
        raise HTTPException(status_code=404, detail="评估结果不存在")
    
    # 获取学生信息
    student = db_service.get_student_by_internal_id(evaluation.student_id)
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    # 获取维度评分
    dimension_scores = db_service.get_dimension_scores_by_evaluation_id(evaluation_id)
    
    # 构建维度评分响应
    dimension_scores_response = [
        DimensionScoreResponse(
            dimension=ds.dimension,
            score=ds.score,
            confidence=ds.confidence,
            evidence=process_evidence(ds.evidence),
            reasoning=ds.reasoning
        )
        for ds in dimension_scores
    ]
    
    # 检查evaluation对象是否有stage和stage_progress属性
    stage = None
    stage_progress = None
    
    if hasattr(evaluation, 'stage'):
        stage = evaluation.stage
    
    if hasattr(evaluation, 'stage_progress'):
        stage_progress = evaluation.stage_progress
    
    return EvaluationResponse(
        evaluation_id=evaluation.evaluation_id,
        student_id=student.student_id,
        overall_score=evaluation.overall_score,
        strengths=process_string_list(evaluation.strengths),
        areas_for_improvement=process_string_list(evaluation.areas_for_improvement),
        recommendations=process_string_list(evaluation.recommendations),
        dimension_scores=dimension_scores_response,
        evaluated_at=evaluation.evaluated_at,
        evaluator_agent=evaluation.evaluator_agent,
        stage=stage,
        stage_progress=stage_progress
    )

@app.get("/students/{student_id}/evaluations", response_model=List[EvaluationResponse])
async def get_student_evaluations(
    student_id: str,
    skip: int = 0,
    limit: int = 100,
    db_service: DatabaseService = Depends(get_database_service)
):
    try:
        # 获取学生的评估结果，按时间排序
        evaluations = db_service.get_evaluation_results_by_student_id_sorted(student_id)
        
        response = []
        for evaluation in evaluations:
            # 获取维度评分
            dimension_scores = db_service.get_dimension_scores_by_evaluation_id(evaluation.evaluation_id)
            
            # 构建维度评分响应
            dimension_scores_response = [
                DimensionScoreResponse(
                    dimension=ds.dimension,
                    score=ds.score,
                    confidence=ds.confidence,
                    evidence=process_evidence(ds.evidence),
                    reasoning=ds.reasoning
                )
                for ds in dimension_scores
            ]
            
            # 检查evaluation对象是否有stage和stage_progress属性
            stage = None
            stage_progress = None
            
            if hasattr(evaluation, 'stage'):
                stage = evaluation.stage
            
            if hasattr(evaluation, 'stage_progress'):
                stage_progress = evaluation.stage_progress
            
            # 获取学生信息
            student = db_service.get_student_by_internal_id(evaluation.student_id)
            if not student:
                continue
            
            response.append(EvaluationResponse(
                evaluation_id=evaluation.evaluation_id,
                student_id=student.student_id,
                overall_score=evaluation.overall_score,
                strengths=process_string_list(evaluation.strengths),
                areas_for_improvement=process_string_list(evaluation.areas_for_improvement),
                recommendations=process_string_list(evaluation.recommendations),
                dimension_scores=dimension_scores_response,
                evaluated_at=evaluation.evaluated_at,
                evaluator_agent=evaluation.evaluator_agent,
                stage=stage,
                stage_progress=stage_progress
            ))
        
        return response
    except Exception as e:
        logger.error(f"获取学生评估记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取学生评估记录失败: {str(e)}")

@app.delete("/evaluations/{evaluation_id}")
async def delete_evaluation(
    evaluation_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """删除评估记录及其相关的维度评分"""
    success = db_service.delete_evaluation_result(evaluation_id)
    if not success:
        raise HTTPException(status_code=404, detail="评估记录不存在")
    
    return {"message": "评估记录已成功删除"}

@app.put("/evaluations/{evaluation_id}", response_model=EvaluationResponse)
async def update_evaluation(
    evaluation_id: str,
    update_data: dict,
    db_service: DatabaseService = Depends(get_database_service)
):
    """更新评估记录"""
    # 更新评估结果
    evaluation = db_service.update_evaluation_result(evaluation_id, **update_data)
    if not evaluation:
        raise HTTPException(status_code=404, detail="评估结果不存在")
    
    # 获取学生信息
    student = db_service.get_student_by_internal_id(evaluation.student_id)
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    # 获取维度评分
    dimension_scores = db_service.get_dimension_scores_by_evaluation_id(evaluation_id)
    
    # 构建维度评分响应
    dimension_scores_response = [
        DimensionScoreResponse(
            dimension=ds.dimension,
            score=ds.score,
            confidence=ds.confidence,
            evidence=process_evidence(ds.evidence),
            reasoning=ds.reasoning
        )
        for ds in dimension_scores
    ]
    
    # 检查evaluation对象是否有stage和stage_progress属性
    stage = None
    stage_progress = None
    
    if hasattr(evaluation, 'stage'):
        stage = evaluation.stage
    
    if hasattr(evaluation, 'stage_progress'):
        stage_progress = evaluation.stage_progress
    
    return EvaluationResponse(
        evaluation_id=evaluation.evaluation_id,
        student_id=student.student_id,
        overall_score=evaluation.overall_score,
        strengths=process_string_list(evaluation.strengths),
        areas_for_improvement=process_string_list(evaluation.areas_for_improvement),
        recommendations=process_string_list(evaluation.recommendations),
        dimension_scores=dimension_scores_response,
        evaluated_at=evaluation.evaluated_at,
        evaluator_agent=evaluation.evaluator_agent,
        stage=stage,
        stage_progress=stage_progress
    )

@app.get("/students/{student_id}/progress-report", response_model=ProgressReportResponse)
async def generate_student_progress_report(
    student_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """
    生成学生的整体进度报告，根据之前的不同进度的作业评价，
    将其输入大模型后获得对该学生的，在时间线上的能力进步
    """
    # 获取学生的所有评估结果，按时间排序
    evaluations = db_service.get_evaluation_results_by_student_id_sorted(student_id)
    
    if not evaluations:
        raise HTTPException(status_code=404, detail="该学生没有评估记录")
    
    # 构建评估历史数据
    evaluation_history = []
    for eval in evaluations:
        dimension_scores = db_service.get_dimension_scores_by_evaluation_id(eval.evaluation_id)
        dimension_data = [
            {
                "dimension": ds.dimension,
                "score": ds.score,
                "confidence": ds.confidence,
                "reasoning": ds.reasoning
            }
            for ds in dimension_scores
        ]
        
        evaluation_history.append({
            "evaluation_id": eval.evaluation_id,
            "evaluated_at": eval.evaluated_at.isoformat(),
            "overall_score": eval.overall_score,
            "strengths": process_string_list(eval.strengths),
            "areas_for_improvement": process_string_list(eval.areas_for_improvement),
            "recommendations": process_string_list(eval.recommendations),
            "dimension_scores": dimension_data
        })
    
    # 构建提示词
    prompt = f"""你是一位专业的教育评估专家，擅长分析学生在时间线上的能力进步。

请根据以下学生的评估历史数据，生成一份详细的整体进度报告：

学生ID: {student_id}

评估历史（按时间顺序）：
{json.dumps(evaluation_history, ensure_ascii=False, indent=2)}

报告要求：
1. 分析学生在各个维度上的能力变化趋势
2. 识别学生的优势和持续改进的领域
3. 提供关于学生能力发展的关键洞察
4. 给出基于历史数据的未来发展建议
5. 报告应该结构清晰，语言专业但易于理解
6. 包含具体的数据支持和分析

请生成一份全面的进度报告，帮助教师和学生了解能力发展情况。"""
    
    try:
        # 调用大模型生成报告
        from src.evaluation.llm_evaluator import llm_evaluator
        report = llm_evaluator.generate_report(prompt)
        
        # 提取关键信息
        key_insights = ["学生能力发展趋势分析", "优势领域识别", "改进空间分析"]
        improvement_areas = ["需要持续关注的能力维度"]
        
        # 构建时间范围
        time_range = {
            "start": evaluations[0].evaluated_at.isoformat(),
            "end": evaluations[-1].evaluated_at.isoformat()
        }
        
        # 保存进度报告到数据库
        db_service.create_progress_report(
            student_id=student_id,
            report=report,
            total_evaluations=len(evaluations),
            time_range=time_range,
            key_insights=key_insights,
            improvement_areas=improvement_areas
        )
        
        return ProgressReportResponse(
            student_id=student_id,
            report=report,
            generated_at=datetime.utcnow(),
            total_evaluations=len(evaluations),
            time_range=time_range,
            key_insights=key_insights,
            improvement_areas=improvement_areas
        )
    except Exception as e:
        logger.error(f"生成进度报告失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"生成进度报告失败: {str(e)}")

@app.get("/students/{student_id}/progress-reports", response_model=List[ProgressReportResponse])
async def get_student_progress_reports(
    student_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """
    获取学生的历史进度报告
    """
    import json
    
    # 获取学生的所有进度报告
    reports = db_service.get_progress_reports_by_student_id(student_id)
    
    if not reports:
        return []
    
    # 构建响应
    response = []
    for report in reports:
        # 解析 JSON 字段
        time_range = {}
        key_insights = []
        improvement_areas = []
        
        try:
            if report.time_range:
                time_range = json.loads(report.time_range)
            if report.key_insights:
                key_insights = json.loads(report.key_insights)
            if report.improvement_areas:
                improvement_areas = json.loads(report.improvement_areas)
        except:
            pass
        
        response.append(ProgressReportResponse(
            student_id=student_id,
            report=report.report,
            generated_at=report.generated_at,
            total_evaluations=report.total_evaluations,
            time_range=time_range,
            key_insights=key_insights,
            improvement_areas=improvement_areas,
            report_id=report.report_id  # 添加report_id
        ))
    
    return response

# Pydantic model for updating progress report
class ProgressReportUpdate(BaseModel):
    generated_at: Optional[datetime] = None
    report: Optional[str] = None
    time_range: Optional[Dict[str, str]] = None
    key_insights: Optional[List[str]] = None
    improvement_areas: Optional[List[str]] = None

@app.put("/progress-reports/{report_id}")
async def update_progress_report(
    report_id: str,
    report_update: ProgressReportUpdate,
    db_service: DatabaseService = Depends(get_database_service)
):
    """
    更新进度报告
    """
    # 构建更新数据
    update_data = {}
    if report_update.generated_at is not None:
        update_data['generated_at'] = report_update.generated_at
    if report_update.report is not None:
        update_data['report'] = report_update.report
    if report_update.time_range is not None:
        update_data['time_range'] = report_update.time_range
    if report_update.key_insights is not None:
        update_data['key_insights'] = report_update.key_insights
    if report_update.improvement_areas is not None:
        update_data['improvement_areas'] = report_update.improvement_areas
    
    # 更新报告
    updated_report = db_service.update_progress_report(report_id, **update_data)
    
    if not updated_report:
        raise HTTPException(status_code=404, detail="报告不存在")
    
    # 构建响应
    import json
    time_range = {}
    key_insights = []
    improvement_areas = []
    
    try:
        if updated_report.time_range:
            time_range = json.loads(updated_report.time_range)
        if updated_report.key_insights:
            key_insights = json.loads(updated_report.key_insights)
        if updated_report.improvement_areas:
            improvement_areas = json.loads(updated_report.improvement_areas)
    except:
        pass
    
    return ProgressReportResponse(
        student_id="",  # 这里需要从报告中获取学生ID，但当前模型没有直接关联
        report=updated_report.report,
        generated_at=updated_report.generated_at,
        total_evaluations=updated_report.total_evaluations,
        time_range=time_range,
        key_insights=key_insights,
        improvement_areas=improvement_areas,
        report_id=updated_report.report_id
    )

@app.delete("/progress-reports/{report_id}")
async def delete_progress_report(
    report_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """删除进度报告"""
    deleted = db_service.delete_progress_report(report_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="进度报告不存在")
    
    return {"message": "进度报告删除成功"}

@app.get("/submissions/{submission_id}/evaluation", response_model=EvaluationResponse)
async def get_submission_evaluation(
    submission_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    # 获取提交的评估结果
    evaluation = db_service.get_evaluation_result_by_submission_id(submission_id)
    if not evaluation:
        raise HTTPException(status_code=404, detail="评估结果不存在")
    
    # 获取学生信息
    student = db_service.get_student_by_internal_id(evaluation.student_id)
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    # 获取维度评分
    dimension_scores = db_service.get_dimension_scores_by_evaluation_id(evaluation.evaluation_id)
    
    # 构建维度评分响应
    dimension_scores_response = [
        DimensionScoreResponse(
            dimension=ds.dimension,
            score=ds.score,
            confidence=ds.confidence,
            evidence=process_evidence(ds.evidence),
            reasoning=ds.reasoning
        )
        for ds in dimension_scores
    ]
    
    # 检查evaluation对象是否有stage属性
    stage = None
    if hasattr(evaluation, 'stage'):
        stage = evaluation.stage
    
    return EvaluationResponse(
            evaluation_id=evaluation.evaluation_id,
            student_id=student.student_id,
            overall_score=evaluation.overall_score,
            strengths=process_string_list(evaluation.strengths),
            areas_for_improvement=process_string_list(evaluation.areas_for_improvement),
            recommendations=process_string_list(evaluation.recommendations),
            dimension_scores=dimension_scores_response,
            evaluated_at=evaluation.evaluated_at,
            evaluator_agent=evaluation.evaluator_agent,
            stage=stage
        )

@app.get("/students/{student_id}/evaluations/comparison")
async def compare_student_evaluations(
    student_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """获取学生的评估结果对比"""
    # 获取学生信息
    student = db_service.get_student_by_id(student_id)
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    # 获取按阶段排序的评估结果
    evaluations = db_service.get_evaluation_results_by_student_id_sorted(student_id)
    
    if not evaluations:
        raise HTTPException(status_code=404, detail="暂无评估记录")
    
    # 构建对比响应
    comparison_data = {
        "student_id": student_id,
        "student_name": student.name,
        "evaluations": []
    }
    
    for evaluation in evaluations:
        # 获取维度评分
        dimension_scores = db_service.get_dimension_scores_by_evaluation_id(evaluation.evaluation_id)
        
        # 构建维度评分响应
        dimension_scores_response = [
            {
                "dimension": ds.dimension,
                "score": ds.score,
                "confidence": ds.confidence
            }
            for ds in dimension_scores
        ]
        
        evaluation_data = {
            "evaluation_id": evaluation.evaluation_id,
            "overall_score": evaluation.overall_score,
            "evaluated_at": evaluation.evaluated_at,
            "dimension_scores": dimension_scores_response,
            "strengths": process_string_list(evaluation.strengths),
            "areas_for_improvement": process_string_list(evaluation.areas_for_improvement),
            "recommendations": process_string_list(evaluation.recommendations)
        }
        
        comparison_data["evaluations"].append(evaluation_data)
    
    # 计算评分变化
    if len(evaluations) > 1:
        score_changes = []
        for i in range(1, len(evaluations)):
            prev_evaluation = evaluations[i-1]
            current_evaluation = evaluations[i]
            
            score_change = {
                "from_evaluation": prev_evaluation.evaluation_id,
                "to_evaluation": current_evaluation.evaluation_id,
                "overall_score_change": current_evaluation.overall_score - prev_evaluation.overall_score,
                "dimension_score_changes": []
            }
            
            # 计算维度评分变化
            prev_dimensions = {ds.dimension: ds.score for ds in db_service.get_dimension_scores_by_evaluation_id(prev_evaluation.evaluation_id)}
            current_dimensions = {ds.dimension: ds.score for ds in db_service.get_dimension_scores_by_evaluation_id(current_evaluation.evaluation_id)}
            
            for dimension in set(prev_dimensions.keys()) | set(current_dimensions.keys()):
                prev_score = prev_dimensions.get(dimension, 0)
                current_score = current_dimensions.get(dimension, 0)
                
                score_change["dimension_score_changes"].append({
                    "dimension": dimension,
                    "change": current_score - prev_score
                })
            
            score_changes.append(score_change)
        
        comparison_data["score_changes"] = score_changes
    
    return comparison_data

# AI 配置相关的 Pydantic 模型
class AIConfigRequest(BaseModel):
    provider: str
    api_key: str
    model: str
    base_url: Optional[str] = None
    temperature: float = 0.7
    max_tokens: int = 2000

class AIConfigResponse(BaseModel):
    provider: str
    model: str
    base_url: str
    temperature: float
    max_tokens: int
    has_api_key: bool

class TestAIResponse(BaseModel):
    success: bool
    message: str
    model: Optional[str] = None
    response_time: Optional[float] = None
    error: Optional[str] = None

# 内存中存储的 AI 配置（实际生产环境应该使用更安全的方式）
_current_ai_config: Dict[str, Any] = None

def get_current_ai_config():
    """获取当前 AI 配置"""
    global _current_ai_config
    if _current_ai_config is None:
        # 使用默认配置
        config = get_ai_config()
        _current_ai_config = {
            "provider": settings.ai_provider,
            "api_key": config.get("api_key", ""),
            "model": config.get("model", settings.ai_model),
            "base_url": config.get("base_url", ""),
            "temperature": config.get("temperature", settings.ai_temperature),
            "max_tokens": config.get("max_tokens", settings.ai_max_tokens)
        }
    return _current_ai_config

# AI 配置相关路由
@app.get("/ai-config")
async def get_ai_configuration():
    """获取当前 AI 配置"""
    config = get_current_ai_config()
    return {
        "provider": config["provider"],
        "model": config["model"],
        "base_url": config["base_url"],
        "temperature": config["temperature"],
        "max_tokens": config["max_tokens"],
        "api_key": config["api_key"],  # 返回 API Key 以便前端使用
        "has_api_key": bool(config["api_key"])
    }

@app.post("/ai-config")
async def update_ai_configuration(config: AIConfigRequest):
    """更新 AI 配置"""
    global _current_ai_config
    
    # 验证提供商
    if config.provider not in AI_PROVIDERS and config.provider != "custom":
        raise HTTPException(status_code=400, detail=f"不支持的 AI 提供商: {config.provider}")
    
    # 获取提供商信息
    provider_info = AI_PROVIDERS.get(config.provider, {})
    
    # 确定 base_url
    if config.provider == "custom":
        if not config.base_url:
            raise HTTPException(status_code=400, detail="自定义提供商必须提供 base_url")
        base_url = config.base_url
    else:
        base_url = config.base_url or provider_info.get("base_url", "")
    
    # 保存配置
    _current_ai_config = {
        "provider": config.provider,
        "api_key": config.api_key,
        "model": config.model,
        "base_url": base_url,
        "temperature": config.temperature,
        "max_tokens": config.max_tokens
    }
    
    # 更新环境变量（供其他模块使用）
    os.environ["AI_PROVIDER"] = config.provider
    os.environ["AI_API_KEY"] = config.api_key
    os.environ["AI_MODEL"] = config.model
    os.environ["AI_BASE_URL"] = base_url
    os.environ["AI_TEMPERATURE"] = str(config.temperature)
    os.environ["AI_MAX_TOKENS"] = str(config.max_tokens)
    
    # 同时设置OpenAI的环境变量，确保兼容性
    os.environ["OPENAI_API_KEY"] = config.api_key
    os.environ["OPENAI_MODEL"] = config.model
    if base_url:
        os.environ["OPENAI_BASE_URL"] = base_url
    
    return {"message": "AI 配置已更新", "provider": config.provider, "model": config.model}

@app.post("/ai-config/reset")
async def reset_ai_configuration():
    """重置 AI 配置为默认值"""
    global _current_ai_config
    _current_ai_config = None
    
    # 清除环境变量
    for key in ["AI_PROVIDER", "AI_API_KEY", "AI_MODEL", "AI_BASE_URL", "OPENAI_API_KEY", "OPENAI_MODEL", "OPENAI_BASE_URL"]:
        if key in os.environ:
            del os.environ[key]
    
    return {"message": "AI 配置已重置为默认值"}

@app.post("/ai-config/test", response_model=TestAIResponse)
async def test_ai_connection():
    """测试 AI 连接"""
    config = get_current_ai_config()
    
    if not config["api_key"]:
        return TestAIResponse(
            success=False,
            message="",
            error="未配置 API Key"
        )
    
    try:
        start_time = time.time()
        
        # 根据提供商使用不同的测试方式
        if config["provider"] in ["openai", "deepseek", "zhipu", "moonshot", "qwen", "custom"]:
            # 使用 OpenAI 兼容格式测试
            from openai import OpenAI
            
            client = OpenAI(
                api_key=config["api_key"],
                base_url=config["base_url"]
            )
            
            response = client.chat.completions.create(
                model=config["model"],
                messages=[
                    {"role": "system", "content": "你是一个 helpful assistant."},
                    {"role": "user", "content": "你好，请回复'连接测试成功'"}
                ],
                temperature=config["temperature"],
                max_tokens=config["max_tokens"]
            )
            
            response_time = time.time() - start_time
            message = response.choices[0].message.content
            
            return TestAIResponse(
                success=True,
                message=message,
                model=config["model"],
                response_time=round(response_time, 2)
            )
        else:
            return TestAIResponse(
                success=False,
                message="",
                error=f"不支持的提供商: {config['provider']}"
            )
            
    except Exception as e:
        return TestAIResponse(
            success=False,
            message="",
            error=f"连接测试失败: {str(e)}"
        )

class AnalyzeThesisRequest(BaseModel):
    content: str
    title: str = ""

@app.post("/analyze_thesis_abstract")
async def analyze_thesis_abstract(request: AnalyzeThesisRequest):
    """
    分析论文摘要并检测论文类型（改进版）
    
    改进点：
    1. 提取摘要时保留可读性
    2. 使用混合检测策略（规则+LLM）
    3. 传递完整内容进行更准确的类型判断
    """
    try:
        abstract = extract_abstract(request.content)
        
        if not abstract:
            return {
                "abstract": "",
                "has_abstract": False,
                "project_type": None,
                "message": "未能从论文中提取到摘要"
            }
        
        type_result = detect_project_type_by_llm(abstract, request.title, request.content)
        
        return {
            "abstract": abstract,
            "has_abstract": True,
            "project_type": type_result.get("type", "mixed"),
            "project_type_name": type_result.get("type_name", "混合类"),
            "confidence": type_result.get("confidence", 0.5),
            "reason": type_result.get("reason", ""),
            "features": type_result.get("features", {})
        }
    except Exception as e:
        logger.error(f"分析论文摘要失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"分析论文摘要失败: {str(e)}")

class ExtractChapterRequest(BaseModel):
    content: str
    chapter_name: str

@app.post("/extract_chapter")
async def api_extract_chapter(request: ExtractChapterRequest):
    """
    提取论文指定章节内容
    
    Args:
        content: 论文全文内容
        chapter_name: 章节名称（摘要、引言、绪论、结论、总结、参考文献、致谢）
        
    Returns:
        提取的章节内容
    """
    try:
        chapter_content = extract_chapter(request.content, request.chapter_name)
        
        return {
            "chapter_name": request.chapter_name,
            "content": chapter_content,
            "has_content": len(chapter_content) > 0
        }
    except Exception as e:
        logger.error(f"提取章节内容失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"提取章节内容失败: {str(e)}")

@app.post("/debug_pdf_extraction")
async def debug_pdf_extraction(
    file: UploadFile = File(...)
):
    """
    调试PDF提取 - 查看提取的原始内容和章节识别结果
    
    用于诊断PDF提取问题
    """
    try:
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_path = tmp_file.name
        
        try:
            from src.utils.pdf_extractor_enhanced import EnhancedPDFExtractor
            
            extractor = EnhancedPDFExtractor(enable_ocr=False)
            result = extractor.extract_with_metadata(tmp_path)
            
            text = result["text"]
            
            chapter_markers = re.findall(r'【章节】([^\n]+)', text)
            title_markers = re.findall(r'【标题】([^\n]+)', text)
            
            lines = text.split('\n')
            first_100_lines = '\n'.join(lines[:100])
            
            chapter_patterns = [
                r'第[一二三四五六七八九十\d]+\s*章[^\n]*',
                r'摘\s*要',
                r'ABSTRACT',
                r'结论',
                r'总结',
            ]
            
            found_chapters = []
            for pattern in chapter_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                found_chapters.extend(matches)
            
            return {
                "metadata": result["metadata"],
                "extraction_log": result["log"],
                "chapter_markers_found": chapter_markers[:20],
                "title_markers_found": title_markers[:20],
                "chapters_found_by_pattern": list(set(found_chapters))[:20],
                "first_100_lines": first_100_lines,
                "total_chars": len(text),
                "total_lines": len(lines),
                "sample_content": {
                    "first_500_chars": text[:500],
                    "middle_500_chars": text[len(text)//2:len(text)//2+500] if len(text) > 1000 else "",
                    "last_500_chars": text[-500:] if len(text) > 500 else ""
                }
            }
        finally:
            os.unlink(tmp_path)
            
    except Exception as e:
        logger.error(f"调试PDF提取失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调试PDF提取失败: {str(e)}")

@app.post("/debug_section_extraction")
async def debug_section_extraction(
    request: Dict = Body(...)
):
    """
    调试章节提取 - 查看章节识别和提取结果
    
    用于诊断章节提取问题
    """
    try:
        content = request.get('content', '')
        
        if not content:
            raise HTTPException(status_code=400, detail="内容不能为空")
        
        from src.evaluation.llm_evaluator import llm_evaluator
        from src.evaluation.sectioned_evaluator import SectionedEvaluator
        
        sectioned_evaluator = SectionedEvaluator(llm_evaluator)
        
        structure = sectioned_evaluator.identify_thesis_structure(content)
        
        sections = sectioned_evaluator.extract_sections(content, structure)
        
        sections_debug = []
        for sec in sections:
            sections_debug.append({
                "title": sec.get("title", ""),
                "type": sec.get("section_type", ""),
                "content_length": len(sec.get("content", "")),
                "content_preview": sec.get("content", "")[:500] if sec.get("content") else "",
                "start_marker": sec.get("start_marker", ""),
                "end_marker": sec.get("end_marker", "")
            })
        
        return {
            "structure": structure,
            "sections_debug": sections_debug,
            "total_sections": len(sections),
            "empty_sections": [s["title"] for s in sections_debug if s["content_length"] < 100]
        }
        
    except Exception as e:
        logger.error(f"调试章节提取失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调试章节提取失败: {str(e)}")

# 主入口
@app.delete("/evaluations/{evaluation_id}")
async def delete_evaluation(
    evaluation_id: str,
    db_service: DatabaseService = Depends(get_database_service)
):
    """
    删除评估记录
    """
    try:
        # 删除评估记录
        if db_service.delete_evaluation_result(evaluation_id):
            return {"message": "评估记录删除成功"}
        else:
            raise HTTPException(status_code=404, detail="评估记录不存在")
    except Exception as e:
        logger.error(f"删除评估记录失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除评估记录失败: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host=settings.api_host,
        port=settings.api_port,
        reload=settings.api_debug,
        reload_dirs=[os.path.dirname(os.path.dirname(os.path.abspath(__file__)))]
    )